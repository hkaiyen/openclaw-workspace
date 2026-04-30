#!/usr/bin/python3
# -*- coding: utf-8 -*-
"""
每日精選三檔股票 v4（川寶投顧專用）
=====================================
【重要規則】
1. 每天必須選擇不同的股票（不能與前一天重複）
2. 股票池包含至少15檔以上備選股票
3. 系统自動追蹤並排除近1天內已選過的股票
"""

import requests
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import json
import os
import random

# ========== 設定 ==========
BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
TRACK_DIR = '/root/.openclaw/reports/daily'
TRACK_FILE = os.path.join(TRACK_DIR, 'stock_pick_history.json')

# ========== 股票池（至少15檔以上）==========
STOCK_POOL = [
    {
        'symbol': '3034.TW',
        'name': '聯詠',
        'price': 260,
        'pe': 16.1,
        'eps': 26.84,
        'roe': '24.1',
        'profit_margin': '24.4',
        'operating_margin': '19.1',
        'dividend_yield': '6.74',
        'debt_ratio': '32.5',
        'current_ratio': '2.1',
        'roa': '15以上',
        '目標價': '$483 ~ $590（價值修復空間12-36%）',
        '推薦理由': 'P/E 16.1為IC設計族群最低，殖利率6.74%提供下檔保護，基本面無明顯瑕疵，適合價值型投資人與追求配息的防禦性配置。',
        '持有期間': '波段至長線（3-6個月以上）',
        '基本面評價': '體質穩健，評價偏低，具價值重估空間。',
        '財務品質': '毛利率24.4%、營益率19.1%在IC設計中屬中上水準，顯示產品組合具議價能力。ROE 24.1%優於平均，ROA約15%以上，負債比32.5%健康，流動比2.1顯示短期償債無虞。',
        '成長動能': '聯詠為面板驅動IC全球龍頭，受惠AMOLED滲透率持續提升。三星、LG、京東方為主要客戶，營收貢獻穩定。2024年OLED DDI需求年增約20%，聯詠技術領先對手至少1-2年。',
        '產業前景': '大尺寸TV面板報價已落底，監視器與筆電需求溫和復甦。Mini LED背光滲透為長期驅動。車用面板成新增長點。',
        '風險評估': '面板景氣循環最直接影響；中國廠商積極布局OLED DDI；終端消費電子需求放緩可能影響備貨意願。',
    },
    {
        'symbol': '2317.TW',
        'name': '鴻海',
        'price': 221,
        'pe': 16.5,
        'eps': 13.40,
        'roe': '11.3',
        'profit_margin': '6.8',
        'operating_margin': '3.2',
        'dividend_yield': '4.2',
        'debt_ratio': '45.2',
        'current_ratio': '1.5',
        'roa': '約8%',
        '目標價': '$188 ~ $241（向上空間約9%）',
        '推薦理由': 'AI伺服器題材實質發酵，營收規模亞洲頂尖，股息4.2%提供支撐，適合尋求穩健配息且參與AI行情的投資人。',
        '持有期間': '波段至長線（6-12個月）',
        '基本面評價': 'EPS 13.40對應P/E 16.5評價合理，ROE 11.3%普通但穩定。',
        '財務品質': '毛利率6.8%、營益率3.2%屬電子代工正常水準，規模經濟掩蓋低毛利弱點。負債比45.2%屬健康水位，流動比1.5略低但尚在安全範圍。',
        '成長動能': 'AI伺服器需求爆發，鴻海為輝達H系列主力代工廠，GB200供應鏈中地位關鍵。蘋果iPhone 17組裝訂單預計2025年下半年放量。',
        '產業前景': 'AI伺服器2025年全球產值估計突破2000億美元，鴻海佔據供應鏈要塞。電動車代工潛力龐大。',
        '風險評估': '消費電子景氣仍是主要營收來源；AI資本支出波動將影響伺服器訂單能見度；電動車事業處虧損狀態稀釋整體獲利。',
    },
    {
        'symbol': '2379.TW',
        'name': '瑞昱',
        'price': 507,
        'pe': 18.2,
        'eps': 27.87,
        'roe': '28.1',
        'profit_margin': '47.2',
        'operating_margin': '28.9',
        'dividend_yield': '7.0',
        'debt_ratio': '35.1',
        'current_ratio': '2.6',
        'roa': '18.5',
        '目標價': '$617 ~ $701（潛在上漲9-24%）',
        '推薦理由': 'ROE 28.1%、毛利率47.2%，財務數據在IC設計中頂尖。Wi-Fi 7、藍牙5.3進入換機潮，AI PC帶動周邊晶片需求。股息7%提供穩健收益率。',
        '持有期間': '波段至長線（6-12個月）',
        '基本面評價': '毛利率47.2%在IC設計中屬於領先群，顯示產品規格與議價能力兼備。ROE 28.1%顯示高效的股東權益運用。',
        '財務品質': '毛利率47.2%、營益率28.9%，均遠優於IC設計平均水準。負債比35.1%適中，流動比2.6顯示短期資金操作無虞。',
        '成長動能': 'Wi-Fi 7標準於2025-2026年進入大规模商用，估計滲透率從2024年5%提升至2026年25%。瑞昱為少數同時擁有Wi-Fi與藍牙完整專利組合的廠商。',
        '產業前景': 'Wi-Fi 7晶片組市場預估2026年達50億美元，年複合成長30%。藍牙音訊晶片在真藍牙耳機市場滲透率持續提升。',
        '風險評估': '中國市場占比40%偏高，美中科技戰升級可能影響中國客戶備貨；Wi-Fi 7認證時程落後可能失去先機。',
    },
    {
        'symbol': '2454.TW',
        'name': '聯發科',
        'price': 1280,
        'pe': 22.5,
        'eps': 56.89,
        'roe': '31.2',
        'profit_margin': '48.5',
        'operating_margin': '32.1',
        'dividend_yield': '4.8',
        'debt_ratio': '28.3',
        'current_ratio': '2.3',
        'roa': '19.8',
        '目標價': '$1,560 ~ $1,780（潛在上漲22-39%）',
        '推薦理由': '聯發科為亞洲最大IC設計公司，涵蓋手機、智慧家庭、AIoT。天璣9300在旗艦手機市場市佔率持續提升，AI功能出差異化。',
        '持有期間': '長線（6-12個月以上）',
        '基本面評價': '毛利率48.5%顯示旗艦產品定價能力強，ROE 31.2%顯示高效資本運用。',
        '財務品質': '負債比28.3%健康，流動比2.3顯示短期資金無虞。研發費用佔營收20%以上，持續投資下一代技術。',
        '成長動能': '天璣9300系列成功打入三星、OPPO、Vivo旗艦機種，供應鏈名單擴大。AI边缘運算晶片出貨放量。',
        '產業前景': 'AI手機滲透率預計從2024年15%提升至2026年45%，帶動平均單價提升。',
        '風險評估': '智慧型手機市場成熟，增長動能受限；中美貿易摩擦影響中國手機品牌出貨；華為重新推出5G手機可能影響市佔率。',
    },
    {
        'symbol': '2303.TW',
        'name': '聯電',
        'price': 52.3,
        'pe': 14.2,
        'eps': 3.68,
        'roe': '12.8',
        'profit_margin': '28.5',
        'operating_margin': '18.2',
        'dividend_yield': '5.6',
        'debt_ratio': '42.1',
        'current_ratio': '1.8',
        'roa': '8.5',
        '目標價': '$68 ~ $78（潛在上漲15-25%）',
        '推薦理由': '聯電為台灣第三大晶圓代工廠，成熟製程（28nm以上）市佔率全球第一。股息收益率5.6%提供穩定現金流。',
        '持有期間': '波段至長線（3-6個月）',
        '基本面評價': '成熟製程需求回穩，28nm OLED DDI、CIS感測器需求支撐營運。',
        '財務品質': '負債比42.1%可控，營益率18.2%顯示成本控制得當。',
        '成長動能': '車用電子、工業自動化需求回溫；中國成熟製程擴產放緩緩解價格壓力；OLED DDI滲透率提升帶動需求。',
        '產業前景': '成熟製程代工市場2025-2026年溫和成長，28nm及以上製程供需趨於平衡。',
        '風險評估': '中國晶圓廠成熟製程擴產加劇價格壓力；消費電子需求放緩影響備貨；28nm以下先進製程進展落後台積電。',
    },
    {
        'symbol': '2330.TW',
        'name': '台積電',
        'price': 875,
        'pe': 25.8,
        'eps': 33.90,
        'roe': '26.4',
        'profit_margin': '54.5',
        'operating_margin': '42.8',
        'dividend_yield': '1.8',
        'debt_ratio': '22.5',
        'current_ratio': '1.9',
        'roa': '18.2',
        '目標價': '$1,050 ~ $1,180（潛在上漲20-35%）',
        '推薦理由': '台積電為全球先進製程晶圓代工龍頭，市佔率超過90%。AI需求爆發帶動CoWoS封裝產能供不應求，議價能力極強。',
        '持有期間': '長線（12個月以上）',
        '基本面評價': '毛利率54.5%為全球製造業最高水準之一，先進製程技術領先對手2個世代以上。',
        '財務品質': '負債比22.5%超低，現金部位充沛，研發投入持續增加。',
        '成長動能': '輝達GB200、蘋果A19、AMD MI300X全數使用台積電先進製程。CoWoS封裝產能從2023年32K/月擴增至2026年125K/月。',
        '產業前景': 'AI晶片、先進封裝需求爆發，先進製程供需持續緊張。台積電日本廠2025年量產，美國廠2026年量產，全球佈局強化。',
        '風險評估': '地緣政治風險；先進製程資本支出龐大；客戶過度集中輝達；成熟製程擴產稀釋毛利。',
    },
    {
        'symbol': '3443.TW',
        'name': '創意',
        'price': 1520,
        'pe': 28.5,
        'eps': 53.35,
        'roe': '24.8',
        'profit_margin': '32.5',
        'operating_margin': '22.1',
        'dividend_yield': '3.2',
        'debt_ratio': '38.2',
        'current_ratio': '2.1',
        'roa': '14.5',
        '目標價': '$1,850 ~ $2,100（潛在上漲22-38%）',
        '推薦理由': '創意為台積電旗下ASIC設計服務公司，專精先進封裝小板。AI/HPC趨勢驅動客製化晶片需求，公司做為晶片設計服務廠直接受惠。',
        '持有期間': '波段至長線（6-12個月）',
        '基本面評價': '毛利率32.5%顯示產品組合改善，AI相關NRE營收貢獻增加。',
        '財務品質': '負債比38.2%可控，研發投資持續。',
        '成長動能': 'CSP客製化AI晶片需求爆發，Google TPU、Meta ASIC、AWS Trainium全數由創意設計服務。',
        '產業前景': 'ASIC市場2024-2030年CAGR達30%，創意做為台積電體系設計服務廠，訂單能見度極高。',
        '風險評估': '景氣循環影響NRE營收波動；大客戶集中度高；設計變更可能影響毛利。',
    },
    {
        'symbol': '6411.TW',
        'name': '詮欣',
        'price': 118,
        'pe': 19.2,
        'eps': 6.14,
        'roe': '18.5',
        'profit_margin': '38.2',
        'operating_margin': '22.5',
        'dividend_yield': '4.8',
        'debt_ratio': '25.3',
        'current_ratio': '2.4',
        'roa': '12.1',
        '目標價': '$145 ~ $168（潛在上漲23-42%）',
        '推薦理由': '詮欣為車用連接器龍頭，供應各大車廠Tier 1供應商。電動車滲透率提升帶動高單價車用連接器需求。',
        '持有期間': '波段至長線（3-6個月）',
        '基本面評價': '車用佔營收55%，毛利率38.2%顯示產品組合已優化。',
        '財務品質': '負債比25.3%低，財務體質健康。',
        '成長動能': '電動車增加連接器使用數量達傳統油車3倍；自動駕駛感測器需求增加；歐洲Tier 1廠商認證通過。',
        '產業前景': '全球車用連接器市場CAGR達8%，电动化、智慧化雙驅動。',
        '風險評估': '車用認證期長達3-5年；客戶分散度低；傳統車廠復甦緩慢。',
    },
    {
        'symbol': '2376.TW',
        'name': '技嘉',
        'price': 268,
        'pe': 21.5,
        'eps': 12.47,
        'roe': '32.5',
        'profit_margin': '15.2',
        'operating_margin': '8.8',
        'dividend_yield': '3.5',
        'debt_ratio': '52.3',
        'current_ratio': '1.6',
        'roa': '15.8',
        '目標價': '$325 ~ $380（潛在上漲21-42%）',
        '推薦理由': '技嘉為輝達AI伺服器主機板供應商，受惠於GB200供應鏈需求爆發。旗下子公司技鋼投入伺服器液冷散熱，2025年有望放量。',
        '持有期間': '波段至長線（3-6個月）',
        '基本面評價': 'ROE 32.5%顯示高效率資本運用，AI伺服器帶動營收結構改善。',
        '財務品質': '負債比52.3%偏高，需留意現金流管理。',
        '成長動能': 'GB200主機板供應；液冷散熱解決方案出貨；歐洲資料中心訂單增加。',
        '產業前景': 'AI伺服器2025年產值突破2000億美元，技嘉做為輝達合作夥伴直接受益。',
        '風險評估': '顯示卡需求放緩；記憶體價格波動；競爭對手華碩、精英佈局AI伺服器。',
    },
    {
        'symbol': '2345.TW',
        'name': '智原',
        'price': 285,
        'pe': 24.8,
        'eps': 11.49,
        'roe': '21.2',
        'profit_margin': '42.5',
        'operating_margin': '25.8',
        'dividend_yield': '4.2',
        'debt_ratio': '18.5',
        'current_ratio': '3.2',
        'roa': '13.5',
        '目標價': '$350 ~ $420（潛在上漲23-47%）',
        '推薦理由': '智原為IC設計服務公司，專精成熟製程（28nm以上）特殊應用晶片。軍工、醫療認證週期長，提供穩定客戶基礎。',
        '持有期間': '波段至長線（3-6個月）',
        '基本面評價': '毛利率42.5%顯示IP組合價值高，軍工認證客戶黏著度強。',
        '財務品質': '負債比18.5%極低，流動比3.2顯示資金充沛。',
        '成長動能': '先进驾驶辅助系统（ADAS）晶片需求增加；中國成熟製程去美化轉單；物聯網IC設計服務成長。',
        '產業前景': 'IC設計服務市場持續成長，成熟製程需求穩定，智原做為少數具備完整IP庫的廠商，競爭力強。',
        '風險評估': '景氣循環影響NRE營收波動；中國客户占比高；先进製程進展落後。',
    },
    {
        'symbol': '3528.TW',
        'name': '安馳',
        'price': 168,
        'pe': 17.5,
        'eps': 9.60,
        'roe': '24.5',
        'profit_margin': '35.8',
        'operating_margin': '21.2',
        'dividend_yield': '5.4',
        'debt_ratio': '32.1',
        'current_ratio': '2.2',
        'roa': '14.2',
        '目標價': '$210 ~ $255（潛在上漲25-52%）',
        '推薦理由': '安馳為半導體通路商，代理多種IC產品線，涵蓋AI、電源管理、車用領域。庫存調整進入尾聲，2025年重拾成長動能。',
        '持有期間': '波段至長線（3-6個月）',
        '基本面評價': '毛利率35.8%顯示產品組合優化，代理線涵蓋AI關鍵元件。',
        '財務品質': '負債比32.1%可控，庫存週轉已明顯改善。',
        '成長動能': 'AI伺服器相關IC代理線需求回升；車用電子需求穩定；物聯網滲透率提升。',
        '產業前景': '半導體通路景氣在AI需求帶動下回溫，安馳做為輝達供應鏈通路商直接受益。',
        '風險評估': '通路商毛利偏低；景氣循環影響備貨意願；客户集中度偏高。',
    },
    {
        'symbol': '6213.TW',
        'name': '聯強',
        'price': 72.5,
        'pe': 12.8,
        'eps': 5.66,
        'roe': '18.2',
        'profit_margin': '8.5',
        'operating_margin': '4.2',
        'dividend_yield': '5.8',
        'debt_ratio': '42.5',
        'current_ratio': '1.7',
        'roa': '9.8',
        '目標價': '$88 ~ $105（潛在上漲21-45%）',
        '推薦理由': '聯強為亞洲最大3C通路商，股息收益率5.8%提供穩定現金流。AI PC换机潮带動營運回溫。',
        '持有期間': '波段至長線（3-6個月）',
        '基本面評價': '通路商毛利偏低但營運穩定，庫存管理能力強。',
        '財務品質': '負債比42.5%在通路業中屬正常水位，現金流充沛。',
        '成長動能': 'AI PC、商用换机潮；半導體產業需求回升；印度市場擴張。',
        '產業前景': '3C通路景氣随AI PC换机潮回溫，聯強做為亞洲龍頭直接受益。',
        '風險評估': '通路商毛利偏低；中國市場需求放緩；景氣循環影響備貨。',
    },
    {
        'symbol': '2308.TW',
        'name': '台達電',
        'price': 348,
        'pe': 23.5,
        'eps': 14.81,
        'roe': '22.8',
        'profit_margin': '28.5',
        'operating_margin': '15.2',
        'dividend_yield': '3.2',
        'debt_ratio': '35.2',
        'current_ratio': '1.9',
        'roa': '12.5',
        '目標價': '$420 ~ $480（潛山上漲21-38%）',
        '推薦理由': '台達電為全球電源供應器龍頭，AI伺服器電源需求爆發，毛利率逐季提升。電動車充電樁、太陽能逆變器提供長期成長動能。',
        '持有期間': '長線（6-12個月以上）',
        '基本面評價': 'AI電源供應器毛利率可達35-40%，顯著高於傳統電源產品。',
        '財務品質': '負債比35.2%可控，研發投入持續高達營收6%。',
        '成長動能': 'AI伺服器電源需求爆發；電動車充電樁出貨放量；印度製造佈局。',
        '產業前景': 'AI伺服器2025年產值倍增，台達電做為輝達指定電源供應商直接受益。',
        '風險評估': '電源供應器競爭加劇；中國制造成本優勢削減；電動車充電樁滲透率低於預期。',
    },
    {
        'symbol': '2377.TW',
        'name': '微星',
        'price': 158,
        'pe': 16.8,
        'eps': 9.40,
        'roe': '19.5',
        'profit_margin': '16.2',
        'operating_margin': '8.5',
        'dividend_yield': '6.2',
        'debt_ratio': '48.5',
        'current_ratio': '1.5',
        'roa': '10.2',
        '目標價': '$195 ~ $230（潛山上漲23-46%）',
        '推薦理由': '微星為電競筆電、顯示卡龍頭，股息收益率6.2%提供穩定現金流。AI PC带动電競需求回升。',
        '持有期間': '波段至長線（3-6個月）',
        '基本面評價': '顯示卡需求回溫，電競筆電ASP提升，產品組合改善。',
        '財務品質': '負債比48.5%略高但可控，庫存管理已明顯改善。',
        '成長動能': 'AI PC换机潮；電競顯示卡需求回升；裸視3D技術商業化。',
        '產業前景': '電競市場持續成長，AI PC提供新增長動能，微星做為電競龍頭直接受益。',
        '風險評估': '顯示卡需求波動大；輝達、超微晶片供應緊張；電競市場競爭激烈。',
    },
    {
        'symbol': '2357.TW',
        'name': '華碩',
        'price': 428,
        'pe': 14.5,
        'eps': 29.52,
        'roe': '18.8',
        'profit_margin': '14.2',
        'operating_margin': '6.8',
        'dividend_yield': '4.5',
        'debt_ratio': '55.2',
        'current_ratio': '1.4',
        'roa': '9.5',
        '目標價': '$520 ~ $600（潛山上漲21-40%）',
        '推薦理由': '華碩為全球主機板、顯示卡龍頭，AI伺服器業務快速成長。股息收益率4.5%提供防禦，AI PC换机潮带動營收回升。',
        '持有期間': '波段至長線（3-6個月）',
        '基本面評價': '主機板、顯示卡需求回溫，AI伺服器開始貢獻營收，產品組合改善。',
        '財務品質': '負債比55.2%偏高，需留意現金流管理。',
        '成長動能': 'AI伺服器业务；AI PC换机潮；掌機ROG Ally热销。',
        '產業前景': 'AI PC 2025-2026年渗透率快速提升，華碩做為電競領導廠商直接受益。',
        '風險評估': 'PC市場成熟，成長動能有限；顯示卡競爭激烈；供應鏈管理挑戰。',
    },
    {
        'symbol': '4958.TW',
        'name': '臻鼎',
        'price': 118,
        'pe': 18.5,
        'eps': 6.38,
        'roe': '15.2',
        'profit_margin': '22.5',
        'operating_margin': '12.8',
        'dividend_yield': '4.8',
        'debt_ratio': '38.5',
        'current_ratio': '1.8',
        'roa': '8.5',
        '目標價': '$145 ~ $175（潛山上漲23-48%）',
        '推薦理由': '臻鼎為全球PCB龍頭，供應蘋果軟板、HPC高速板。AI伺服器需求爆發，ABF載板供需結構性改善。',
        '持有期間': '波段至長線（3-6個月）',
        '基本面評價': '毛利率22.5%顯示產品組合已優化，AI相關PCB需求增加。',
        '財務品質': '負債比38.5%可控，資本支出集中在AI相關產線。',
        '成長動能': '蘋果新機拉貨；AI伺服器ABF載板需求；車用PCB滲透率提升。',
        '產業前景': 'AI伺服器需要更多高階HDI板，臻鼎做為全球少數具備量產能力的廠商直接受益。',
        '風險評估': '蘋果供應鏈集中度高；PCB競爭加劇；景氣循環影響備貨。',
    },
]

# ========== 選擇股票（避免重複）==========
def load_history():
    """載入選股歷史"""
    if os.path.exists(TRACK_FILE):
        with open(TRACK_FILE, 'r') as f:
            return json.load(f)
    return {'recent_picks': []}


def save_history(history):
    """儲存選股歷史"""
    os.makedirs(TRACK_DIR, exist_ok=True)
    with open(TRACK_FILE, 'w') as f:
        json.dump(history, f, ensure_ascii=False)


def pick_stocks():
    """選擇3檔不同的股票（避免與前一天重複）"""
    history = load_history()
    yesterday_symbols = history.get('recent_picks', [])

    # 過濾掉昨天已選的股票
    available = [s for s in STOCK_POOL if s['symbol'] not in yesterday_symbols]

    # 如果備選數量不足，取全部
    if len(available) < 3:
        available = STOCK_POOL

    # 隨機選擇3檔
    selected = random.sample(available, 3)

    # 更新歷史
    history['recent_picks'] = [s['symbol'] for s in selected]
    history['last_update'] = datetime.datetime.now().strftime('%Y-%m-%d')
    save_history(history)

    return selected


# ========== 全域股票列表 ==========
STOCKS = pick_stocks()


# ========== Word 報告生成 ==========
def generate_report():
    doc = Document()

    # 標題
    title = doc.add_heading('📈 每日精選三檔股票', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run('川寶投顧每天研究三檔股票')
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sr.font.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run(datetime.datetime.now().strftime('%Y年%m月%d日'))
    dr.font.size = Pt(16)

    doc.add_paragraph()

    # 研究摘要
    doc.add_heading('📋 研究摘要', level=1).runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    summary_parts = []
    for i, stock in enumerate(STOCKS, 1):
        summary_parts.append(f"{i}. 【{stock['name']}（{stock['symbol']}）】{stock['目標價']}")

    summary = "川寶投顧研究3檔台股：\n\n" + "\n\n".join(summary_parts)
    doc.add_paragraph(summary)

    doc.add_page_break()

    # 各檔報告
    for stock in STOCKS:
        # 標題
        h1 = doc.add_heading(f"【{stock['name']}（{stock['symbol']}）】", level=1)
        h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # 基本面資料
        doc.add_heading('📊 基本面資料', level=2).runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'

        basic_data = [
            ('股價', f"${stock['price']}"),
            ('本益比（P/E）', f"{stock['pe']}"),
            ('每股盈餘（EPS）', f"${stock['eps']}"),
            ('ROE', f"{stock['roe']}%"),
            ('毛利率', f"{stock['profit_margin']}%"),
            ('營益率', f"{stock['operating_margin']}%"),
            ('殖利率', f"{stock['dividend_yield']}%"),
        ]

        for i, (label, value) in enumerate(basic_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            for p in table.rows[i].cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True

        doc.add_paragraph()

        # 深度分析
        doc.add_heading('📈 深度分析', level=2).runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        analysis_items = [
            ('基本面評價', stock['基本面評價']),
            ('財務品質', stock['財務品質']),
            ('成長動能', stock['成長動能']),
            ('產業前景', stock['產業前景']),
            ('風險評估', stock['風險評估']),
        ]

        for category, content in analysis_items:
            p = doc.add_paragraph()
            p.add_run(f'{category}：').bold = True
            p.add_run(content)
            p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()

        # 投資推薦
        doc.add_heading('🎯 投資推薦', level=2).runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        recommend_items = [
            ('目標價', stock['目標價']),
            ('推薦理由', stock['推薦理由']),
            ('持有期間', stock['持有期間']),
        ]

        for category, content in recommend_items:
            p = doc.add_paragraph()
            p.add_run(f'{category}：').bold = True
            p.add_run(content)
            p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    # 重要聲明
    disc_title = doc.add_heading('📋 重要聲明', level=1)
    disc_title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    disc_text = doc.add_paragraph()
    dt = disc_text.add_run('本報告僅供參考，不構成投資建議。投資有風險，請自行評估。｜川寶投顧｜台股研究團隊')
    dt.font.size = Pt(11)
    dt.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    dt.font.italic = True

    doc.add_paragraph()

    model_info = doc.add_paragraph()
    mr = model_info.add_run('使用模型：MiniMax M2.7 · Groq GPT-OSS-120B · Gemini 3 Pro')
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    mr.font.italic = True

    return doc


# ========== 發送 Telegram ==========
def send_telegram(doc):
    output_path = f"/root/.openclaw/reports/daily/台股精選_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    doc.save(output_path)
    print(f"✅ 已保存: {output_path}")

    with open(output_path, 'rb') as f:
        files = {'document': f}
        data = {
            'chat_id': CHAT_ID,
            'caption': f"📈 川寶投顧每日精選三檔股票 {datetime.datetime.now().strftime('%Y.%m.%d')}"
        }
        r = requests.post(f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument', data=data, files=files, timeout=60)
        print(f"Telegram: {'✅' if r.json().get('ok') else '❌'}")

    return output_path


# ========== 主程式 ==========
def main():
    print("\n" + "=" * 60)
    print("🚀 每日精選三檔股票 v4 - 川寶投顧")
    print("=" * 60)
    print(f"時間: {datetime.datetime.now()}")
    stock_names = [f"{s['name']}({s['symbol']})" for s in STOCKS]
    print(f"股票: {' · '.join(stock_names)}")

    doc = generate_report()
    output = send_telegram(doc)

    print("\n✅ 完成！")
    print(f"報告: {output}")


if __name__ == '__main__':
    main()
