def upload_to_website(file_path):
    """將報告上傳到網站分類（自動分類 + 自動更新首頁）"""
    import subprocess
    import glob
    from docx import Document
    
    try:
        doc = Document(file_path)
        today = datetime.datetime.now()
        file_name = os.path.basename(file_path)
        
        # === 1. 自動分類 ===
        category = "研究"  # 預設
        if "資產報酬率" in file_name or "股市" in file_name or "股票" in file_name or "台股" in file_name:
            category = "股市"
        elif "房地產" in file_name or "房價" in file_name or "房市" in file_name:
            category = "房地產"
        elif "促銷" in file_name or "活動" in file_name or "優惠" in file_name:
            category = "促銷"
        
        # === 2. 產生Markdown ===
        md_filename = f"{file_name.replace('.docx', '')}_{today.strftime('%Y%m%d_%H%M')}.md"
        md_path = f"/root/.openclaw/workspace/reports_site/docs/reports/{category}/{md_filename}"
        
        # 讀取docx轉換為markdown（統一格式）
        md_content = []
        md_content.append(f"# 📊 {file_name.replace('.docx', '')}\n")
        md_content.append(f"**報告日期：** {today.strftime('%Y年%m月%d日')}\n")
        md_content.append(f"**分類：** {category}\n")
        md_content.append("---\n\n")
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                md_content.append(f"{text}\n\n")
        
        for table in doc.tables:
            rows_data = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(" | ".join(cells))
            md_content.append("\n| " + " | ".join(["---"] * len(table.columns)) + " |\n")
            for row in rows_data:
                md_content.append(f"| {row} |\n")
            md_content.append("\n")
        
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write("".join(md_content))
        
        print(f"   📤 已上傳到：{category}/{md_filename}")
        
        # === 3. 自動更新首頁 index.md ===
        update_index_page()
        
        # === 4. Git commit and push ===
        try:
            subprocess.run(['git', 'add', '.'], cwd='/root/.openclaw/workspace/reports_site', check=True, capture_output=True)
            subprocess.run(['git', 'commit', '-m', f'自動更新：{category}報告 {today.strftime("%Y%m%d %H:%M")}'], 
                         cwd='/root/.openclaw/workspace/reports_site', check=True, capture_output=True)
            subprocess.run(['git', 'push', 'origin', 'master'], 
                          cwd='/root/.openclaw/workspace/reports_site', check=True, capture_output=True, timeout=30)
            print("   ✅ 已推送到 GitHub，網站將自動更新")
            return True
        except Exception as e:
            print(f"   ⚠️ Git推送失敗: {e}")
            return False
            
    except Exception as e:
        print(f"   ⚠️ 上傳失敗: {e}")
        return False


def update_index_page():
    """自動更新首頁的最新報告連結"""
    categories = {
        "股市": "📈 股市與投資",
        "房地產": "🏠 房地產",
        "促銷": "🛒 促銷活動",
        "研究": "📚 研究報告"
    }
    
    index_content = ["<style>\n.md-content p, .md-content, .md-typeset p, .md-typeset {\n    font-size: 22px !important;\n    line-height: 1.8 !important;\n}\n.md-typeset h1, .md-content h1 { font-size: 1.6em !important; font-weight: bold !important; }\n.md-typeset h2, .md-content h2 { font-size: 1.4em !important; font-weight: bold !important; }\n.md-typeset h3, .md-content h3 { font-size: 1.5em !important; font-weight: bold !important; }\n.md-typeset table, .md-typeset td, .md-typeset th { font-size: 22px !important; }\n.md-typeset a, .md-content a { font-size: 22px !important; }\n.md-nav__link, .md-nav__item, .md-nav__link--active { font-size: 24px !important; }\n.md-typeset ul, .md-typeset ol { font-size: 22px !important; }\n</style>\n\n"]
    index_content.append("# 川寶每日報告\n\n")
    index_content.append("歡迎使用川寶團隊每日報告網站！\n\n")
    index_content.append("## 📊 最新報告\n\n")
    
    # 遍歷每個分類，取得最新3份報告
    for cat_key, cat_name in categories.items():
        cat_dir = f"/root/.openclaw/workspace/reports_site/docs/reports/{cat_key}"
        index_content.append(f"### {cat_name}\n")
        
        if os.path.exists(cat_dir):
            # 取得所有md檔案，按修改時間排序，取最新3個
            md_files = glob.glob(f"{cat_dir}/*.md")
            md_files.sort(key=os.path.getmtime, reverse=True)
            
            for md_file in md_files[:3]:
                if os.path.basename(md_file) != "index.md":
                    title = os.path.basename(md_file).replace('.md', '')
                    link = f"reports/{cat_key}/{os.path.basename(md_file)}"
                    index_content.append(f"- [{title}]({link})\n")
        index_content.append("\n")
    
    index_content.append("---\n\n")
    index_content.append("## 🔄 自動更新\n\n")
    index_content.append("本網站每日自動更新，報告會在生成後自動同步到這裡。\n\n")
    index_content.append("## 📝 報告分類\n\n")
    index_content.append("- **📈 股市與投資** - 股市分析、資產報酬率、股票策略\n")
    index_content.append("- **🏠 房地產** - 房價指數、房地產趨勢\n")
    index_content.append("- **🛒 促銷活動** - 最新優惠、母親節檔期\n")
    index_content.append("- **📚 研究報告** - 深度研究、趨勢分析\n")
    
    # 寫入index.md
    index_path = "/root/.openclaw/workspace/reports_site/docs/index.md"
    with open(index_path, 'w', encoding='utf-8') as f:
        f.write("".join(index_content))
    
    print("   📝 已更新首頁 index.md")
