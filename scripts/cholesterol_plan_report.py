#!/usr/bin/python3
"""三個月降膽固醇飲食運動生活計劃報告"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os, requests

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
BLUE = RGBColor(0x1F, 0x49, 0x7D)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)
ORANGE = RGBColor(0xFF, 0x66, 0x00)

def set_cell(cell, text, bold=False, color=None):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_header(doc, text, level=1, color=BLUE):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = color
    return h

def add_bold_para(doc, text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True

def add_point(doc, text):
    p = doc.add_paragraph()
    p.add_run('• ' + text)

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/三個月降膽固醇計劃_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ===== 封面 =====
    h0 = doc.add_heading('🏥 三個月降膽固醇計劃', 0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h0.runs[0].font.size = Pt(28)
    h0.runs[0].font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('飲食 · 運動 · 生活作息完整指南')
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('製定日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(11)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== 免責聲明 =====
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr2 = disc.add_run('⚠️ 本計劃僅供參考，執行前請諮詢醫師意見')
    dr2.font.size = Pt(11)
    dr2.font.color.rgb = RED
    dr2.bold = True

    doc.add_paragraph()

    # ===== 第一章：計劃概覽 =====
    add_header(doc, '第一章：計劃概覽')

    table0 = doc.add_table(rows=5, cols=2)
    table0.style = 'Table Grid'
    headers0 = ['項目', '內容']
    for i, h in enumerate(headers0):
        table0.rows[0].cells[i].text = h
        table0.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    overview = [
        ('計劃名稱', '三個月降膽固醇健康計劃'),
        ('目標', 'LDL（壞膽固醇）下降 20%'),
        ('預期總膽固醇下降', '15-20%'),
        ('預期時間', '12週（約三個月）'),
    ]
    for i, (k, v) in enumerate(overview):
        table0.rows[i+1].cells[0].text = k
        set_cell(table0.rows[i+1].cells[1], v, True)

    doc.add_paragraph()

    # ===== 第二章：飲食計劃 =====
    add_header(doc, '第二章：第一個月 - 飲食調整計劃', color=GREEN)

    add_bold_para(doc, '一、每日飲食原則')

    diet_table = doc.add_table(rows=6, cols=3)
    diet_table.style = 'Table Grid'
    headers_diet = ['食物類別', '建議', '避免']
    for i, h in enumerate(headers_diet):
        diet_table.rows[0].cells[i].text = h
        diet_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    diet_data = [
        ('蛋白質', '魚（鮭魚、鮪魚）、雞胸肉、豆類', '紅肉、內臟、加工肉品'),
        ('蔬菜', '每天 300g 以上', '-'),
        ('全穀類', '糙米、燕麥、藜麥', '白米、白麵包'),
        ('堅果', '每天一小把（無鹽）', '鹽烤堅果'),
        ('油脂', '橄欖油、酪梨油', '豬油、牛油、椰子油'),
    ]
    for i, row in enumerate(diet_data):
        for j, val in enumerate(row):
            diet_table.rows[i+1].cells[j].text = val

    doc.add_paragraph()
    add_bold_para(doc, '二、每週執行原則')

    weekly_points = [
        '禮拜一～五：無肉日（以豆類+魚代替肉類）',
        '禮拜六日：可以吃一次瘦肉（巴掌大小）',
        '每天：一杯燕麥片 + 10顆杏仁',
        '每天喝水量：2000-2500ml',
        '避免：甜食、含糖飲料、酒精',
    ]
    for p in weekly_points:
        add_point(doc, p)

    doc.add_paragraph()
    add_bold_para(doc, '三、範例一日三餐')

    meal_table = doc.add_table(rows=4, cols=2)
    meal_table.style = 'Table Grid'
    meal_table.rows[0].cells[0].text = '餐次'
    meal_table.rows[0].cells[1].text = '建議'
    for c in meal_table.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True

    meals = [
        ('早餐', '燕麥片 + 藍莓 + 無糖豆漿'),
        ('午餐', '烤鮭魚 + 糙米 + 大量蔬菜'),
        ('晚餐', '豆腐味噌湯 + 涼拌花椰菜'),
    ]
    for i, (m, food) in enumerate(meals):
        meal_table.rows[i+1].cells[0].text = m
        meal_table.rows[i+1].cells[1].text = food

    doc.add_paragraph()

    # ===== 第三章：運動計劃 =====
    add_header(doc, '第三章：第二個月 - 運動計劃', color=ORANGE)

    add_bold_para(doc, '一、每週運動目標')

    exercise_table = doc.add_table(rows=4, cols=3)
    exercise_table.style = 'Table Grid'
    ex_headers = ['運動類型', '頻率', '時間']
    for i, h in enumerate(ex_headers):
        exercise_table.rows[0].cells[i].text = h
        exercise_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    ex_data = [
        ('快走/慢跑', '每週 5 天', '每次 30 分鐘'),
        ('阻力訓練', '每週 2 天', '每次 20 分鐘'),
        ('伸展', '每天', '10 分鐘'),
    ]
    for i, row in enumerate(ex_data):
        for j, val in enumerate(row):
            exercise_table.rows[i+1].cells[j].text = val

    doc.add_paragraph()
    add_bold_para(doc, '二、8週進度表')

    progress_table = doc.add_table(rows=5, cols=2)
    progress_table.style = 'Table Grid'
    progress_table.rows[0].cells[0].text = '週數'
    progress_table.rows[0].cells[1].text = '目標'
    for c in progress_table.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True

    progress = [
        ('第 5 週', '快走 30 分鐘（可說話程度）'),
        ('第 6 週', '快走 30 分鐘 + 啞鈴訓練'),
        ('第 7 週', '慢跑 5 分鐘 + 快走 25 分鐘'),
        ('第 8 週', '慢跑 10 分鐘 + 快走 20 分鐘'),
    ]
    for i, (w, t) in enumerate(progress):
        progress_table.rows[i+1].cells[0].text = w
        progress_table.rows[i+1].cells[1].text = t

    doc.add_paragraph()
    add_bold_para(doc, '三、詳細運動處方')

    add_point(doc, '【快走/慢跑】每週5次')
    p = doc.add_paragraph()
    p.add_run('暖身 5 分鐘 → 快走/慢跑 25-30 分鐘 → 緩和 5 分鐘')
    p.paragraph_format.left_indent = Cm(0.5)

    p2 = doc.add_paragraph()
    p2.add_run('強度：可以說完整句子但有點喘').italic = True
    p2.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()
    add_point(doc, '【阻力訓練】每週2次')
    p3 = doc.add_paragraph()
    p3.add_run('• 深蹲 15下 × 3組\n• 橋式 15下 × 3組\n• 啞鈴划船 12下 × 3組')
    p3.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()

    # ===== 第四章：生活作息 =====
    add_header(doc, '第四章：第三個月 - 生活作息優化', color=GREEN)

    add_bold_para(doc, '一、睡眠優化')

    sleep_table = doc.add_table(rows=4, cols=2)
    sleep_table.style = 'Table Grid'
    sleep_table.rows[0].cells[0].text = '項目'
    sleep_table.rows[0].cells[1].text = '目標'
    for c in sleep_table.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True

    sleep_data = [
        ('睡眠時間', '22:30 前就寢'),
        ('睡眠時數', '7-8 小時'),
        ('睡前藍光', '睡前 30 分鐘不看手機'),
    ]
    for i, (k, v) in enumerate(sleep_data):
        sleep_table.rows[i+1].cells[0].text = k
        sleep_table.rows[i+1].cells[1].text = v

    doc.add_paragraph()
    add_bold_para(doc, '二、壓力管理')

    stress_points = [
        '每天：冥想 5-10 分鐘',
        '每週：至少一天完全休息',
        '每天：深呼吸練習（4-7-8原則）',
    ]
    for p in stress_points:
        add_point(doc, p)

    doc.add_paragraph()
    add_bold_para(doc, '三、每日習慣檢查表')

    habit_items = [
        '☐ 早上：喝溫水 300ml',
        '☐ 早餐：燕麥 + 水果',
        '☐ 午餐前：喝一杯水',
        '☐ 下午：快走 15 分鐘',
        '☐ 晚餐：蔬菜佔一半',
        '☐ 睡前：冥想 5 分鐘',
        '☐ 睡覺：22:30 前',
    ]
    for item in habit_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.3)

    doc.add_paragraph()

    # ===== 第五章：進度追蹤 =====
    add_header(doc, '第五章：三個月進度追蹤')

    track_table = doc.add_table(rows=6, cols=3)
    track_table.style = 'Table Grid'
    track_headers = ['週數', '里程碑', '目標']
    for i, h in enumerate(track_headers):
        track_table.rows[0].cells[i].text = h
        track_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    track_data = [
        ('第 2 週', '飲食調整', '開始執行健康飲食'),
        ('第 4 週', '初次見效', '感覺精神變好'),
        ('第 6 週', '運動養成', '每週運動 3 次以上'),
        ('第 8 週', '中期檢測', '預估 LDL 下降 10-15%'),
        ('第 12 週', '最終檢測', '總膽固醇下降 15-20%'),
    ]
    for i, row in enumerate(track_data):
        for j, val in enumerate(row):
            track_table.rows[i+1].cells[j].text = val

    doc.add_paragraph()

    # ===== 第六章：預期效果 =====
    add_header(doc, '第六章：預期效果')

    effect_table = doc.add_table(rows=6, cols=3)
    effect_table.style = 'Table Grid'
    effect_headers = ['指標', '預期改善', '說明']
    for i, h in enumerate(effect_headers):
        effect_table.rows[0].cells[i].text = h
        effect_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    effect_data = [
        ('總膽固醇', '↓ 15-20%', '明顯改善'),
        ('LDL（壞膽固醇）', '↓ 20-25%', '核心目標'),
        ('HDL（好膽固醇）', '↑ 5-10%', '需要運動配合'),
        ('三酸甘油酯', '↓ 20-30%', '飲食控制有效'),
        ('體重', '↓ 3-5 kg', '健康減重'),
    ]
    for i, (k, v, note) in enumerate(effect_data):
        effect_table.rows[i+1].cells[0].text = k
        set_cell(effect_table.rows[i+1].cells[1], v, True, GREEN)
        effect_table.rows[i+1].cells[2].text = note

    doc.add_paragraph()

    # ===== 第七章：重要提醒 =====
    add_header(doc, '第七章：重要提醒', level=1, color=RED)

    warnings = [
        '⚠️ 每 4 週回診：驗血追蹤膽固醇數值',
        '⚠️ 如在服用藥物：飲食運動不能取代藥物，請勿擅自停藥',
        '⚠️ 不要快速減重：每週減 0.5-1kg 為宜，急速減重可能影響肌肉量',
        '⚠️ 執行前建議先去醫院驗血，建立基準線',
        '⚠️ 如有不適症狀（頭暈、胸悶、呼吸困難）請立即就醫',
    ]
    for w in warnings:
        p = doc.add_paragraph(w)
        p.paragraph_format.left_indent = Cm(0.3)

    doc.add_paragraph()

    # ===== 第八章：科學依據 =====
    add_header(doc, '第八章：科學依據')

    science_points = [
        '📚 美國心臟協會研究顯示：改變飲食與運動可在 12 週內降低 LDL 15-30%',
        '📚 地中海飲食（魚、橄欖油、堅果）被證實可有效降低 LDL',
        '📚 每週 150 分鐘中等強度運動可提升 HDL 5-10%',
        '📚 燕麥中的 β-聚葡萄糖可降低 LDL 5-10%',
        '📚 深海魚類的 Omega-3 脂肪酸有助於降低三酸甘油酯',
    ]
    for s in science_points:
        p = doc.add_paragraph(s)
        p.paragraph_format.left_indent = Cm(0.3)

    doc.add_paragraph()

    # ===== 頁尾 =====
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 三個月降膽固醇計劃')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '🏥 三個月降膽固醇計劃\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '📋 計劃重點：\n'
        '• 第一個月：飲食調整（減少飽和脂肪）\n'
        '• 第二個月：運動計劃（每週150分鐘）\n'
        '• 第三個月：生活作息優化\n\n'
        '📊 預期效果：\n'
        '• LDL 下降 20-25%\n'
        '• 總膽固醇下降 15-20%\n'
        '• 三酸甘油酯下降 20-30%\n\n'
        '⚠️ 執行前請先諮詢醫師意見\n\n'
        '小安製 ❤️'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生三個月降膽固醇計劃報告...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')