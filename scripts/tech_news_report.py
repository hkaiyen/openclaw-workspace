#!/usr/bin/python3
"""
科技新知每日快報
收集 AI、半導體、科技巨頭、新創趨勢資訊，翻譯成繁體中文
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import datetime
import requests
import re
import urllib.parse
import json
import time
import os

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'

BLUE = RGBColor(0x1F, 0x49, 0x7D)
ORANGE = RGBColor(0xFF, 0x66, 0x00)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)

# ===== RSS 來源 =====
RSS_SOURCES = {
    'BBC Tech': 'https://feeds.bbci.co.uk/news/technology/rss.xml',
    'TechCrunch': 'https://techcrunch.com/feed/',
    'Ars Technica': 'https://feeds.arstechnica.com/arstechnica/index',
}

# ===== 翻譯函數（小歐 Groq GPT-OSS-120B）=====
def translate_to_chinese(text):
    if not text or len(text) < 2:
        return text
    api_key = os.environ.get('GROQ_API_KEY', '')
    if not api_key:
        return text
    
    payload = {
        "model": "openai/gpt-oss-120b",
        "messages": [
            {"role": "system", "content": "你是小歐，專業股票翻譯分析師。將以下英文翻譯成繁體中文，只回傳翻譯結果，不要加任何解釋。"},
            {"role": "user", "content": text[:500]}
        ],
        "temperature": 0.3,
        "max_tokens": 200
    }
    
    for attempt in range(3):  # 重試3次
        try:
            req = urllib.request.Request(
                "https://api.groq.com/openai/v1/chat/completions",
                data=json.dumps(payload).encode(),
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                    "User-Agent": "curl/8.0"
                },
                method="POST"
            )
            with urllib.request.urlopen(req, timeout=20) as resp:
                data = json.loads(resp.read().decode())
                result = data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
                if result and len(result) > 3 and result != text:
                    return result
        except Exception as e:
            if attempt == 2:
                print(f"  翻譯錯誤（{attempt+1}次）: {text[:30]}...")
    return text  # 失敗時回傳原文

# ===== 工具函數 =====
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def make_header_cell(cell, text, bg='1F497D'):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, bg)

def add_data_cell(cell, text):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# ===== 抓取 RSS 新聞 =====
def fetch_news(url, source_name, max_items=12):
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
            'Accept': 'application/rss+xml, application/xml, text/xml, */*'
        }
        resp = requests.get(url, headers=headers, timeout=12)
        resp.encoding = 'utf-8'
        
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(resp.text, 'html.parser')
        items = soup.find_all('item')
        
        news_list = []
        for item in items[:max_items]:
            title_elem = item.find('title')
            title_text = title_elem.get_text(strip=True) if title_elem else ""
            title_text = re.sub(r'<!\[CDATA\[|\]\]>', '', title_text)
            link_elem = item.find('link')
            link_text = link_elem.get_text(strip=True) if link_elem else ""
            
            if title_text and len(title_text) > 5:
                news_list.append({
                    'title': title_text,
                    'link': link_text,
                    'source': source_name
                })
        return news_list
    except Exception as e:
        print(f'  RSS 錯誤 {source_name}: {e}')
        return []

# ===== 關鍵字分類 =====
CATEGORIES = {
    '🤖 AI / 人工智慧': ['AI', 'artificial intelligence', 'machine learning', 'ChatGPT', 'OpenAI', 'Anthropic', 'deep learning', 'neural', 'LLM', '生成式', '人工智慧', '機器學習'],
    '💻 半導體 / 晶片': ['chip', 'semiconductor', 'GPU', 'Nvidia', 'NVIDIA', 'TSMC', 'Intel', 'AMD', 'Qualcomm', 'processor', 'silicon', '摩爾', '晶片'],
    '🍎 蘋果 / Google / 微軟': ['Apple', 'Google', 'Microsoft', 'Meta', 'Amazon', 'iPhone', 'Android', 'Windows', 'Alphabet', 'Tim Cook', 'Sundar', 'Satya'],
    '🚗 電動車 / 新能源': ['Tesla', 'EV', 'electric vehicle', 'battery', 'renewable', 'solar', 'wind', 'energy', '電動車', '特斯拉', '綠能'],
    '🪙 區塊鏈 / Web3': ['Bitcoin', 'crypto', 'blockchain', 'Web3', 'NFT', 'Ethereum', 'blockchain', 'web3', '加密貨幣'],
    '🚀 新創 / 創投': ['startup', 'funding', 'venture', 'IPO', 'acquisition', 'billion', 'million', ' Series ', ' Series A', 'Series B', 'unicorn', '新創', '獨角獸'],
    '🔐 資安 / 隱私': ['security', 'hack', 'breach', 'privacy', 'cyber', 'ransomware', 'vulnerability', '資安', '駭客', '隱私'],
    '📱 產品 / 裝置': ['iPad', 'MacBook', 'Apple Watch', 'Pixel', 'Samsung', 'Galaxy', 'surface', 'VR', 'AR', 'mixed reality', ' headset'],
}

def categorize(news_item):
    title = news_item.get('title', '') + news_item.get('body', '')
    matched = []
    for category, keywords in CATEGORIES.items():
        for kw in keywords:
            if kw.lower() in title.lower():
                matched.append(category)
                break
    return matched if matched else ['🌐 一般科技']

# ===== 主程式 =====
def main():
    today = datetime.datetime.now()
    print(f"[{today.strftime('%Y-%m-%d %H:%M:%S')}] 開始生成科技新知快報...")
    
    # 收集所有新聞
    all_news = []
    for source_name, url in RSS_SOURCES.items():
        print(f"  抓取 {source_name}...")
        news = fetch_news(url, source_name, max_items=12)
        all_news.extend(news)
        time.sleep(1)
    
    # 去重
    seen = set()
    unique_news = []
    for item in all_news:
        key = item['title'][:40]
        if key not in seen:
            seen.add(key)
            unique_news.append(item)
    
    print(f"  收集到 {len(unique_news)} 則新聞")
    
    # 分類
    categorized = {cat: [] for cat in CATEGORIES.keys()}
    categorized['🌐 一般科技'] = []
    
    for item in unique_news:
        cats = categorize(item)
        for cat in cats:
            if cat in categorized and len(categorized[cat]) < 8:
                if item not in categorized[cat]:
                    categorized[cat].append(item)
    
    # 小歐翻譯
    print(f"  小歐翻譯中...")
    for cat in categorized:
        for item in categorized[cat]:
            if '_zh' not in item:
                item['_zh'] = translate_to_chinese(item['title'])
                time.sleep(0.5)
    
    # 生成報告
    output_path = f"/root/.openclaw/reports/daily/科技新知_{today.strftime('%Y%m%d_%H%M')}.docx"
    
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # 標題
    title = doc.add_heading('🚀 科技新知每日快報', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"📅 {today.strftime('%Y年%m月%d日')} | 🤖 AI · 💻 晶片 · 🍎 科技巨頭 · 🚀 新創")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    
    doc.add_paragraph()
    
    # 各分類
    for cat, items in categorized.items():
        if not items:
            continue
        
        h = doc.add_heading(cat, level=1)
        h.runs[0].font.color.rgb = BLUE
        h.runs[0].font.size = Pt(14)
        
        for item in items[:8]:
            p = doc.add_paragraph()
            title_text = item.get('_zh', item['title'])
            p.add_run(f"• {title_text}")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.3)
            
            source_run = p.add_run(f"  [{item['source']}]")
            source_run.font.size = Pt(9)
            source_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        doc.add_paragraph()
    
    # 資料來源
    h_src = doc.add_heading('📋 資料來源', level=1)
    h_src.runs[0].font.color.rgb = BLUE
    for src in RSS_SOURCES:
        p = doc.add_paragraph(f"• {src}")
    
    # 免責
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disc.add_run('【聲明】本報告資訊僅供參考，版權歸各來源所有')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    dr.font.italic = True
    
    doc.save(output_path)
    print(f"  ✅ 已儲存: {output_path}")
    
    # 發送到 Telegram（小歐翻譯）
    caption = f"🚀 科技新知_{today.strftime('%Y年%m月%d日 %H:%M')}\n\n共 {len(unique_news)} 則資訊\n🤖 AI · 💻 晶片 · 🍎 巨頭 · 🚀 新創\n\n小歐翻譯"
    result = subprocess.run([
        'curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
        '-F', f'chat_id={CHAT_ID}',
        '-F', f'document=@{output_path}',
        '-F', f'caption={caption}'
    ], capture_output=True)
    
    try:
        res = json.loads(result.stdout)
        if res.get('ok'):
            print("  ✅ 已發送到 Telegram")
        else:
            print(f"  ⚠️ Telegram 發送失敗: {res.get('description')}")
    except:
        print("  ⚠️ 無法解析回應")
    
    print(f"[{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] 完成！")

if __name__ == '__main__':
    main()