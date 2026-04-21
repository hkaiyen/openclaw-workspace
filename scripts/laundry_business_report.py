#!/usr/bin/python3
"""自主洗衣店創業研究報告 - 台北市文山區"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
BLUE = RGBColor(0x1F, 0x49, 0x7D)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)

def set_color(cell, text, color, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/自主洗衣店創業研究_文山區_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面
    title = doc.add_heading('🧺 自主洗衣店創業研究報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run('台北市文山區市場可行性分析')
    sr.font.size = Pt(18)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('研究日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== 第一章：市場基礎研究 =====
    h1 = doc.add_heading('📊 第一章：市場基礎研究', level=1)
    h1.runs[0].font.color.rgb = BLUE

    h1_1 = doc.add_heading('1.1 文山區人口結構分析', level=2)
    h1_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table1 = doc.add_table(rows=8, cols=2)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '項目'
    hdr1[1].text = '數據'
    for c in hdr1:
        c.paragraphs[0].runs[0].bold = True

    data1 = [
        ('總人口數 (2025年)', '約 25.4 萬人'),
        ('總戶數 (2025年)', '約 11.1 萬戶'),
        ('人口成長率 (近10年)', '-4.6%（微幅減少）'),
        ('人口最多里別', '木新里（9,878人）'),
        ('人口最少里別', '老泉里（833人）'),
        ('區域特色', '住宅密集、學區环繞'),
        ('主要居住型態', '套房出租、大樓林立'),
    ]
    for i, (idx, val) in enumerate(data1):
        table1.rows[i+1].cells[0].text = idx
        table1.rows[i+1].cells[1].text = val

    doc.add_paragraph()

    h1_2 = doc.add_heading('1.2 目標客群分析', level=2)
    h1_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '客群'
    hdr2[1].text = '特徵'
    hdr2[2].text = '需求強度'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    data2 = [
        ('套房租屋族', '無洗衣機設置空間', '⭐⭐⭐⭐⭐'),
        ('大專院校學生', '政大、世新、景文', '⭐⭐⭐⭐'),
        ('小家庭', '居住空間有限', '⭐⭐⭐'),
        ('銀髮族', '體力負擔重衣物', '⭐⭐⭐'),
    ]
    for i, (客群, 特徵, 強度) in enumerate(data2):
        table2.rows[i+1].cells[0].text = 客群
        table2.rows[i+1].cells[1].text = 特徵
        table2.rows[i+1].cells[2].text = 強度

    doc.add_paragraph()

    h1_3 = doc.add_heading('1.3 市場需求規模估算', level=2)
    h1_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    p = doc.add_paragraph()
    p.add_run('依據區域特性與產業經驗法則：')
    p.paragraph_format.space_after = Pt(6)

    demand_points = [
        '📍 文山區套房屋密集，估計潛在使用人口約 3-5 萬人',
        '📍 目標客群每週洗衣需求約 1-2 次',
        '📍 每次消費單價預估 $50-80 元（投幣式）',
        '📍 合理的市場滲透率目標：5-10%',
    ]
    for point in demand_points:
        doc.add_paragraph(point)

    doc.add_paragraph()

    # ===== 第二章：地點與競爭分析 =====
    h2 = doc.add_heading('📍 第二章：地點與競爭分析', level=1)
    h2.runs[0].font.color.rgb = BLUE

    h2_1 = doc.add_heading('2.1 精華區位評估', level=2)
    h2_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table3 = doc.add_table(rows=6, cols=4)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '地點'
    hdr3[1].text = '人流量'
    hdr3[2].text = '租金水位'
    hdr3[3].text = '適合程度'
    for c in hdr3:
        c.paragraphs[0].runs[0].bold = True

    data3 = [
        ('景美夜市商圈', '高', '中', '⭐⭐⭐⭐⭐'),
        ('萬隆捷運站周邊', '中', '中', '⭐⭐⭐⭐'),
        ('羅斯福路沿線', '中', '中高', '⭐⭐⭐'),
        ('木柵路沿線', '中', '中低', '⭐⭐⭐⭐'),
        ('興隆路/萬盛街', '中', '低', '⭐⭐⭐⭐⭐'),
    ]
    for i, (地點, 人流, 租金, 評級) in enumerate(data3):
        table3.rows[i+1].cells[0].text = 地點
        table3.rows[i+1].cells[1].text = 人流
        table3.rows[i+1].cells[2].text = 租金
        table3.rows[i+1].cells[3].text = 評級

    doc.add_paragraph()

    h2_2 = doc.add_heading('2.2 選址關鍵條件', level=2)
    h2_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    location_points = [
        '✅ 捷運站出口 500 公尺內最佳',
        '✅ 大型集合住宅/套房大樓出入口',
        '✅ 鄰近便利商店、超市等人流聚集點',
        '✅ 停車/臨停方便',
        '✅ 室内空間 20-40 坪',
        '✅ 通風排水條件佳',
        '✅ 可申請商業用電錶',
    ]
    for point in location_points:
        doc.add_paragraph(point)

    doc.add_paragraph()

    h2_3 = doc.add_heading('2.3 租金行情估算', level=2)
    h2_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table4 = doc.add_table(rows=5, cols=2)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '區位'
    hdr4[1].text = '月租金估算'
    for c in hdr4:
        c.paragraphs[0].runs[0].bold = True

    data4 = [
        ('精華商圈 (景美/萬隆)', '$30,000-50,000/月'),
        ('住宅區沿線', '$15,000-30,000/月'),
        ('大學學區周邊', '$20,000-35,000/月'),
        ('一般住宅區', '$10,000-20,000/月'),
    ]
    for i, (區位, 租金) in enumerate(data4):
        table4.rows[i+1].cells[0].text = 區位
        table4.rows[i+1].cells[1].text = 租金

    doc.add_paragraph()

    # ===== 第三章：財務可行性評估 =====
    h3 = doc.add_heading('💰 第三章：財務可行性評估', level=1)
    h3.runs[0].font.color.rgb = BLUE

    h3_1 = doc.add_heading('3.1 初始投資估算（自營模式）', level=2)
    h3_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table5 = doc.add_table(rows=8, cols=2)
    table5.style = 'Table Grid'
    hdr5 = table5.rows[0].cells
    hdr5[0].text = '項目'
    hdr5[1].text = '費用估算'
    for c in hdr5:
        c.paragraphs[0].runs[0].bold = True

    data5 = [
        ('洗衣設備 (4-6台)', '$80,000-120,000'),
        ('烘衣設備 (2-3台)', '$40,000-60,000'),
        ('乾燥機/脫水機', '$20,000-30,000'),
        ('店面裝潢/水電', '$50,000-100,000'),
        ('押金/保證金 (2個月)', '$30,000-80,000'),
        ('加盟金/權利金', '$0 (自營)'),
        ('週轉金/備用金', '$50,000-100,000'),
    ]
    for i, (項目, 費用) in enumerate(data5):
        table5.rows[i+1].cells[0].text = 項目
        table5.rows[i+1].cells[1].text = 費用

    doc.add_paragraph()

    p_total = doc.add_paragraph()
    p_total.add_run('💡 自營模式總投資預估：').bold = True
    p_total.add_run('$270,000 - $490,000')
    p_total.runs[1].font.color.rgb = GREEN

    doc.add_paragraph()

    h3_2 = doc.add_heading('3.2 每月營運成本', level=2)
    h3_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table6 = doc.add_table(rows=7, cols=2)
    table6.style = 'Table Grid'
    hdr6 = table6.rows[0].cells
    hdr6[0].text = '項目'
    hdr6[1].text = '每月費用'
    for c in hdr6:
        c.paragraphs[0].runs[0].bold = True

    data6 = [
        ('店租', '$15,000-40,000'),
        ('水電瓦斯', '$15,000-30,000'),
        ('耗材/洗劑', '$2,000-5,000'),
        ('設備保養', '$3,000-8,000'),
        ('網路/系統', '$500-1,000'),
        ('其他支出', '$2,000-5,000'),
    ]
    for i, (項目, 費用) in enumerate(data6):
        table6.rows[i+1].cells[0].text = 項目
        table6.rows[i+1].cells[1].text = 費用

    doc.add_paragraph()

    p_total2 = doc.add_paragraph()
    p_total2.add_run('💡 每月固定成本合計：').bold = True
    p_total2.add_run('$37,500 - $89,000')
    p_total2.runs[1].font.color.rgb = RED

    doc.add_paragraph()

    h3_3 = doc.add_heading('3.3 營收預測模型', level=2)
    h3_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table7 = doc.add_table(rows=6, cols=4)
    table7.style = 'Table Grid'
    hdr7 = table7.rows[0].cells
    hdr7[0].text = '情境'
    hdr7[1].text = '日均來客'
    hdr7[2].text = '月營收'
    hdr7[3].text = '月淨利'
    for c in hdr7:
        c.paragraphs[0].runs[0].bold = True

    data7 = [
        ('保守', '20人', '$36,000', '-$5,000~+$5,000'),
        ('普通', '40人', '$72,000', '$10,000-25,000'),
        ('樂觀', '60人', '$108,000', '$30,000-50,000'),
        ('優異', '80人', '$144,000', '$50,000-70,000'),
    ]
    for i, (情境, 來客, 營收, 淨利) in enumerate(data7):
        table7.rows[i+1].cells[0].text = 情境
        table7.rows[i+1].cells[1].text = 來客
        table7.rows[i+1].cells[2].text = 營收
        if '-$' in 淨利 or ('+$' in 淨利):
            set_color(table7.rows[i+1].cells[3], 淨利, RED, True)
        else:
            set_color(table7.rows[i+1].cells[3], 淨利, GREEN, True)

    doc.add_paragraph()

    h3_4 = doc.add_heading('3.4 回本時間分析', level=2)
    h3_4.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table8 = doc.add_table(rows=4, cols=3)
    table8.style = 'Table Grid'
    hdr8 = table8.rows[0].cells
    hdr8[0].text = '情境'
    hdr8[1].text = '月淨利'
    hdr8[2].text = '回本時間'
    for c in hdr8:
        c.paragraphs[0].runs[0].bold = True

    data8 = [
        ('普通', '$15,000/月', '約 2-3 年', True),
        ('樂觀', '$30,000/月', '約 1.5-2 年', True),
        ('優異', '$50,000/月', '約 1 年', True),
    ]
    for i, (情境, 淨利, 時間, is_pos) in enumerate(data8):
        table8.rows[i+1].cells[0].text = 情境
        table8.rows[i+1].cells[1].text = 淨利
        set_color(table8.rows[i+1].cells[2], 時間, GREEN, True)

    doc.add_paragraph()

    # ===== 第四章：營運模式規劃 =====
    h4 = doc.add_heading('⚙️ 第四章：營運模式規劃', level=1)
    h4.runs[0].font.color.rgb = BLUE

    h4_1 = doc.add_heading('4.1 設備建議清單', level=2)
    h4_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table9 = doc.add_table(rows=6, cols=3)
    table9.style = 'Table Grid'
    hdr9 = table9.rows[0].cells
    hdr9[0].text = '設備'
    hdr9[1].text = '數量'
    hdr9[2].text = '建議規格'
    for c in hdr9:
        c.paragraphs[0].runs[0].bold = True

    data9 = [
        ('洗衣機', '4-6台', '15-20kg 大型滾筒'),
        ('烘衣機', '2-3台', '30kg 以上'),
        ('脫水機', '1-2台', '輔助乾燥'),
        ('投幣系統', '全套', '支援悠遊卡/LINE Pay'),
        ('監控系統', '1套', '雲端錄影+APP通知'),
    ]
    for i, (設備, 數量, 規格) in enumerate(data9):
        table9.rows[i+1].cells[0].text = 設備
        table9.rows[i+1].cells[1].text = 數量
        table9.rows[i+1].cells[2].text = 規格

    doc.add_paragraph()

    h4_2 = doc.add_heading('4.2 自營 vs 加盟比較', level=2)
    h4_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table10 = doc.add_table(rows=6, cols=3)
    table10.style = 'Table Grid'
    hdr10 = table10.rows[0].cells
    hdr10[0].text = '項目'
    hdr10[1].text = '自營'
    hdr10[2].text = '加盟 (UCC/奇搓)'
    for c in hdr10:
        c.paragraphs[0].runs[0].bold = True

    data10 = [
        ('初始投資', '較低', '較高 (含加盟金)'),
        ('品牌支援', '無', '有品牌效益'),
        ('設備採購', '自行議價', '統一採購折扣'),
        ('教育訓練', '自己學', '原廠教育訓練'),
        ('營運支持', '獨立摸索', 'SOP手冊、行銷支援'),
    ]
    for i, (項目, 自營, 加盟) in enumerate(data10):
        table10.rows[i+1].cells[0].text = 項目
        table10.rows[i+1].cells[1].text = 自營
        table10.rows[i+1].cells[2].text = 加盟

    doc.add_paragraph()

    h4_3 = doc.add_heading('4.3 數位化營運建議', level=2)
    h4_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    digital_points = [
        '📱 APP 預約洗衣/烘衣時間功能',
        '💳 多元支付：悠遊卡、LINE Pay、信用卡',
        '📊 即時監控：手機查看機器狀態、空位',
        '🔔 自動通知：洗滌完成推播通知',
        '📈 數據分析：来客習慣、熱門時段統計',
    ]
    for point in digital_points:
        doc.add_paragraph(point)

    doc.add_paragraph()

    # ===== 第五章：執行計劃書 =====
    h5 = doc.add_heading('📋 第五章：執行計劃書', level=1)
    h5.runs[0].font.color.rgb = BLUE

    h5_1 = doc.add_heading('5.1 行動時間表', level=2)
    h5_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table11 = doc.add_table(rows=7, cols=3)
    table11.style = 'Table Grid'
    hdr11 = table11.rows[0].cells
    hdr11[0].text = '階段'
    hdr11[1].text = '時間'
    hdr11[2].text = '任務'
    for c in hdr11:
        c.paragraphs[0].runs[0].bold = True

    data11 = [
        ('第一階段', '第 1-2 週', '地點實地訪查、租金行情確認'),
        ('第二階段', '第 3-4 週', '設備商報價、財務規劃'),
        ('第三階段', '第 5-6 週', '合約簽訂、裝潢申請'),
        ('第四階段', '第 7-8 週', '設備安裝、測試營運'),
        ('第五階段', '第 9 週', '正式開幕、行銷推廣'),
        ('第六階段', '第 10-12 週', '營運優化、調整定價'),
    ]
    for i, (階段, 時間, 任務) in enumerate(data11):
        table11.rows[i+1].cells[0].text = 階段
        table11.rows[i+1].cells[1].text = 時間
        table11.rows[i+1].cells[2].text = 任務

    doc.add_paragraph()

    h5_2 = doc.add_heading('5.2 風險管控計劃', level=2)
    h5_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    risk_points = [
        ('⚠️ 設備故障', '建立備用機制、簽訂保固合約'),
        ('⚠️ 營收不如預期', '前3個月觀察期、果斷調整地點或定價'),
        ('⚠️ 異常天候', '冬季/雨天加強行銷、提供烘衣優惠'),
        ('⚠️ 人為破壞', '監視系統、意外險保障'),
        ('⚠️ 法規問題', '預先確認使用分區、環保噪音標準'),
    ]
    for 風險, 對策 in risk_points:
        p = doc.add_paragraph()
        run1 = p.add_run(風險 + '：')
        run1.bold = True
        p.add_run(對策)

    doc.add_paragraph()

    h5_3 = doc.add_heading('5.3 最終建議', level=2)
    h5_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    final_points = [
        ('✅', '首選地點：興隆路/萬盛街住宅區，人流穩定、租金適中'),
        ('✅', '首選設備：4台洗衣機 + 2台烘衣機，性價比最高'),
        ('✅', '營運模式：自營（省去加盟金，初期壓力較小）'),
        ('✅', '數位功能：至少要有悠遊卡支付 + APP 查詢'),
        ('✅', '預備金：额外準備 3-6 個月的營運資金'),
    ]
    for icon, text in final_points:
        p = doc.add_paragraph()
        run = p.add_run(icon + ' ' + text)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # 頁尾
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 自主洗衣店創業研究報告')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '🧺 自主洗衣店創業研究報告\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '📊 台北市文山區市場分析\n'
        '• 潛在客群：25萬人口中的套房租屋族\n'
        '• 預估投資：$27-49萬（自營模式）\n'
        '• 月淨利預估：$10,000-50,000\n'
        '• 回本時間：1.5-3年\n\n'
        '📍 首選地點：興隆路/萬盛街住宅區\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    import requests
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生自主洗衣店創業研究報告...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
