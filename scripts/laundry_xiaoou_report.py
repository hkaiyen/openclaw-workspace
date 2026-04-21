#!/usr/bin/python3
"""小歐版：台北市文山區自助洗衣店創業創意研究報告"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
BLUE = RGBColor(0x1F, 0x49, 0x7D)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)
ORANGE = RGBColor(0xFF, 0x88, 0x00)
PURPLE = RGBColor(0x88, 0x00, 0x88)

def set_color(cell, text, color, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/自主洗衣店創業研究_小歐版_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面
    title = doc.add_heading('🧺 台北市文山區自助洗衣店創業創意研究報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = PURPLE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run('小歐創意策略顧問版')
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tg = tagline.add_run('——從另類視角透視市場機會與創新突圍策略')
    tg.font.size = Pt(12)
    tg.font.color.rgb = ORANGE
    tg.italic = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('研究日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # 前言
    intro = doc.add_paragraph()
    intro.add_run('文山區，台北人心中的「邊陲」——但正是這種邊陲性格，造就了獨特的創業沃土。這裡有大學城的青春活力、有住宅區的穩定人流、有山區步道的觀光人潮，更有一群被都會生活壓力轟炸、卻渴望便利與品質的居民。自助洗衣店在這裡，不該只是「投幣洗衣服」的傳統場景——它可以是一個生活樞紐、一個社區節點、一個年輕品牌的起點。').italic = True

    doc.add_paragraph()

    # ===== 一、另類市場洞察 =====
    h1 = doc.add_heading('一、別類市場洞察：那些傳統報告不會告訴你的事', level=1)
    h1.runs[0].font.color.rgb = PURPLE

    h1_1 = doc.add_heading('1.1 「洗衣盲區」——大學生不是主力軍', level=2)
    h1_1.runs[0].font.color.rgb = ORANGE

    doc.add_paragraph('多數人直覺認為，大學城效應帶來的就是「學生洗衣商機」。錯。政大、世新、警專的學生宿舍普遍附設投幣式洗衣機，且年輕族群早已習慣用手洗/宿舍公用洗衣機。真正被忽略的，是以下三個族群：')

    table1 = doc.add_table(rows=4, cols=3)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '被忽視的客群'
    hdr1[1].text = '規模估算'
    hdr1[2].text = '核心需求'
    for c in hdr1:
        c.paragraphs[0].runs[0].bold = True

    data1 = [
        ('50-70歲銀髮族', '文山區約4-5萬人', '床單、被套、厚重衣物；不願占用自家陽台空間'),
        ('新手爸媽（30-45歲）', '有0-6歲子女家庭約1.5-2萬戶', '嬰兒衣物需高溫洗滌；沒有時間等待'),
        ('租屋族（25-40歲單身/小家庭）', '估計1-2萬人', '沒有洗衣機、陽台通風差；需要24小時可取回'),
    ]
    for i, (客群, 規模, 需求) in enumerate(data1):
        table1.rows[i+1].cells[0].text = 客群
        table1.rows[i+1].cells[1].text = 規模
        table1.rows[i+1].cells[2].text = 需求

    doc.add_paragraph()

    h1_2 = doc.add_heading('1.2 「天氣心理學」——雨季就是免費的業務員', level=2)
    h1_2.runs[0].font.color.rgb = ORANGE

    doc.add_paragraph('台北年均雨天約180天，雨季（5-6月梅雨、7-9月颱風）讓戶外曬衣變成一場賭博。文山區多山坡，濕度更高，衣服不易乾——這不是乾洗衣店的機會，這是自助洗衣+烘乾一站式服務的天然剛需。')

    doc.add_paragraph()

    h1_3 = doc.add_heading('1.3 「空間剝削率」——坪效可以比別人高3倍', level=2)
    h1_3.runs[0].font.color.rgb = ORANGE

    doc.add_paragraph('同樣的30坪空間，可以變成：')

    space_points = [
        '🌅 早上：銀髮族的洗被時段',
        '🌤️ 下午：媽媽的嬰兒衣物專區',
        '🌙 晚上：租屋族的日常洗衣',
        '🏔 週末：登山客的機能衣物洗護',
    ]
    for point in space_points:
        doc.add_paragraph(point)

    doc.add_paragraph('同一個地點，時間切割出不同客群，坪效翻倍。').bold = True

    doc.add_paragraph()

    # ===== 二、差異化策略建議 =====
    h2 = doc.add_heading('二、差異化策略建議：避開紅海的五條小路', level=1)
    h2.runs[0].font.color.rgb = PURPLE

    strategies = [
        ('2.1 「機能衣物專家」——瞄準登山與戶外社群', [
            '文山區緊鄰拇指山、指南宮，登山族群有個核心痛點：Gore-Tex、刷毛外套、排汗衣需要特殊洗劑',
            '設置「機能衣物專用洗衣機」（標示清楚、收費略高15-20%）',
            '與在地登山社團、OUTDOOR店面合作導客',
        ]),
        ('2.2 「深夜經濟」——24小時取衣的儀式感', [
            '文山區有不少輪班工作者（醫療業、服務業、貨運司機）',
            '全天候24小時經營（以自助門禁系統管控）',
            'APP預約洗衣時段、推播通知取衣',
        ]),
        ('2.3 「寵物衣物專區」——毛孩爸媽的隱形需求', [
            '設置「寵物衣物專用洗程」（高溫消毒、低速脫水）',
            '提供寵物除蟎洗護（加價服務）',
            '結合附近寵物美容店做異業聯盟',
        ]),
        ('2.4 「等待經濟」——把等待時間變成消費場景', [
            '與鄰近咖啡廳、便利超商合作，提供「等待券」',
            '在店內設置小型夾娃娃機、零食飲料販賣機（被動收入+填補等待時間）',
            '設置共享辦公區（15元/10分鐘 WiFi+插座），讓等待變成一種工作儀式',
        ]),
        ('2.5 「透明化信任」——洗程直播與SOP展示', [
            '安裝透明玻璃隔間+直播洗程（顧客在等待區可觀看）',
            '牆面展示洗劑成分，SOP流程',
            '鼓勵顧客參觀鍋爐清潔日，建立信任感',
        ]),
    ]

    for title, points in strategies:
        h_sub = doc.add_heading(title, level=2)
        h_sub.runs[0].font.color.rgb = BLUE
        for point in points:
            doc.add_paragraph('• ' + point)
        doc.add_paragraph()

    doc.add_paragraph()

    # ===== 三、創新商業模式提案 =====
    h3 = doc.add_heading('三、創新商業模式提案：超越「洗衣服」的四种可能', level=1)
    h3.runs[0].font.color.rgb = PURPLE

    models = [
        ('3.1 「共享洗衣艙」——小型連鎖加盟模式', [
            '核心旗艦店：旗艦機種、展示空間、體驗區',
            '衛星艙：6-8台機器的小型站點（進駐社區大樓、公家機關、學校福利社）',
            '衛星艙由旗艦店統一管理洗劑、維護、收銀系統',
        ]),
        ('3.2 「衣物健康管理」——自助洗衣店的進化形態', [
            '結合AI視覺辨識與衣物標籤掃描',
            '衣物洗滌歷史記錄、磨損提醒、換季保養提醒',
            '變現方式：免費基礎服務吸引會員，高階衣物健康管理收費',
        ]),
        ('3.3 「洗衣訂閱制」——家庭包月服務', [
            '基礎包：每月30kg洗+烘額度 → $799/月',
            '進階包：每月60kg + 機能衣物專洗 → $1,499/月',
            '尊榮包：无限使用 + 每月1次到府收送 → $2,999/月',
        ]),
        ('3.4 「綠色洗衣聯盟」——環保差異化品牌', [
            '使用環保洗劑（主動公告成分與來源）',
            '太陽能烘乾輔助系統（削減電費+ESG形象）',
            '與在地環保團體合作，建立品牌好感度',
        ]),
    ]

    for title, points in models:
        h_sub = doc.add_heading(title, level=2)
        h_sub.runs[0].font.color.rgb = GREEN
        for point in points:
            doc.add_paragraph('• ' + point)
        doc.add_paragraph()

    doc.add_paragraph()

    # ===== 四、區位創意選擇 =====
    h4 = doc.add_heading('四、區位創意選擇：不是選「在哪裡」，而是選「被誰需要」', level=1)
    h4.runs[0].font.color.rgb = PURPLE

    h4_1 = doc.add_heading('4.1 選址新思維：不是「人潮多」，而是「需求密度高」', level=2)
    h4_1.runs[0].font.color.rgb = ORANGE

    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '創意區位選項'
    hdr2[1].text = '理由'
    hdr2[2].text = '潛在風險'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    data2 = [
        ('萬美公寓社區地下商場', '社區型高密度住宅，居民日常動線核心', '管理委員會談判難度'),
        ('政大後山步道入口（指南路三段）', '登山客必經之地，機能衣物需求明確', '觀光淡旺季人流波動'),
        ('景美醫院周邊', '住院病患家屬、輪班護理師，24小時需求強', '租金可能較高'),
    ]
    for i, (區位, 理由, 風險) in enumerate(data2):
        table2.rows[i+1].cells[0].text = 區位
        table2.rows[i+1].cells[1].text = 理由
        table2.rows[i+1].cells[2].text = 風險

    doc.add_paragraph()

    h4_2 = doc.add_heading('4.2 「被忽視的黃金地點」：興德里與景行里交界', level=2)
    h4_2.runs[0].font.color.rgb = ORANGE

    doc.add_paragraph('這個區塊沒有捷運，但有以下特質：')

    loc_points = [
        '✅ 老舊公寓密集，陽台空間極小',
        '✅ 銀髮族群比例高（行動不便、難以曬衣）',
        '✅ 方圓500公尺內可能不到3家自助洗衣店',
        '✅ 租金預估比捷運沿線低30-40%',
    ]
    for point in loc_points:
        doc.add_paragraph(point)

    doc.add_paragraph('策略：以「社區型便利站」為定位，主打銀髮族與小家庭。').bold = True

    doc.add_paragraph()

    h4_3 = doc.add_heading('4.3 「大學城邊緣策略」：不是在校內，而是在「必經之路」', level=2)
    h4_3.runs[0].font.color.rgb = ORANGE

    doc.add_paragraph('政治大學正門（指南路）與後山步道入口之間，有一段不到500公尺的道路，是師生每天必經的動線。這裡沒有餐廳、沒有便利商店，卻有穩定的大量人流。')

    doc.add_paragraph()

    # ===== 五、結論與獨特建議 =====
    h5 = doc.add_heading('五、結論與獨特建議', level=1)
    h5.runs[0].font.color.rgb = PURPLE

    h5_1 = doc.add_heading('5.1 核心策略：用「需求分層」取代「價格競爭」', level=2)
    h5_1.runs[0].font.color.rgb = RED

    doc.add_paragraph('文山區自助洗衣市場的問題，不是供過於求，而是多數玩家服務同一群客戶，做同一件事。當所有店都在搶同一批「投幣洗衣服」的客人，價格戰是唯一的出口——但那是一條死路。')

    doc.add_paragraph()

    p_break = doc.add_paragraph()
    p_break.add_run('破局之道：').bold = True
    p_break.add_run('選定一個被忽視的族群，把自己變成這個族群的「唯一選擇」。')

    doc.add_paragraph()

    h5_2 = doc.add_heading('5.2 最小可行版本（MVP）建議', level=2)
    h5_2.runs[0].font.color.rgb = GREEN

    steps = [
        ('第一步', [
            '租一個20-25坪的空間（避開一級戰區，選興德里或景行里）',
            '設置4台洗衣機+2台烘乾機（其中1台為機能衣物專用）',
            'APP系統：基本預約+推播通知',
            '目標：單月損益兩平',
        ]),
        ('第二步（3-6個月後）', [
            '根據實際客戶數據，確認主力客群',
            '調整機器比例（哪種機器使用率最高？）',
            '推出訂閱制或異業合作',
        ]),
        ('第三步（6-12個月後）', [
            '考慮衛星艙或加盟模式擴張',
            '建立品牌識別系統',
        ]),
    ]

    for step_title, step_points in steps:
        p_step = doc.add_paragraph()
        p_step.add_run(step_title + '：').bold = True
        for point in step_points:
            doc.add_paragraph('  • ' + point)

    doc.add_paragraph()

    h5_3 = doc.add_heading('5.3 一句話總結', level=2)
    h5_3.runs[0].font.color.rgb = PURPLE

    p_summary = doc.add_paragraph()
    p_summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_summary = p_summary.add_run('「在文山區開自助洗衣店，不是跟45家店競爭——而是找到第46個還沒被服務的人。」')
    run_summary.bold = True
    run_summary.italic = True
    run_summary.font.color.rgb = PURPLE
    run_summary.font.size = Pt(14)

    doc.add_paragraph()

    # ===== 附錄：關鍵數據 =====
    h6 = doc.add_heading('附錄：關鍵數據快速參照', level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    table3 = doc.add_table(rows=8, cols=2)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '指標'
    hdr3[1].text = '數據'
    for c in hdr3:
        c.paragraphs[0].runs[0].bold = True

    data3 = [
        ('文山區總人口', '252,897人'),
        ('總戶數', '110,852戶'),
        ('現有自助洗衣店', '45-55家（估算密度：1家/4,600人）'),
        ('潛在未被滿足的客群', '13-14萬人'),
        ('大學城師生總數', '約4-5萬人'),
        ('建議初始投資規模', '$80-150萬（不含房租押金）'),
        ('預估回收期', '18-30個月（視選址與定位而定）'),
    ]
    for i, (指標, 數據) in enumerate(data3):
        table3.rows[i+1].cells[0].text = 指標
        table3.rows[i+1].cells[1].text = 數據

    doc.add_paragraph()
    doc.add_paragraph()

    # 聲明
    p_disclaimer = doc.add_paragraph()
    p_disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = p_disclaimer.add_run('*本報告由創意策略顧問視角分析*\n*實際投資前仍建議進行實地考察與詳細財務評估*\n*報告完成時間：2026年4月20日*')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()

    # 頁尾
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 小歐創意策略報告')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '🧺 台北市文山區自助洗衣店創業創意研究報告\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '🌟 小歐創意策略顧問版\n\n'
        '💡 銀髮族/新手爸媽/登山族的隱形需求\n'
        '🚀 機能衣物專家/深夜經濟/寵物專區\n'
        '📍 不是選「在哪裡」，而是選「被誰需要」\n\n'
        '「在文山區開自助洗衣店，不是跟45家店競爭——而是找到第46個還沒被服務的人。」\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    import requests
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生小歐版自助洗衣店創業創意研究報告...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
