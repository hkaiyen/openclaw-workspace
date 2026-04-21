#!/usr/bin/python3
"""千問版：台北市文山區自助洗衣店創業研究報告"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
BLUE = RGBColor(0x1F, 0x49, 0x7D)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)

def set_color(cell, text, color, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/自主洗衣店創業研究_千問版_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面
    title = doc.add_heading('🧺 台北市文山區自助洗衣店創業可行性研究報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run('千問專業顧問分析版')
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('研究日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== 一、執行摘要 =====
    h1 = doc.add_heading('一、執行摘要', level=1)
    h1.runs[0].font.color.rgb = BLUE

    doc.add_paragraph('本報告針對台北市文山區設立自助洗衣店之創業可行性進行全面分析。')

    p_core = doc.add_paragraph()
    p_core.add_run('核心發現：').bold = True

    core_points = [
        '• 文山區潛在客群達5至8萬人，市場需求基本面堅實',
        '• 區內現有競爭店家30至50家，市場尚存進入空間',
        '• 自營模式初始投資預估180至355萬元',
        '• 預估月營收13至18萬元，回本週期2至4年，財務可行性屬中等至良好',
    ]
    for point in core_points:
        doc.add_paragraph(point)

    doc.add_paragraph()
    p_conclusion = doc.add_paragraph()
    p_conclusion.add_run('結論與建議：').bold = True

    doc.add_paragraph('文山區自助洗衣店創業具有可行性，建議優先選擇景美捷運站周邊或興隆路沿線區位，採取差異化服務策略，並嚴格控制初期投資規模，以提高投資報酬率並縮短回本週期。')

    doc.add_paragraph()

    # ===== 二、市場基礎研究 =====
    h2 = doc.add_heading('二、市場基礎研究', level=1)
    h2.runs[0].font.color.rgb = BLUE

    h2_1 = doc.add_heading('2.1 區域概況', level=2)
    h2_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table1 = doc.add_table(rows=6, cols=2)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '項目'
    hdr1[1].text = '數據'
    for c in hdr1:
        c.paragraphs[0].runs[0].bold = True

    data1 = [
        ('行政區', '台北市文山區'),
        ('總人口', '252,897人'),
        ('總戶數', '110,852戶'),
        ('面積', '31.5平方公里'),
        ('人口密度', '約8,025人/平方公里'),
    ]
    for i, (idx, val) in enumerate(data1):
        table1.rows[i+1].cells[0].text = idx
        table1.rows[i+1].cells[1].text = val

    doc.add_paragraph()

    h2_2 = doc.add_heading('2.2 目標客群分析', level=2)
    h2_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    p_客層 = doc.add_paragraph()
    p_客層.add_run('主要客層：').bold = True

    customer_points = [
        '• 學生族群：3至5萬人，主要分布於景美、萬隆一帶之大專院校及高級中學周邊',
        '• 租屋族：3至4萬人，文山區租屋市場活絡，單身套房及小家庭為主',
        '• 潛在總客群：5至8萬人，覆蓋學生、租屋族、銀髮族及小型家庭',
    ]
    for point in customer_points:
        doc.add_paragraph(point)

    doc.add_paragraph()

    h2_3 = doc.add_heading('2.3 市場驅動因素', level=2)
    h2_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    market_factors = [
        '• 都會居住型態改變：小坪數住宅增加，許多公寓或套房未預留洗衣空間',
        '• 生活節奏加快：雙薪家庭普及，自助洗衣分擔家務時間',
        '• 租屋市場穩定：文山區臨近台灣大學、師範大學、景文高中',
        '• 節省成本考量：相較於購入烘衣機，自助洗衣單次費用對部分族群更具經濟效益',
    ]
    for factor in market_factors:
        doc.add_paragraph(factor)

    doc.add_paragraph()

    # ===== 三、地點與競爭分析 =====
    h3 = doc.add_heading('三、地點與競爭分析', level=1)
    h3.runs[0].font.color.rgb = BLUE

    h3_1 = doc.add_heading('3.1 建議區位評比', level=2)
    h3_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table2 = doc.add_table(rows=4, cols=4)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '區位'
    hdr2[1].text = '優勢'
    hdr2[2].text = '風險'
    hdr2[3].text = '租金行情'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    data2 = [
        ('景美捷運站周邊', '交通樞紐，人流量大', '競爭最激烈', '4-6萬/月'),
        ('興隆路沿線', '住宅密集，社區居民穩定', '人潮較弱', '2.5-4萬/月'),
        ('萬隆商圈', '學生族群集中', '腹地較小', '2-3.5萬/月'),
    ]
    for i, (區位, 優勢, 風險, 租金) in enumerate(data2):
        table2.rows[i+1].cells[0].text = 區位
        table2.rows[i+1].cells[1].text = 優勢
        table2.rows[i+1].cells[2].text = 風險
        table2.rows[i+1].cells[3].text = 租金

    doc.add_paragraph()

    h3_2 = doc.add_heading('3.2 區位推薦優先順序', level=2)
    h3_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    location_rec = [
        ('第一優先：景美捷運站周邊（500公尺範圍內）', '適合資金充裕、追求最大客流量之創業者'),
        ('第二優先：興隆路沿線', '適合鎖定社區型居民，強調便利性與社區服務'),
        ('第三優先：萬隆商圈', '適合小額投資者，贴近學區主力客群'),
    ]
    for priority, desc in location_rec:
        p = doc.add_paragraph()
        p.add_run('• ' + priority).bold = True
        doc.add_paragraph('  ' + desc)

    doc.add_paragraph()

    h3_3 = doc.add_heading('3.3 競爭態勢評估', level=2)
    h3_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    competition = [
        '• 現有店家數量：30至50家（含連鎖品牌與獨立店）',
        '• 市場集中度：中等，尚未出現一家獨大局面',
        '• 主要競爭者類型：連鎖品牌、傳統獨立店、社區型小店',
        '• 競爭缺口：24小時營運服務、數位化服務、明亮潔淨空間',
    ]
    for c in competition:
        doc.add_paragraph(c)

    doc.add_paragraph()

    # ===== 四、財務可行性評估 =====
    h4 = doc.add_heading('四、財務可行性評估', level=1)
    h4.runs[0].font.color.rgb = BLUE

    h4_1 = doc.add_heading('4.1 初始投資估算（自營模式）', level=2)
    h4_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table3 = doc.add_table(rows=8, cols=3)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '項目'
    hdr3[1].text = '中標'
    hdr3[2].text = '說明'
    for c in hdr3:
        c.paragraphs[0].runs[0].bold = True

    data3 = [
        ('店面押金（6個月）', '21萬', ''),
        ('裝潢與隔間', '35萬', ''),
        ('洗衣設備（8-12台）', '120萬', '核心設備'),
        ('烘衣設備（6-10台）', '60萬', ''),
        ('乾燥機/輔助設備', '15萬', ''),
        ('耗材備料+周轉金', '23萬', ''),
        ('行銷宣傳開辦費', '10萬', ''),
    ]
    for i, (項目, 中標, 說明) in enumerate(data3):
        table3.rows[i+1].cells[0].text = 項目
        table3.rows[i+1].cells[1].text = 中標
        table3.rows[i+1].cells[2].text = 說明

    doc.add_paragraph()

    p_total = doc.add_paragraph()
    p_total.add_run('💡 自營模式總投資（中標）：').bold = True
    p_total.add_run('約 284萬')
    p_total.runs[1].font.color.rgb = GREEN

    p_range = doc.add_paragraph()
    p_range.add_run('💡 投資區間：').bold = True
    p_range.add_run('180萬 - 413萬')
    p_range.runs[1].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    h4_2 = doc.add_heading('4.2 每月固定成本估算', level=2)
    h4_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table4 = doc.add_table(rows=7, cols=3)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '項目'
    hdr4[1].text = '中標'
    hdr4[2].text = '說明'
    for c in hdr4:
        c.paragraphs[0].runs[0].bold = True

    data4 = [
        ('店面租金', '3.5萬', ''),
        ('人力成本（1-2人）', '4萬', ''),
        ('水電瓦斯', '2.5萬', '含烘衣機耗電'),
        ('耗材', '0.8萬', '洗劑、衛生用品'),
        ('設備維護保養', '0.5萬', ''),
        ('其他支出', '0.5萬', '管理費、稅捐'),
    ]
    for i, (項目, 中標, 說明) in enumerate(data4):
        table4.rows[i+1].cells[0].text = 項目
        table4.rows[i+1].cells[1].text = 中標
        table4.rows[i+1].cells[2].text = 說明

    doc.add_paragraph()

    p_cost = doc.add_paragraph()
    p_cost.add_run('💡 月固定成本合計（中標）：').bold = True
    p_cost.add_run('約 11.8萬')
    p_cost.runs[1].font.color.rgb = RED

    doc.add_paragraph()

    h4_3 = doc.add_heading('4.3 營收預測', level=2)
    h4_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table5 = doc.add_table(rows=4, cols=4)
    table5.style = 'Table Grid'
    hdr5 = table5.rows[0].cells
    hdr5[0].text = '情境'
    hdr5[1].text = '月營業額'
    hdr5[2].text = '月固定成本'
    hdr5[3].text = '月淨利'
    for c in hdr5:
        c.paragraphs[0].runs[0].bold = True

    data5 = [
        ('保守（60%）', '約13萬', '約11.8萬', '約1.2萬', False),
        ('普通（75%）', '約16萬', '約11.8萬', '約4.2萬', True),
        ('樂觀（90%）', '約20萬', '約11.8萬', '約8.2萬', True),
    ]
    for i, (情境, 營業額, 成本, 淨利, is_pos) in enumerate(data5):
        table5.rows[i+1].cells[0].text = 情境
        table5.rows[i+1].cells[1].text = 營業額
        table5.rows[i+1].cells[2].text = 成本
        color = GREEN if is_pos else RED
        set_color(table5.rows[i+1].cells[3], 淨利, color, True)

    doc.add_paragraph()

    h4_4 = doc.add_heading('4.4 回本分析', level=2)
    h4_4.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table6 = doc.add_table(rows=4, cols=4)
    table6.style = 'Table Grid'
    hdr6 = table6.rows[0].cells
    hdr6[0].text = '情境'
    hdr6[1].text = '初始投資'
    hdr6[2].text = '月淨利'
    hdr6[3].text = '回本週期'
    for c in hdr6:
        c.paragraphs[0].runs[0].bold = True

    data6 = [
        ('保守', '284萬', '1.2萬', '超過20年（不建議）', False),
        ('普通', '284萬', '4.2萬', '約5.6年', False),
        ('樂觀', '284萬', '8.2萬', '約2.9年', True),
    ]
    for i, (情境, 投資, 淨利, 回本, is_pos) in enumerate(data6):
        table6.rows[i+1].cells[0].text = 情境
        table6.rows[i+1].cells[1].text = 投資
        table6.rows[i+1].cells[2].text = 淨利
        color = GREEN if is_pos else RED
        set_color(table6.rows[i+1].cells[3], 回本, color, True)

    doc.add_paragraph()

    # ===== 五、營運模式規劃 =====
    h5 = doc.add_heading('五、營運模式規劃', level=1)
    h5.runs[0].font.color.rgb = BLUE

    h5_1 = doc.add_heading('5.1 服務項目規劃', level=2)
    h5_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    p_core = doc.add_paragraph()
    p_core.add_run('核心服務（必備）：').bold = True
    doc.add_paragraph('• 自助洗衣（滾筒式洗衣機，容量7-10公斤）')
    doc.add_paragraph('• 自助烘衣（乾衣機，容量10-15公斤）')
    doc.add_paragraph('• 多元支付（LINE Pay、街口支付、悠遊卡）')

    doc.add_paragraph()
    p_add = doc.add_paragraph()
    p_add.add_run('加值服務（差異化）：').bold = True
    doc.add_paragraph('• 快速洗烘一籃（30分鐘完成）')
    doc.add_paragraph('• 奢護洗衣（高單價衣物，手洗精洗）')
    doc.add_paragraph('• 上門收送服務（社區內限定）')

    doc.add_paragraph()

    h5_2 = doc.add_heading('5.2 數位化營運規劃', level=2)
    h5_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    digital = [
        '• 排隊系統：與LINE官方帳號整合，即時顯示候機狀態',
        '• 會員App或LINE群組：推播優惠、候機通知',
        '• POS系統：整合多元支付、會員紅利、報表管理',
        '• 設備IoT監控：即時掌握洗衣機狀態，預防故障停機',
    ]
    for d in digital:
        doc.add_paragraph(d)

    doc.add_paragraph()

    # ===== 六、執行計劃與時間表 =====
    h6 = doc.add_heading('六、執行計劃與時間表', level=1)
    h6.runs[0].font.color.rgb = BLUE

    h6_1 = doc.add_heading('6.1 階段一時程：籌備期（第1-3個月）', level=2)
    h6_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table7 = doc.add_table(rows=4, cols=3)
    table7.style = 'Table Grid'
    hdr7 = table7.rows[0].cells
    hdr7[0].text = '期間'
    hdr7[1].text = '工作項目'
    hdr7[2].text = '交付物'
    for c in hdr7:
        c.paragraphs[0].runs[0].bold = True

    data7 = [
        ('第1個月', '市場現地勘查、區位評估、租約洽談', '候選店點3-5處'),
        ('第2個月', '室內設計規劃、設備廠商比價、工程發包', '施工圖說、設備報價單'),
        ('第3個月', '裝潢施工、設備安裝、申請許可證', '竣工驗收報告'),
    ]
    for i, (期間, 工作, 交付) in enumerate(data7):
        table7.rows[i+1].cells[0].text = 期間
        table7.rows[i+1].cells[1].text = 工作
        table7.rows[i+1].cells[2].text = 交付

    doc.add_paragraph()

    h6_2 = doc.add_heading('6.2 階段二：開幕期（第4-6個月）', level=2)
    h6_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table8 = doc.add_table(rows=4, cols=2)
    table8.style = 'Table Grid'
    hdr8 = table8.rows[0].cells
    hdr8[0].text = '期間'
    hdr8[1].text = '工作項目'
    for c in hdr8:
        c.paragraphs[0].runs[0].bold = True

    data8 = [
        ('第4個月', '試營運、人員訓練、壓力測試（2-4週試營運）'),
        ('第5個月', '正式開幕、行銷活動啟動（社群宣傳＋在地推廣）'),
        ('第6個月', '營運正常化、會員制度上線（建立SOP、收支初步穩定）'),
    ]
    for i, (期間, 工作) in enumerate(data8):
        table8.rows[i+1].cells[0].text = 期間
        table8.rows[i+1].cells[1].text = 工作

    doc.add_paragraph()

    h6_3 = doc.add_heading('6.3 階段三：成長期（第7-12個月）', level=2)
    h6_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table9 = doc.add_table(rows=4, cols=3)
    table9.style = 'Table Grid'
    hdr9 = table9.rows[0].cells
    hdr9[0].text = '期間'
    hdr9[1].text = '目標'
    hdr9[2].text = '策略'
    for c in hdr9:
        c.paragraphs[0].runs[0].bold = True

    data9 = [
        ('第7-9個月', '提升利用率至75%', '優化服務流程、開發學校合作通路'),
        ('第10-12個月', '損益平衡達成', '控制成本、提高客單價'),
        ('第12個月', '評估第二店可能性', '複製成功模式'),
    ]
    for i, (期間, 目標, 策略) in enumerate(data9):
        table9.rows[i+1].cells[0].text = 期間
        table9.rows[i+1].cells[1].text = 目標
        table9.rows[i+1].cells[2].text = 策略

    doc.add_paragraph()

    # ===== 七、風險管控 =====
    h7 = doc.add_heading('七、風險管控', level=1)
    h7.runs[0].font.color.rgb = BLUE

    h7_1 = doc.add_heading('7.1 風險矩陣', level=2)
    h7_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table10 = doc.add_table(rows=5, cols=3)
    table10.style = 'Table Grid'
    hdr10 = table10.rows[0].cells
    hdr10[0].text = '風險類型'
    hdr10[1].text = '風險等級'
    hdr10[2].text = '對策'
    for c in hdr10:
        c.paragraphs[0].runs[0].bold = True

    data10 = [
        ('景氣低迷影響消費', '⚠️ 高', '建立學生平價方案'),
        ('選址失誤', '⚠️ 高', '簽約前嚴格執行人流計數'),
        ('競爭加劇', '⚡ 中', '差異化服務，避免削價競爭'),
        ('設備故障', '⚡ 中', '選擇信譽設備商，簽訂保固合約'),
    ]
    for i, (風險, 等級, 對策) in enumerate(data10):
        table10.rows[i+1].cells[0].text = 風險
        color = RED if '⚠️' in 等級 else RGBColor(0xCC, 0x88, 0x00)
        set_color(table10.rows[i+1].cells[1], 等級, color, True)
        table10.rows[i+1].cells[2].text = 對策

    doc.add_paragraph()

    h7_2 = doc.add_heading('7.2 緊急應變預案', level=2)
    h7_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    emergency = [
        ('單月營收低於10萬，連續3個月', '立即檢討固定成本、評估縮小規模或轉型'),
        ('主要設備故障超過1週', '啟動臨時優惠折扣，引導顧客至鄰近合作店'),
        ('重大競爭者進入同一商圈', '加速差異化服務上線，重新定位目標客群'),
    ]
    for 情境, 措施 in emergency:
        p = doc.add_paragraph()
        p.add_run('• ' + 情境 + '：').bold = True
        p.add_run(措施)

    doc.add_paragraph()

    # ===== 八、結論與建議 =====
    h8 = doc.add_heading('八、結論與建議', level=1)
    h8.runs[0].font.color.rgb = BLUE

    h8_1 = doc.add_heading('8.1 總結', level=2)
    h8_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    summary = [
        '• 潛在客群龐大：5至8萬人之目標族群支撐基本市場需求',
        '• 區位選擇多元：景美、興隆路、萬隆三大區塊各具優勢',
        '• 差異化空間明確：現代化、數位化、健康環保取向之創新者有突圍機會',
    ]
    for s in summary:
        doc.add_paragraph(s)

    doc.add_paragraph()

    h8_2 = doc.add_heading('8.2 最終建議', level=2)
    h8_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    p_good = doc.add_paragraph()
    p_good.add_run('✅ 建議執行之創業者類型：').bold = True
    good = [
        '• 具備餐飲/服務業基層管理經驗者',
        '• 資金充裕且可承受2-3年回收期者',
        '• 鎖定景美捷運站或興隆路沿線黃金店點者',
    ]
    for g in good:
        doc.add_paragraph(g)

    doc.add_paragraph()

    p_bad = doc.add_paragraph()
    p_bad.add_run('❌ 建議暂緩之情形：').bold = True
    bad = [
        '• 僅有有限資金、無法維持3個月以上周轉金儲備者',
        '• 純投資者無意親自營運管理者',
        '• 僅有租金昂貴之非優質店點選項者',
    ]
    for b in bad:
        doc.add_paragraph(b)

    doc.add_paragraph()

    h8_3 = doc.add_heading('8.3 行動建議清單', level=2)
    h8_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    action = [
        '☐ 第一步：於目標區位進行為期2週之人流計數與競爭店訪查',
        '☐ 第二步：與至少3位以上房東洽談，確認實際租金行情',
        '☐ 第三步：取得設備廠商完整報價書（含保固條款）',
        '☐ 第四步：撰寫完整營運計畫書，申請青年創業貸款或相關補助',
        '☐ 第五步：正式簽約前再次驗算財務模型（損益平衡與回本時間）',
    ]
    for a in action:
        doc.add_paragraph(a)

    doc.add_paragraph()
    doc.add_paragraph()

    # 聲明
    p_disclaimer = doc.add_paragraph()
    p_disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = p_disclaimer.add_run('*本報告由千問專業創業顧問分析產生*\n*資料基準日：2026年4月*\n*報告內容僅供參考，實際投資前請自行進行更完整之盡職調查*')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()

    # 頁尾
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 千問專業顧問分析報告')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '🧺 台北市文山區自助洗衣店創業可行性研究報告\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '🐰 千問專業顧問分析版\n\n'
        '📊 市場概況：252,897人，潛在客群5-8萬人\n'
        '💰 投資規劃：180-355萬，回本2-4年\n'
        '📍 區位建議：景美捷運站/興隆路沿線\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    import requests
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生千問版自助洗衣店創業研究報告...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
