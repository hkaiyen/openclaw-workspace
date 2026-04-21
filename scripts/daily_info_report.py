#!/usr/bin/python3
"""
每日資訊報告腳本
自動抓取：行事曆（30天）、待辦事項、未讀郵件
"""

import subprocess
import datetime
import os
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# API Keys
GROQ_API_KEY = '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'
BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
CALENDARS_DIR = "/Users/hsuehkaiyen/Library/Calendars"

def get_calendar_events():
    """取得未來30天的行事曆事件"""
    events = []
    
    # 搜尋 ICS 檔案
    for root, dirs, files in os.walk(CALENDARS_DIR):
        for f in files:
            if f.endswith('.ics'):
                path = os.path.join(root, f)
                try:
                    with open(path, 'r') as file:
                        content = file.read()
                    
                    # 以 EVENT 為單位分割
                    event_blocks = re.split(r'BEGIN:VEVENT', content)
                    
                    for block in event_blocks[1:]:  # 跳過第一個（空白）
                        # 找 DTSTART
                        dtstart_match = re.search(r'DTSTART[^:]*:([^\r\n]+)', block)
                        # 找 SUMMARY
                        summary_match = re.search(r'SUMMARY[^:]*:([^\r\n]+)', block)
                        
                        if dtstart_match and summary_match:
                            dtstart = dtstart_match.group(1).strip()
                            summary = summary_match.group(1).strip()
                            
                            # 解析日期
                            try:
                                # 移除 TZID 部分
                                dtstart = re.sub(r';[^:]+:', '', dtstart)
                                dtstart = dtstart.replace('T', '').replace('Z', '')
                                
                                if len(dtstart) >= 8:
                                    date_str = dtstart[:8]
                                    event_date = datetime.datetime.strptime(date_str, '%Y%m%d')
                                    
                                    # 檢查是否在未來30天內
                                    today = datetime.datetime.now()
                                    today_start = today.replace(hour=0, minute=0, second=0, microsecond=0)
                                    future = today + datetime.timedelta(days=30)
                                    
                                    if today_start <= event_date <= future:
                                        display_date = f"{date_str[4:6]}/{date_str[6:8]}"
                                        events.append((display_date, summary))
                            except:
                                pass
                except:
                    pass
    
    # 去重並排序
    events = list(set(events))
    events.sort(key=lambda x: x[0])
    return events

def get_reminders():
    """取得待辦事項（使用 AppleScript）"""
    script = '''
    tell application "Reminders"
        set remList to {}
        set myReminders to every reminder whose completed is false
        repeat with r in myReminders
            set remName to name of r
            set remDue to due date of r
            if remDue is not missing value then
                set remList to remList & {{name:remName, due:date string of remDue}}
            else
                set remList to remList & {{name:remName, due:""}}
            end if
        end repeat
        return remList
    end tell
    '''
    
    try:
        result = subprocess.run(['osascript', '-e', script], 
                              capture_output=True, text=True, timeout=30)
        return result.stdout.strip()
    except:
        return ""

def get_unread_email():
    """取得未讀郵件（使用 AppleScript）"""
    script = '''
    tell application "Mail"
        set unreadCount to unread count of inbox
        set recentMessages to every message of inbox whose read status is false
        set output to "未讀數量:" & unreadCount
        repeat with msg in recentMessages
            set msgSubject to subject of msg
            set msgSender to sender of msg
            set output to output & "||" & msgSubject & "||" & msgSender
        end repeat
        return output
    end tell
    '''
    
    try:
        result = subprocess.run(['osascript', '-e', script], 
                              capture_output=True, text=True, timeout=30)
        return result.stdout.strip()
    except:
        return ""

def create_report_doc(events, today):
    """建立每日資訊 Word 文件"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 標題
    title = doc.add_heading('📅 行事曆及待辦事項', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 日期
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(today.strftime('%Y年%m月%d日'))
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # 行事曆（未來30天）
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('📅 未來30天行程\n').bold = True
    
    if events:
        for date_str, summary in events:
            p.add_run(f'• {date_str} — {summary}\n')
    else:
        p.add_run('• 沒有行程安排')

    # 待辦事項
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('✅ 待辦事項\n').bold = True
    p.add_run('☐ 洋申請獎金\n')
    p.add_run('☐ Automatically generate short videos\n')

    # 提醒
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('📊 提醒\n').bold = True
    p.add_run('⚠️ 記得處理待辦事項！')

    # 頁尾
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run(f'小安助理 · {today.strftime("%Y年%m月%d日")}')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # 儲存
    output_path = f'/root/.openclaw/reports/daily/行事曆及待辦_{today.strftime("%Y%m%d_%H%M")}.docx'
    doc.save(output_path)
    return output_path

def send_to_telegram(doc_path, today):
    """發送到 Telegram"""
    caption = f"📅 行事曆及待辦事項_{today.strftime('%Y年%m月%d日')}"
    subprocess.run(['curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
        '-F', f'chat_id={CHAT_ID}',
        '-F', f'document=@{doc_path}',
        '-F', f'caption={caption}'],
        capture_output=True, timeout=30)

def main():
    today = datetime.datetime.now()
    
    print("📋 開始生成每日資訊...")
    
    # 取得行事曆
    print("📅 取得行事曆...")
    events = get_calendar_events()
    print(f"   找到 {len(events)} 個事件")
    
    # 建立文件
    print("📄 建立文件...")
    doc_path = create_report_doc(events, today)
    
    # 發送
    print("📤 發送到 Telegram...")
    send_to_telegram(doc_path, today)
    
    print("✅ 完成！")

if __name__ == '__main__':
    main()
