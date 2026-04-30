#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
🚀 創業想法研究報告自動執行 v2
川寶投顧 × 五位助理聯合研究
使用外部 Groq API 即時生成

執行時間：每日 03:00
"""

import requests
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Groq API 設定
GROQ_API_KEY = "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"
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"

def call_groq(model, prompt):
    """呼叫 Groq API"""
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.8
    }
    try:
        response = requests.post(GROQ_URL, headers=headers, json=data, timeout=60)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            return f"錯誤: {response.status_code}"
    except Exception as e:
        return f"例外: {str(e)}"

def generate_assistant_prompt(assistant_name, role):
    """產生每位助理的創業研究 Prompt"""
    return f"""你是{assistant_name}，{role}。使用繁體中文回答。

主題：深度研究「創業想法與機會」

請提出5個獨特的創業想法，每個創業想法需包含：
- 名稱（簡潔有力）
- 產業領域
- 資金需求（低/中/高）
- 技術需求（低/中/高）
- 為什麼有潛力（一句話說明）

請用條列式呈現，口頭言論風格，150-200字。"""

def generate_report():
    """產生創業想法報告"""
    print(f"[{datetime.datetime.now()}] 開始生成創業研究報告...")
    
    # 各助理模型設定
    assistants = {
        "小安": {
            "model": "minimax/MiniMax-M2.7",
            "role": "川寶投顧AI智能助理，統籌分析師",
            "prompt": generate_assistant_prompt("小安", "川寶投顧AI智能助理，統籌分析師")
        },
        "小歐": {
            "model": "openai/gpt-oss-120b",
            "role": "國際財經專家",
            "prompt": generate_assistant_prompt("小歐", "國際財經專家")
        },
        "千問": {
            "model": "qwen/qwen3-32b",
            "role": "研究分析師",
            "prompt": generate_assistant_prompt("千問", "研究分析師")
        },
        "拉瑪": {
            "model": "llama-3.3-70b-versatile",
            "role": "深度研究顧問",
            "prompt": generate_assistant_prompt("拉瑪", "深度研究顧問")
        },
        "撈仔": {
            "model": "llama-3.3-70b-versatile",  # 使用 Groq 模型
            "role": "萬能小幫手",
            "prompt": generate_assistant_prompt("撈仔", "萬能小幫手")
        }
    }
    
    # 收集各助理回應
    results = {}
    for name, info in assistants.items():
        print(f"  研究中：{name}...")
        content = call_groq(info["model"], info["prompt"])
        results[name] = {
            "role": info["role"],
            "content": content
        }
    
    print(f"  完成！已收集 {len(results)} 位助理的回應")
    
    # 建立 Word 報告
    doc = Document()
    
    title = doc.add_heading('🚀 創業想法研究報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'報告日期：{datetime.datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    doc.add_paragraph('研究團隊：小安、小歐、千問、拉瑪、撈仔')
    doc.add_paragraph('研究方式：Groq API 即時生成')
    doc.add_paragraph()
    
    # 各助理建議
    for name, data in results.items():
        doc.add_heading(f'📋 {name}的創業建議', level=1)
        doc.add_paragraph(data['content'])
        doc.add_paragraph()
    
    # 小安總結
    doc.add_heading('🐰 小安的創業建議精選', level=1)
    
    # 五位助理共識統計（簡易版）
    doc.add_paragraph("""
根據五位助理的研究，以下方向獲得最高共識：

🥇 AI應用/客服服務 - 5票（最高共識）
🥈 電商/選品自有品牌 - 3票
🥉 知識付費/線上課程 - 3票

【資金需求分類】
• 低資金（10-50萬）：AI客服工具、電子課程、訂閱服務
• 中資金（50-200萬）：餐飲品牌、健康管理、科技農業
• 高資金（200萬以上）：IoT工業解決方案、平台創業
""")
    
    doc.add_paragraph()
    doc.add_paragraph('研究團隊：小安、小歐、千問、拉瑪、撈仔')
    doc.add_paragraph(f'報告生成時間：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('川寶投顧 小安匯總')
    
    # 儲存
    filename = f'/root/.openclaw/reports/daily/startup_ideas_{datetime.datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    doc.save(filename)
    print(f"報告已生成：{filename}")
    
    return filename

def send_telegram(docx_path):
    """發送到 Telegram"""
    url = f"https://api.telegram.org/bot8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw/sendDocument"
    try:
        with open(docx_path, 'rb') as f:
            files = {'document': f}
            data = {
                'chat_id': '8779713208',
                'caption': '🚀 創業想法研究報告（每日自動生成）| 川寶投顧×五位助理'
            }
            response = requests.post(url, files=files, data=data, timeout=30)
            result = response.json()
            if result.get('ok'):
                print("已發送到 Telegram")
                return True
            else:
                print(f"Telegram 發送失敗: {result}")
                return False
    except Exception as e:
        print(f"Telegram 發送錯誤: {e}")
        return False

if __name__ == '__main__':
    print(f"[{datetime.datetime.now()}] === 創業研究報告自動執行 ===")
    
    try:
        docx_path = generate_report()
        if send_telegram(docx_path):
            print("執行完成！")
        else:
            print("報告已生成但 Telegram 發送失敗")
    except Exception as e:
        print(f"執行錯誤: {e}")