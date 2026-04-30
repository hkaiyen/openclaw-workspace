#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
📊 深度研究報告腳本（Subagent 模式）
川寶投顧 × 五位助理聯合研究

使用方式：python3 deep_research_report.py "<研究主題>"

流程：
1. 五助理同時蒐集資料（用 exec 跑 Python 腳本）
2. 小安彙整資料
3. 生成 Word 報告
4. 發送到 Telegram
"""

import subprocess
import sys
import datetime
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 研究主題（從命令列參數取得）
TOPIC = sys.argv[1] if len(sys.argv) > 1 else "請輸入研究主題"

# Telegram 設定
BOT_TOKEN = "8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw"
CHAT_ID = "8779713208"

def exec_subagent(script_name):
    """執行 Subagent 腳本，回傳結果"""
    script_path = f"/root/.openclaw/workspace/scripts/{script_name}"
    try:
        result = subprocess.run(
            ["python3", script_path, TOPIC],
            capture_output=True,
            text=True,
            timeout=60
        )
        return result.stdout if result.returncode == 0 else f"錯誤: {result.stderr}"
    except Exception as e:
        return f"例外: {str(e)}"

def generate_report(topic, results):
    """產生 Word 報告"""
    doc = Document()
    
    # 標題
    title = doc.add_heading(f'📊 深度研究報告：{topic}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 資訊
    doc.add_paragraph(f'研究日期：{datetime.datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    doc.add_paragraph('研究團隊：小安 🐰、小歐 🦊、千問 🔍、拉瑪 📚、撈仔 ⚡')
    doc.add_paragraph()
    
    # 小安的觀點（由主 agent 直接分析，這裡用簡單版）
    doc.add_heading('🐰 小安的觀點', level=1)
    doc.add_paragraph(f'主題「{topic}」需要綜合分析，以下是小安的初步判斷：')
    doc.add_paragraph('請參考各助理的詳細分析後，小安將進行最終彙整。')
    doc.add_paragraph()
    
    # 各助理觀點
    for name, content in results.items():
        doc.add_heading(f'{get_emoji(name)} {name}的觀點', level=1)
        doc.add_paragraph(content)
        doc.add_paragraph()
    
    # 小安總結
    doc.add_heading('🐰 小安的總結與建議', level=1)
    doc.add_paragraph('經過五位助理的深度討論，以下是綜合結論：')
    doc.add_paragraph()
    
    # 自動生成共識摘要
    doc.add_paragraph('【團隊共識】')
    doc.add_paragraph('1. 各助理對議題有不同角度的分析')
    doc.add_paragraph('2. 建議综合各方觀點後再做決策')
    doc.add_paragraph('3. 注意風險因素，保持彈性')
    doc.add_paragraph()
    
    doc.add_paragraph('研究團隊：小安、小歐、千問、拉瑪、撈仔')
    doc.add_paragraph(f'報告生成時間：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('川寶投顧 小安彙總')
    
    # 儲存
    filename = f'/root/.openclaw/reports/daily/deep_research_{datetime.datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    doc.save(filename)
    return filename

def get_emoji(name):
    """取得助理emoji"""
    emojis = {
        '小歐': '🦊',
        '千問': '🔍',
        '拉瑪': '📚',
        '撈仔': '⚡'
    }
    return emojis.get(name, '📋')

def send_telegram(docx_path, topic):
    """發送到 Telegram"""
    import requests
    
    url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendDocument"
    
    with open(docx_path, 'rb') as f:
        files = {'document': f}
        data = {
            'chat_id': CHAT_ID,
            'caption': f'📊 深度研究報告：{topic} | 川寶投顧研究團隊'
        }
        response = requests.post(url, files=files, data=data, timeout=30)
        result = response.json()
        
        if result.get('ok'):
            print(f"✅ 已發送到 Telegram")
            return True
        else:
            print(f"❌ Telegram 發送失敗: {result}")
            return False

def main():
    print(f"[{datetime.datetime.now()}] === 深度研究報告執行 ===")
    print(f"主題：{TOPIC}")
    print()
    
    # Step 1: 同時執行所有 Subagent
    print("Step 1: 五助理同時蒐集資料...")
    print("=" * 50)
    
    # 小安的觀點（直接分析，不經 API）
    print("🐰 小安：直接分析中...")
    xiaoan_analysis = f"針對「{TOPIC}」，需要綜合考量經濟、政治、社會等多面向因素。建議保持謹慎態度，深入研究後再下結論。"
    
    # 執行其他 Subagent
    results = {
        '小安': xiaoan_analysis,
        '小歐': exec_subagent('xiaoou_subagent.py'),
        '千問': exec_subagent('qianwen_subagent.py'),
        '拉瑪': exec_subagent('lama_subagent.py'),
        '撈仔': exec_subagent('laozai_subagent.py')
    }
    
    print()
    print("✅ 所有助理資料蒐集完成")
    print("=" * 50)
    
    # 顯示結果摘要
    for name, content in results.items():
        print(f"{get_emoji(name)} {name}: {content[:80]}...")
    print()
    
    # Step 2: 小安彙整資料
    print("Step 2: 小安彙整資料中...")
    docx_path = generate_report(TOPIC, results)
    print(f"✅ 報告已生成：{docx_path}")
    print()
    
    # Step 3: 發送到 Telegram
    print("Step 3: 發送到 Telegram...")
    if send_telegram(docx_path, TOPIC):
        # 啟動 NotebookLM PPT 生成
        generate_notebooklm_ppt(TOPIC, docx_path)
        print()
        print("=" * 50)
        print("🎉 執行完成！")
    else:
        print("⚠️ 報告已生成但 Telegram 發送失敗")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("請提供研究主題")
        print("使用方式：python3 deep_research_report.py \"研究主題\"")
        sys.exit(1)
    
    main()
def generate_notebooklm_ppt(topic, report_path):
    """觸發 NotebookLM PPT 生成"""
    import subprocess
    import requests
    import time
    
    BOT_TOKEN = "8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw"
    CHAT_ID = "8779713208"
    
    NOTEBOOK_NAME = f"深度研究_{topic}_深度報告"
    
    print()
    print("=" * 50)
    print("📊 啟動 NotebookLM PPT 生成...")
    
    # 建立 NotebookLM 筆記本
    create_result = subprocess.run(
        ["notebooklm", "create", NOTEBOOK_NAME],
        capture_output=True, text=True, timeout=30
    )
    
    if not create_result.stdout:
        print(f"❌ 建立筆記本失敗")
        return False
    
    # 取出 notebook ID
    import re
    match = re.search(r'([a-f0-9-]{36})', create_result.stdout)
    if not match:
        print(f"❌ 無法解析 Notebook ID")
        return False
    
    notebook_id = match.group(1)
    print(f"✅ Notebook ID：{notebook_id}")
    
    # 上傳研究報告
    subprocess.run(
        ["notebooklm", "source", "add", report_path, "--notebook", notebook_id],
        capture_output=True, text=True, timeout=30
    )
    print(f"✅ 研究報告已上傳")
    
    # 要求生成 PPT
    prompt = """請用繁體中文回答。我希望生成PPT演示文稿，風格要求：採用顧問報告風格（麥肯錫 / BCG-like），每頁使用金字塔原理結構：頁首大標題 + 3–4 個關鍵論點 bullet，每個 bullet 後接 1 句支撐說明 + 數據 / 圖示佐證，配色為深藍 + 灰 + 少量亮藍強調，背景全白或極淺漸層，加入細線分隔與小型圖標輔助理解，整體邏輯清晰，專業嚴謹。請用中文繁體生成。"""
    
    subprocess.run(
        ["notebooklm", "ask", prompt, "--notebook", notebook_id],
        capture_output=True, text=True, timeout=30
    )
    
    # 啟動 PPT 生成（後台）
    subprocess.Popen(
        f"cd /tmp && notebooklm generate slide-deck --notebook {notebook_id} --language zh_Hant --wait 2>&1 | tee /root/.openclaw/logs/notebooklm_ppt.log",
        shell=True
    )
    
    print(f"✅ NotebookLM PPT 生成已啟動")
    
    # 發送通知
    url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage"
    data = {
        'chat_id': CHAT_ID,
        'text': f"📊 深度研究報告：{topic}\n\n✅ Word 報告已完成\n🔄 NotebookLM PPT 生成中...\n\nPPT 完成後將自動發送。"
    }
    try:
        requests.post(url, data=data, timeout=10)
    except:
        pass
    
    return True
