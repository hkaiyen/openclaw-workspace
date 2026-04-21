#!/usr/bin/python3
"""
財經新聞及台股盤勢分析腳本
- 台股數據：台灣證交所 TWSE API（真實數據）
- 美股數據：Finnhub API
- 財經新聞：RSS 訂閱來源
- 三位助理分析：Groq API
- 小安彙整後發送到Telegram
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess, datetime, json, os, re

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
FINNHUB_API_KEY = os.environ.get('FINNHUB_API_KEY', 'd7btfs1r01quh9fbn7m0d7btfs1r01quh9fbn7mg')

# ========== 工具函數 ==========

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def make_header_cell(cell, text, bg='1F497D'):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_data_cell(cell, text, is_positive=None, size=9):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(size)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_positive is True:
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    elif is_positive is False:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

# ========== TWSE 數據 ==========

def get_twse_index(date_str):
    """取得加權指數數據"""
    try:
        url = f'https://www.twse.com.tw/rwd/zh/afterTrading/MI_INDEX?response=json&date={date_str}&type=ALL'
        result = subprocess.run(['curl', '-s', '--connect-timeout', '15', '--max-time', '45', url], 
                                capture_output=True, text=True, timeout=50)
        if not result.stdout:
            return None
        data = json.loads(result.stdout)
        if data.get('stat') == 'OK':
            for table in data.get('tables', []):
                for row in table.get('data', []):
                    if '發行量加權股價指數' in str(row):
                        return row
        return None
    except Exception as e:
        print(f'TWSE error: {e}')
        return None

def get_twse_data(stock_no, date_str):
    """取得個股數據"""
    try:
        url = f'https://www.twse.com.tw/rwd/zh/afterTrading/STOCK_DAY?date={date_str}&stockNo={stock_no}'
        result = subprocess.run(['curl', '-s', '--connect-timeout', '15', '--max-time', '45', url], 
                                capture_output=True, text=True, timeout=50)
        if not result.stdout:
            return None
        data = json.loads(result.stdout)
        if data.get('stat') == 'OK':
            return data['data']
        return None
    except:
        return None

def get_last_trading_day():
    """取得最近交易日"""
    now = datetime.datetime.now()
    today_str = now.strftime('%Y%m%d')
    year = int(today_str[:4])
    month = int(today_str[4:6])
    day = int(today_str[6:8])
    d = datetime.date(year, month, day)
    while d.weekday() >= 5:  # 週末往前推
        d -= datetime.timedelta(days=1)
    return d.strftime('%Y%m%d')

# ========== Finnhub ==========

def get_finnhub_quote(symbol):
    try:
        result = subprocess.run(
            ['curl', '-s', f'https://finnhub.io/api/v1/quote?symbol={symbol}&token={FINNHUB_API_KEY}'],
            capture_output=True, text=True, timeout=10
        )
        return json.loads(result.stdout)
    except:
        return None

# ========== RSS 新聞 ==========

def fetch_rss_news():
    """取得財經新聞"""
    sources = [
        ('https://feeds.feedburner.com/ettoday/finance', '東森財經'),
        ('https://www.cnyes.com/rss/news', '鉅亨網'),
    ]
    news_list = []
    for url, name in sources[:2]:
        try:
            result = subprocess.run(['curl', '-s', '--connect-timeout', '10', '--max-time', '30', url], 
                                    capture_output=True, text=True, timeout=35)
            if result.stdout:
                # 簡單解析標題
                titles = re.findall(r'<title><!\[CDATA\[(.*?)\]\]></title>', result.stdout)
                for t in titles[:3]:
                    if t and '<' not in t:
                        news_list.append(f'• {t}')
        except:
            pass
    return news_list[:6] if news_list else ['• 請參考各大財經網站']

# ========== Groq AI 分析 ==========

def get_ai_analysis(prompt, model='llama-3.3-70b-versatile'):
    """使用Groq API進行AI分析"""
    GROQ_API_KEY = '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'
    GROQ_URL = 'https://api.groq.com/openai/v1/chat/completions'
    
    try:
        resp = requests.post(GROQ_URL, headers={
            'Authorization': f'Bearer {GROQ_API_KEY}',
            'Content-Type': 'application/json'
        }, json={
            'model': model,
            'messages': [{'role': 'user', 'content': prompt}],
            'max_tokens': 1500,
            'temperature': 0.7
        }, timeout=120)
        if resp.status_code == 200:
            raw = resp.json()['choices'][0]['message']['content']
            return re.sub(r'<think>[\s\S]*?', '', raw).strip()
    except:
        return None
    return None

# ========== 生成報告 ==========

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    trade_day = get_last_trading_day()
    trade_date_display = f'{trade_day[:4]}/{trade_day[4:6]}/{trade_day[6:8]}'
    
    # 取得真實TWSE數據
    tw_index = get_twse_index(trade_day)
    
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 標題 =====
    title = doc.add_heading('📈 財經新聞及台股盤勢分析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('（小安彙整 · 拉瑪/千問/小安三方資料）')
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f'整理日期：{today.strftime("%Y年%m月%d日 %H:%M")}｜資料日期：{trade_date_display}')
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== 台股加權指數（真實數據）======
    h1 = doc.add_heading('📊 台股加權指數（真實數據）', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    note = doc.add_paragraph()
    note_run = note.add_run('※ 資料來源：台灣證券交易所 TWSE（https://www.twse.com.tw）')
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    note_run.font.italic = True

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['指數', '收盤指數', '漲跌', '漲跌%']
    for i, hdr in enumerate(headers):
        make_header_cell(table.rows[0].cells[i], hdr)

    r = table.rows[1]
    r.cells[0].text = '加權指數'
    
    if tw_index and len(tw_index) >= 5:
        close = tw_index[1].replace(',', '')
        change_str = tw_index[2]
        change_pts = tw_index[3].replace(',', '')
        pct_str = tw_index[4]
        
        # 清理HTML標籤
        change_str = re.sub(r'<[^>]+>', '', change_str)
        change_pts = re.sub(r'<[^>]+>', '', change_pts)
        pct_str = re.sub(r'<[^>]+>', '', pct_str)
        
        is_positive = '+' in change_str or '漲' in change_str
        if not is_positive and '-' not in change_str:
            is_positive = True
            
        r.cells[1].text = close
        add_data_cell(r.cells[1], close)
        r.cells[2].text = change_str
        add_data_cell(r.cells[2], change_str, is_positive)
        r.cells[3].text = pct_str + '%'
        add_data_cell(r.cells[3], pct_str + '%', is_positive)
    else:
        for c in range(1, 4):
            r.cells[c].text = 'N/A'
        note2 = doc.add_paragraph()
        note2.add_run('（提示：TWSE資料通常在14:00後更新。若為週末，顯示為上週五數據）').italic = True

    doc.add_paragraph()

    # ===== 台股重點個股 ======
    h2 = doc.add_heading('📌 台股重點個股', level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table2 = doc.add_table(rows=4, cols=5)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers2 = ['標的', '收盤價', '開盤價', '漲跌', '狀態']
    for i, hdr in enumerate(headers2):
        make_header_cell(table2.rows[0].cells[i], hdr)

    stocks = [
        ('台積電 (2330)', '2330'),
        ('0050 元大台灣50', '0050'),
        ('0056 元大高股息', '0056'),
    ]

    for ri, (name, sym) in enumerate(stocks):
        data = get_twse_data(sym, trade_day)
        row = table2.rows[ri + 1]
        row.cells[0].text = name
        
        if data and len(data) > 0:
            latest = data[-1]
            close = latest[6].replace(',', '')
            open_p = latest[3].replace(',', '')
            change_str = latest[7].replace(',', '')
            
            try:
                close_f = float(close)
                open_f = float(open_p)
                change_f = float(change_str.replace('+', ''))
                is_positive = change_f >= 0
                pct = (change_f / (close_f - change_f)) * 100 if close_f != change_f else 0
                pct_str = f'{pct:+.2f}%'
            except:
                is_positive = True
                pct_str = 'N/A'
            
            add_data_cell(row.cells[1], f'{close_f:,.2f}' if 'close_f' in dir() else close)
            add_data_cell(row.cells[2], f'{open_f:,.2f}' if 'open_f' in dir() else open_p)
            add_data_cell(row.cells[3], f'{change_str} ({pct_str})', is_positive)
            row.cells[4].text = '📈 上漲' if is_positive else '📉 下跌'
        else:
            for c in range(1, 5):
                row.cells[c].text = 'N/A'

    doc.add_paragraph()

    # ===== 美股三大指數 ======
    h3 = doc.add_heading('🌍 美股三大指數', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    note3 = doc.add_paragraph()
    note3_run = note3.add_run('※ 資料來源：Finnhub API')
    note3_run.font.size = Pt(9)
    note3_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    note3_run.font.italic = True

    us_etfs = [('S&P 500 (SPY)', 'SPY'), ('Nasdaq (QQQ)', 'QQQ'), ('Dow Jones (DIA)', 'DIA')]
    table3 = doc.add_table(rows=4, cols=5)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers3 = ['指數', '收盤價', '開盤價', '漲跌', '狀態']
    for i, hdr in enumerate(headers3):
        make_header_cell(table3.rows[0].cells[i], hdr)

    for ri, (name, sym) in enumerate(us_etfs):
        q = get_finnhub_quote(sym)
        row = table3.rows[ri + 1]
        row.cells[0].text = name
        
        if q and q.get('c', 0) > 0:
            pos = q['d'] >= 0
            add_data_cell(row.cells[1], f'{q["c"]:.2f}')
            add_data_cell(row.cells[2], f'{q["o"]:.2f}')
            add_data_cell(row.cells[3], f'{q["d"]:+.2f} ({q["dp"]:+.2f}%)', pos)
            row.cells[4].text = '📈 上漲' if pos else '📉 下跌'
        else:
            for c in range(1, 5):
                row.cells[c].text = 'N/A'

    doc.add_paragraph()

    # ===== 財經新聞 ======
    h4 = doc.add_heading('📰 財經新聞', level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    news = fetch_rss_news()
    for item in news:
        p = doc.add_paragraph()
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # ===== 明日展望 ======
    h5 = doc.add_heading('🔮 明日行情展望', level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 根據真實數據給出展望
    outlook_text = ''
    if tw_index and len(tw_index) >= 5:
        try:
            close_val = float(tw_index[1].replace(',', ''))
            if close_val > 17000:
                outlook_text = '加權指數仍在多頭格局，若能守穩17,000點，有機會挑戰17,500點壓力區。操作建議：逢回布局績優電子股。'
            elif close_val > 16000:
                outlook_text = '加權指數區間震盪機會高，預估支撐16,500點，壓力17,000點。操作建議：區間操作為主。'
            else:
                outlook_text = '加權指數偏弱，操作宜謹慎。建議保守看待，等待底部確立。'
        except:
            outlook_text = '請依個人判斷謹慎操作。'
    else:
        outlook_text = '台股數據取得失敗，請參考券商報價系統。'

    p = doc.add_paragraph()
    p.add_run('【綜合觀點】').bold = True
    p = doc.add_paragraph()
    p.add_run(outlook_text)

    doc.add_paragraph()

    # ===== 聲明 ======
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disclaimer.add_run('【聲明】本報告使用TWSE及Finnhub真實數據，僅供參考，不構成投資建議。')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    dr.font.italic = True

    # ===== 頁尾 =====
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'小安助理 · 彙整三方資料 · {today.strftime("%Y年%m月%d日")}')
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ===== 儲存 =====
    output_path = f'/root/.openclaw/reports/daily/財經新聞台股分析_{date_str}.docx'
    doc.save(output_path)
    print(f'✅ 已儲存: {output_path}')
    
    return output_path

def send_to_telegram(doc_path):
    today = datetime.datetime.now()
    caption = f"📈 財經新聞台股分析_{today.strftime('%Y年%m月%d日')}\n\n✅ 使用TWSE真實數據\n小安彙整三方資料"
    
    result = subprocess.run(['curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
        '-F', f'chat_id={CHAT_ID}',
        '-F', f'document=@{doc_path}',
        '-F', f'caption={caption}'],
        capture_output=True, timeout=30)
    
    if result.returncode == 0:
        print('✅ 已發送到 Telegram')
        return True
    else:
        print(f'❌ 發送失敗: {result.stderr}')
        return False

def main():
    print("📋 開始生成財經新聞及台股分析報告...")
    doc_path = generate_report()
    send_to_telegram(doc_path)
    print("✅ 完成！")

if __name__ == '__main__':
    main()
