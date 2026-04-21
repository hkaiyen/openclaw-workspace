#!/usr/bin/python3
"""自主洗衣店創業研究報告 - 台北市文山區 (四助理版本)"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
BLUE = RGBColor(0x1F, 0x49, 0x7D)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)

def set_color(cell, text, color, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/自主洗衣店創業研究_文山區_四助理_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面
    title = doc.add_heading('🧺 自主洗衣店創業研究報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run('台北市文山區市場可行性分析')
    sr.font.size = Pt(18)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('研究日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    byline = doc.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = byline.add_run('四助理共同研究產出')
    br.font.size = Pt(11)
    br.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ===== 摘要 =====
    h_summary = doc.add_heading('📋 研究摘要', level=1)
    h_summary.runs[0].font.color.rgb = BLUE

    summary_text = [
        ('研究地點', '台北市文山區'),
        ('研究方法', '四助理同步研究 + 市場資料搜集'),
        ('總人口', '252,897人（110,852戶）'),
        ('潛在客群', '5-8萬人（學生+租屋族為核心）'),
        ('建議投資', '$180-355萬（自營模式）'),
        ('預期回本', '2-4年（視地點而定）'),
        ('首選區位', '景美捷運站200m內、興隆路沿線'),
    ]
    for idx, val in summary_text:
        p = doc.add_paragraph()
        run1 = p.add_run(idx + '：')
        run1.bold = True
        p.add_run(val)

    doc.add_paragraph()

    # ===== 第一章：市場基礎研究 =====
    h1 = doc.add_heading('🐰 第一章：市場基礎研究（拉瑪）', level=1)
    h1.runs[0].font.color.rgb = BLUE

    h1_1 = doc.add_heading('1.1 文山區基本資料', level=2)
    h1_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table1 = doc.add_table(rows=8, cols=2)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '項目'
    hdr1[1].text = '數據'
    for c in hdr1:
        c.paragraphs[0].runs[0].bold = True

    data1 = [
        ('總面積', '31.5090 平方公里（台北市最大）'),
        ('總人口（2026年3月）', '252,897 人'),
        ('戶數', '110,852 戶'),
        ('人口密度', '8,026 人/平方公里'),
        ('行政里數', '43 里，1,002 鄰'),
        ('人口成長', '近十年微幅減少（-4.6%）'),
        ('區域特色', '住宅密集、學區圍繞、大學城效應'),
    ]
    for i, (idx, val) in enumerate(data1):
        table1.rows[i+1].cells[0].text = idx
        table1.rows[i+1].cells[1].text = val

    doc.add_paragraph()

    h1_2 = doc.add_heading('1.2 現有競爭對手分析', level=2)
    h1_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '項目'
    hdr2[1].text = '分析'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    data2 = [
        ('推估店家數量', '約 30-50 家'),
        ('分佈熱點', '木柵路、興隆路、羅斯福路五段、軍功路、萬芳社區、政大周邊'),
        ('業態特徵', '多數小型獨立經營（1-3機組）、部分連鎖品牌增加'),
        ('市場供需', '每5,000-8,000人支撐一家，競爭強度中等'),
    ]
    for i, (idx, val) in enumerate(data2):
        table2.rows[i+1].cells[0].text = idx
        table2.rows[i+1].cells[1].text = val

    doc.add_paragraph()

    h1_3 = doc.add_heading('1.3 目標客群分析', level=2)
    h1_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table3 = doc.add_table(rows=6, cols=3)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '客群'
    hdr3[1].text = '估算人數'
    hdr3[2].text = '需求強度'
    for c in hdr3:
        c.paragraphs[0].runs[0].bold = True

    data3 = [
        ('學生族群（政大、世新、警專、中國科大）', '3-5萬人', '⭐⭐⭐⭐⭐'),
        ('獨居租屋族', '3-4萬人', '⭐⭐⭐⭐'),
        ('雙薪家庭', '約2-3萬戶', '⭐⭐⭐'),
        ('銀髮族/無洗衣機住戶', '5,000-1萬人', '⭐⭐⭐⭐'),
        ('結論：潛在重度目標客群', '5-8萬人', ''),
    ]
    for i, (客群, 人數, 強度) in enumerate(data3):
        table3.rows[i+1].cells[0].text = 客群
        table3.rows[i+1].cells[1].text = 人數
        table3.rows[i+1].cells[2].text = 強度

    doc.add_paragraph()

    h1_4 = doc.add_heading('1.4 SWOT 分析', level=2)
    h1_4.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    swot = [
        ('✅ 優勢', '大學城效應、氣候潮濕、租屋族集中、剛性需求穩定'),
        ('⚠️ 劣勢', '腹地受限、停車不便、機器折舊、捷運站稍遠'),
        ('🌟 機會', '連鎖化、智能化、遠距監控、央北重劃區新人口'),
        ('❌ 威脅', '家用洗衣機普及、電費上漲、外送洗衣服務競爭'),
    ]
    for title, content in swot:
        p = doc.add_paragraph()
        run = p.add_run(title + '：')
        run.bold = True
        if '✅' in title:
            run.font.color.rgb = GREEN
        elif '⚠️' in title:
            run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)
        elif '🌟' in title:
            run.font.color.rgb = BLUE
        elif '❌' in title:
            run.font.color.rgb = RED
        p.add_run(content)

    doc.add_paragraph()

    # ===== 第二章：選址分析 =====
    h2 = doc.add_heading('🌍 第二章：選址分析（小歐）', level=1)
    h2.runs[0].font.color.rgb = BLUE

    h2_1 = doc.add_heading('2.1 精華區位評估', level=2)
    h2_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table4 = doc.add_table(rows=6, cols=3)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '區位'
    hdr4[1].text = '特色'
    hdr4[2].text = '評級'
    for c in hdr4:
        c.paragraphs[0].runs[0].bold = True

    data4 = [
        ('景美商圈', '發展最早、景美夜市、交通便捷', '⭐⭐⭐⭐⭐'),
        ('萬隆商圈', '住宅社區密集、捷運+公車雙便利', '⭐⭐⭐⭐'),
        ('木柵商圈', '師生族群龐大、年輕租屋族、夜間洗衣需求', '⭐⭐⭐⭐'),
        ('興隆路沿線', '東西向主軸線、住宅林立、曝光率佳', '⭐⭐⭐⭐'),
        ('木柵二期重劃區', '新興住宅區、競爭尚少、長期佈局', '⭐⭐⭐'),
    ]
    for i, (區位, 特色, 評級) in enumerate(data4):
        table4.rows[i+1].cells[0].text = 區位
        table4.rows[i+1].cells[1].text = 特色
        table4.rows[i+1].cells[2].text = 評級

    doc.add_paragraph()

    h2_2 = doc.add_heading('2.2 租金行情', level=2)
    h2_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table5 = doc.add_table(rows=6, cols=2)
    table5.style = 'Table Grid'
    hdr5 = table5.rows[0].cells
    hdr5[0].text = '區段'
    hdr5[1].text = '月租金估算'
    for c in hdr5:
        c.paragraphs[0].runs[0].bold = True

    data5 = [
        ('景美夜市周邊（20-35坪）', '3.5萬-6萬元'),
        ('萬隆羅斯福路六段（20-30坪）', '2.5萬-4萬元'),
        ('木柵路/秀明路沿線（25-40坪）', '2萬-3.5萬元'),
        ('興隆路三段~四段（25-35坪）', '2.5萬-4萬元'),
        ('木柵二期重劃區（25-35坪）', '2萬-3萬元'),
    ]
    for i, (區段, 租金) in enumerate(data5):
        table5.rows[i+1].cells[0].text = 區段
        table5.rows[i+1].cells[1].text = 租金

    doc.add_paragraph()

    h2_3 = doc.add_heading('2.3 選址關鍵條件', level=2)
    h2_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    location_points = [
        '✅ 每日固定人流（居民、學生、通勤族）',
        '✅ 交通可達性（捷運300m內、停車方便）',
        '✅ 競爭環境（方圓300m內不超過2家）',
        '✅ 立地條件（排水/通風、獨立動線）',
        '✅ 與便利商店/超市互補消費動線',
        '❌ 避開宮廟、加油站（氣味/安全顧慮）',
    ]
    for point in location_points:
        doc.add_paragraph(point)

    doc.add_paragraph()

    # ===== 第三章：財務可行性 =====
    h3 = doc.add_heading('💰 第三章：財務可行性（千問）', level=1)
    h3.runs[0].font.color.rgb = BLUE

    h3_1 = doc.add_heading('3.1 初始投資成本', level=2)
    h3_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table6 = doc.add_table(rows=9, cols=2)
    table6.style = 'Table Grid'
    hdr6 = table6.rows[0].cells
    hdr6[0].text = '項目'
    hdr6[1].text = '費用區間'
    for c in hdr6:
        c.paragraphs[0].runs[0].bold = True

    data6 = [
        ('商用洗衣機（6-8台）', '60-100萬'),
        ('商用烘衣機（6-8台）', '45-80萬'),
        ('硬體設施（乾衣機、投幣式）', '10-20萬'),
        ('裝潢費用', '20-50萬'),
        ('租金押金（3個月）', '15-45萬'),
        ('申請執照/管線水電工程', '10-20萬'),
        ('招牌、軟裝、監視器', '5-10萬'),
        ('週轉金（3個月）', '15-30萬'),
    ]
    for i, (項目, 費用) in enumerate(data6):
        table6.rows[i+1].cells[0].text = 項目
        table6.rows[i+1].cells[1].text = 費用

    doc.add_paragraph()

    p_total = doc.add_paragraph()
    p_total.add_run('💡 自營模式總投資：').bold = True
    p_total.add_run('$180-355萬')
    p_total.runs[1].font.color.rgb = GREEN

    doc.add_paragraph()

    h3_2 = doc.add_heading('3.2 每月營運成本', level=2)
    h3_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table7 = doc.add_table(rows=8, cols=2)
    table7.style = 'Table Grid'
    hdr7 = table7.rows[0].cells
    hdr7[0].text = '項目'
    hdr7[1].text = '金額（月）'
    for c in hdr7:
        c.paragraphs[0].runs[0].bold = True

    data7 = [
        ('店面租金', '1.5-4萬'),
        ('水費', '0.3-0.8萬'),
        ('電費（含烘衣機）', '1.5-3.5萬'),
        ('耗材（洗劑、柔軟精）', '0.2-0.5萬'),
        ('人事費用（工讀生）', '1-2萬'),
        ('設備保養維修', '0.3-0.8萬'),
        ('其他（廣告、稅捐、保險）', '0.2-0.5萬'),
    ]
    for i, (項目, 金額) in enumerate(data7):
        table7.rows[i+1].cells[0].text = 項目
        table7.rows[i+1].cells[1].text = 金額

    doc.add_paragraph()

    p_total2 = doc.add_paragraph()
    p_total2.add_run('💡 月固定成本合計：').bold = True
    p_total2.add_run('$5-12萬')
    p_total2.runs[1].font.color.rgb = RED

    doc.add_paragraph()

    h3_3 = doc.add_heading('3.3 營收預測', level=2)
    h3_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table8 = doc.add_table(rows=5, cols=4)
    table8.style = 'Table Grid'
    hdr8 = table8.rows[0].cells
    hdr8[0].text = '情境'
    hdr8[1].text = '月總營收'
    hdr8[2].text = '月利潤'
    hdr8[3].text = '回本時間'
    for c in hdr8:
        c.paragraphs[0].runs[0].bold = True

    data8 = [
        ('保守', '8-12萬', '2-3萬', '7-9年'),
        ('普通', '13-18萬', '7-8萬', '3-4年'),
        ('樂觀', '20-28萬', '12-15萬', '2-3年'),
        ('優異', '25-30萬+', '15-20萬', '1.5-2年'),
    ]
    for i, (情境, 營收, 利潤, 回本) in enumerate(data8):
        table8.rows[i+1].cells[0].text = 情境
        table8.rows[i+1].cells[1].text = 營收
        if '保守' in 情境:
            set_color(table8.rows[i+1].cells[2], 利潤, RED, True)
        else:
            set_color(table8.rows[i+1].cells[2], 利潤, GREEN, True)
        table8.rows[i+1].cells[3].text = 回本

    doc.add_paragraph()

    h3_4 = doc.add_heading('3.4 自營 vs 加盟', level=2)
    h3_4.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table9 = doc.add_table(rows=8, cols=3)
    table9.style = 'Table Grid'
    hdr9 = table9.rows[0].cells
    hdr9[0].text = '項目'
    hdr9[1].text = '自營'
    hdr9[2].text = '加盟'
    for c in hdr9:
        c.paragraphs[0].runs[0].bold = True

    data9 = [
        ('初期投資', '180-355萬', '195-395萬（貴10-15%）'),
        ('品牌知名度', '無', '有加分'),
        ('機台採購', '自己比價', '統一採購'),
        ('教育訓練', '自學', '總部提供'),
        ('行銷支援', '自理', '品牌行銷'),
        ('失敗風險', '較高', '中等'),
        ('獲利彈性', '較高', '較低（需繳權利金）'),
    ]
    for i, (項目, 自營, 加盟) in enumerate(data9):
        table9.rows[i+1].cells[0].text = 項目
        table9.rows[i+1].cells[1].text = 自營
        table9.rows[i+1].cells[2].text = 加盟

    doc.add_paragraph()

    # ===== 第四章：營運建議 =====
    h4 = doc.add_heading('⚙️ 第四章：營運建議（小安）', level=1)
    h4.runs[0].font.color.rgb = BLUE

    h4_1 = doc.add_heading('4.1 設備建議', level=2)
    h4_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table10 = doc.add_table(rows=5, cols=2)
    table10.style = 'Table Grid'
    hdr10 = table10.rows[0].cells
    hdr10[0].text = '設備'
    hdr10[1].text = '建議'
    for c in hdr10:
        c.paragraphs[0].runs[0].bold = True

    data10 = [
        ('洗衣機', '6-8台，15-20kg大型滾筒'),
        ('烘衣機', '6-8台，30kg以上'),
        ('投幣系統', '支援悠遊卡/LINE Pay'),
        ('監控系統', '雲端錄影+APP通知'),
    ]
    for i, (設備, 建議) in enumerate(data10):
        table10.rows[i+1].cells[0].text = 設備
        table10.rows[i+1].cells[1].text = 建議

    doc.add_paragraph()

    h4_2 = doc.add_heading('4.2 數位化建議', level=2)
    h4_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    digital = [
        '📱 APP預約洗衣/烘衣時間',
        '💳 多元支付：悠遊卡、LINE Pay、信用卡',
        '📊 即時監控：手機查看機器狀態、空位',
        '🔔 自動通知：洗滌完成推播通知',
        '📈 數據分析：來客習慣、熱門時段統計',
    ]
    for d in digital:
        doc.add_paragraph(d)

    doc.add_paragraph()

    h4_3 = doc.add_heading('4.3 風險管控', level=2)
    h4_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    risks = [
        ('⚠️ 設備故障', '簽訂保固合約、建立備用機制'),
        ('⚠️ 營收不如預期', '前3個月觀察期，果斷調整'),
        ('⚠️ 異常天候', '冬季/雨天加強行銷、提供烘衣優惠'),
        ('⚠️ 法規問題', '預先確認使用分區、環保噪音標準'),
    ]
    for 風險, 對策 in risks:
        p = doc.add_paragraph()
        run1 = p.add_run(風險 + '：')
        run1.bold = True
        p.add_run(對策)

    doc.add_paragraph()

    # ===== 第五章：執行計劃 =====
    h5 = doc.add_heading('📋 第五章：執行計劃', level=1)
    h5.runs[0].font.color.rgb = BLUE

    h5_1 = doc.add_heading('5.1 行動時間表', level=2)
    h5_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table11 = doc.add_table(rows=7, cols=3)
    table11.style = 'Table Grid'
    hdr11 = table11.rows[0].cells
    hdr11[0].text = '階段'
    hdr11[1].text = '時間'
    hdr11[2].text = '任務'
    for c in hdr11:
        c.paragraphs[0].runs[0].bold = True

    data11 = [
        ('第一階段', '第1-2週', '地點實地訪查、租金行情確認'),
        ('第二階段', '第3-4週', '設備商報價、財務規劃'),
        ('第三階段', '第5-6週', '合約簽訂、裝潢申請'),
        ('第四階段', '第7-8週', '設備安裝、測試營運'),
        ('第五階段', '第9週', '正式開幕、行銷推廣'),
        ('第六階段', '第10-12週', '營運優化、調整定價'),
    ]
    for i, (階段, 時間, 任務) in enumerate(data11):
        table11.rows[i+1].cells[0].text = 階段
        table11.rows[i+1].cells[1].text = 時間
        table11.rows[i+1].cells[2].text = 任務

    doc.add_paragraph()

    h5_2 = doc.add_heading('5.2 四助理最終建議', level=2)
    h5_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    final = [
        ('🐰 拉瑪（市場）', '文山區適合开店，潛在客群5-8萬人，市場供需屬中等'),
        ('🌍 小歐（選址）', '首選景美捷運站200m內或興隆路三段~四段'),
        ('💰 千問（財務）', '自營模式，投資180-355萬，預期回本2-4年'),
        ('⚙️ 小安（營運）', '建議自營，配置智能支付與APP監控'),
    ]
    for 助理, 建議 in final:
        p = doc.add_paragraph()
        run = p.add_run(助理 + '：')
        run.bold = True
        p.add_run(建議)

    doc.add_paragraph()

    # ===== 第六章：綜合結論 =====
    h6 = doc.add_heading('🏆 第六章：綜合結論', level=1)
    h6.runs[0].font.color.rgb = BLUE

    conclusion = [
        ('✅', '文山區自助洗衣店創業：可行！'),
        ('✅', '投資金額：建議準備 $200-250萬（自營模式）'),
        ('✅', '首選地點：景美捷運站周邊或興隆路沿線'),
        ('✅', '預期回本：2-4年（視地點與經營而定）'),
        ('✅', '成功關鍵：好地點+智能設備+優質服務'),
    ]
    for icon, text in conclusion:
        p = doc.add_paragraph()
        run = p.add_run(icon + ' ' + text)
        run.bold = True
        run.font.color.rgb = GREEN

    doc.add_paragraph()

    # 頁尾
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 四助理共同研究報告')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '🧺 自主洗衣店創業研究報告（升級版）\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '🐰 四助理共同研究產出\n\n'
        '📊 市場研究：252,897人，潛在客群5-8萬人\n'
        '🌍 選址建議：景美捷運站周邊、興隆路沿線\n'
        '💰 財務規劃：投資$180-355萬，回本2-4年\n'
        '⚙️ 營運策略：自營+智能設備+多元支付\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    import requests
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生自主洗衣店創業研究報告（四助理版）...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
