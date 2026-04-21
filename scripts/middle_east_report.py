#!/usr/bin/python3
"""
中東局勢追蹤報告生成腳本 - 即時版
新聞來源：BBC + Al Jazeera RSS（翻譯為中文）
國際+臺灣最新資訊
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import datetime
import requests
import re
import time
import urllib.parse
import json

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
BLUE = RGBColor(0x1F, 0x49, 0x7D)
RED = RGBColor(0xCC, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x80, 0x00)
ORANGE = RGBColor(0xFF, 0x66, 0x00)

def translate_to_chinese(text):
    """使用 MyMemory API 翻譯為中文"""
    if not text or len(text) < 2:
        return text
    try:
        encoded = urllib.parse.quote(text[:500])
        url = f"https://api.mymemory.translated.net/get?q={encoded}&langpair=en|zh-TW"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode())
            result = data.get("responseData", {}).get("translatedText", "")
            if result and result != text:
                return result
    except:
        pass
    return text

def search_duckduckgo(query, max_results=8):
    """用 DuckDuckGo 搜尋即時資訊"""
    try:
        from duckduckgo_search import DDGS
        results = []
        with DDGS() as ddgs:
            # 新聞搜尋
            news_results = list(ddgs.news(query, region='wt-wt', max_results=max_results))
            for r in news_results:
                results.append({
                    'title': r.get('title', ''),
                    'body': r.get('body', '')[:150],
                    'url': r.get('url', ''),
                    'source': r.get('source', '未知'),
                    'date': r.get('date', '')
                })
        return results
    except Exception as e:
        print(f'DuckDuckGo error: {e}')
        return []

def fetch_rss_news(url, source_name, max_items=30):
    """抓取 RSS 新聞"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
            'Accept': 'application/rss+xml, application/xml, text/xml, */*'
        }
        resp = requests.get(url, headers=headers, timeout=10)
        resp.encoding = 'utf-8'
        
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(resp.text, 'html.parser')
        items = soup.find_all('item')
        
        news_list = []
        for item in items[:max_items]:
            title_elem = item.find('title')
            title_text = re.sub(r'<!\[CDATA\[|\]\]>', '', title_elem.get_text(strip=True)) if title_elem else ""
            link_elem = item.find('link')
            link_text = link_elem.get_text(strip=True) if link_elem else ""
            
            if title_text:
                news_list.append({
                    'title': title_text,
                    'link': link_text,
                    'source': source_name
                })
        return news_list
    except Exception as e:
        print(f'RSS error for {source_name}: {e}')
        return []

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ===== 標題 =====
    title = doc.add_heading(f'🌍 中東局勢追蹤', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(26)
    title.runs[0].font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('即時國際新聞 · 臺灣相關資訊')
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    sr.font.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run(f'🕐 更新時間：{today.strftime("%Y年%m月%d日 %H:%M")}')
    dr.font.size = Pt(11)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== RSS 新聞來源 =====
    rss_sources = [
        ('https://feeds.bbci.co.uk/news/world/rss.xml', 'BBC World'),
        ('https://www.aljazeera.com/xml/rss/all.xml', 'Al Jazeera'),
    ]

    all_rss_news = []
    for url, name in rss_sources:
        news = fetch_rss_news(url, name, max_items=50)
        all_rss_news.extend(news)
        time.sleep(1)

    # 去重
    seen_titles = set()
    unique_news = []
    for item in all_rss_news:
        key = item['title'][:40]
        if key not in seen_titles:
            seen_titles.add(key)
            unique_news.append(item)

    # 翻譯所有標題為中文（一次翻完，避免顯示時多次 API 呼叫）
    print(f'   翻譯 {len(unique_news)} 則新聞標題...')
    for item in unique_news:
        original = item.get('title', '')
        if original:
            translated = translate_to_chinese(original)
            item['_title_zh'] = translated
            time.sleep(0.3)

    # ===== 搜尋即時資訊（DuckDuckGo 已停用，改用 RSS）=====
    # search_queries = [...]  # 已停用
    # 
    # for query, count in search_queries:
    #     print(f'  搜尋: {query}')
    #     results = search_duckduckgo(query, max_results=count)
    #     search_results.extend(results)
    #     time.sleep(3)
    # 
    # seen_urls = set()
    # unique_search = []
    # for item in search_results:
    #     url_key = item.get('url', '')[:50]
    #     if url_key and url_key not in seen_urls:
    #         seen_urls.add(url_key)
    #         unique_search.append(item)
    unique_search = []  # DuckDuckGo 已停用，改用 BBC + Al Jazeera

    # ===== 第一章：最新局勢摘要 =====
    h1 = doc.add_heading('📊 最新局勢摘要', level=1)
    h1.runs[0].font.color.rgb = BLUE

    keywords_summary = ['伊朗', '以色列', '加薩', '停火', '談判', '中東', '美伊', '葉門', '荷莫茲', 'Gaza', 'Israel', 'Iran', 'Middle East', 'Trump', 'Hezbollah', 'Lebanon', 'Syria']
    summary_news = [n for n in unique_news if any(k in n['title'] for k in keywords_summary)][:6]

    if summary_news:
        for item in summary_news:
            p = doc.add_paragraph()
            p.add_run(f'• {item.get("_title_zh", item["title"])}').bold = False
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.3)
            source_run = p.add_run(f' [{item["source"]}]')
            source_run.font.size = Pt(9)
            source_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    else:
        p = doc.add_paragraph('（暫無最新消息）')

    doc.add_paragraph()

    # ===== 第二章：國際最新動態 =====
    h2 = doc.add_heading('🌏 國際最新動態', level=1)
    h2.runs[0].font.color.rgb = BLUE

    # 顯示所有 RSS 新聞（不限中東）
    intl_news = unique_news[:12]
    if intl_news:
        for item in intl_news:
            p = doc.add_paragraph()
            p.add_run(f'• {item.get("_title_zh", item["title"])}')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.3)
            if item.get('source'):
                src = p.add_run(f' - {item["source"]}')
                src.font.size = Pt(9)
                src.font.color.rgb = RGBColor(0x66, 0x99, 0xCC)
    else:
        p = doc.add_paragraph('（暫無最新消息）')

    doc.add_paragraph()

    # ===== 第三章：美伊核子談判 =====
    h3 = doc.add_heading('⚔️ 美伊核子談判與制裁', level=1)
    h3.runs[0].font.color.rgb = ORANGE

    iranian_keywords = ['伊朗', '制裁', '核子', '濃縮鈾', '川普', '拜登', '外交', '談判', '協議', 'Iran', 'nuclear', 'sanction', 'Trump', 'uranium', 'negotiation', 'deal']
    iranian_news = [n for n in unique_news if any(k in n['title'] for k in iranian_keywords)][:12]
    seen_iran = set()
    final_iranian = []
    for item in iranian_news:
        key = item.get('title', item.get('url', ''))[:40]
        if key not in seen_iran:
            seen_iran.add(key)
            final_iranian.append(item)

    if final_iranian:
        for item in final_iranian[:8]:
            p = doc.add_paragraph()
            p.add_run(f'• {item.get("_title_zh", item.get("title"))}')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.3)
    else:
        p = doc.add_paragraph('（暫無最新消息）')

    doc.add_paragraph()

    # ===== 第四章：以色列與加薩 =====
    h4 = doc.add_heading('🇮🇱 以色列與加薩走廊', level=1)
    h4.runs[0].font.color.rgb = RED

    isr_keywords = ['以色列', '加薩', '哈瑪斯', '巴勒斯坦', '停火', '人質', '納坦雅胡', '拉法', 'Israel', 'Gaza', 'Hamas', 'Palestine', 'ceasefire', 'hostage', 'Netanyahu', 'Rafah']
    isr_news = [n for n in unique_news if any(k in n['title'] for k in isr_keywords)][:12]

    seen_isr = set()
    final_isr = []
    for item in isr_news:
        key = item.get('title', item.get('url', ''))[:40]
        if key not in seen_isr:
            seen_isr.add(key)
            final_isr.append(item)

    if final_isr:
        for item in final_isr[:8]:
            p = doc.add_paragraph()
            p.add_run(f'• {item.get("_title_zh", item.get("title", ""))}')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.3)
    else:
        p = doc.add_paragraph('（暫無最新消息）')

    doc.add_paragraph()

    # ===== 第五章：能源與航運 =====
    h5 = doc.add_heading('⚡ 能源市場與航運影響', level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    energy_keywords = ['油價', '原油', 'OPEC', '沙烏地', '石油', '天然氣', '荷莫茲', '航運', '油輪', '能源', 'Oil', 'crude', 'energy', 'Yemen', 'Hormuz']
    energy_news = [n for n in unique_news if any(k in n['title'] for k in energy_keywords)][:12]

    seen_energy = set()
    final_energy = []
    for item in energy_news:
        key = item.get('title', item.get('url', ''))[:40]
        if key not in seen_energy:
            seen_energy.add(key)
            final_energy.append(item)

    if final_energy:
        for item in final_energy[:7]:
            p = doc.add_paragraph()
            p.add_run(f'• {item.get("_title_zh", item.get("title", ""))}')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.3)
    else:
        p = doc.add_paragraph('（暫無最新消息）')

    doc.add_paragraph()

    # ===== 第六章：臺灣相關資訊 =====
    h6 = doc.add_heading('🇹🇼 臺灣相關資訊', level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    tw_keywords = ['臺灣', '台灣', '两岸', '兩岸', '美臺', '美台', '軍演', '中共', '中國', '南海']
    tw_news = [n for n in unique_news if any(k in n['title'] for k in tw_keywords)][:6]

    if tw_news:
        for item in tw_news:
            p = doc.add_paragraph()
            p.add_run(f'• {item.get("_title_zh", item["title"])}')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.3)
            src = p.add_run(f' [{item["source"]}]')
            src.font.size = Pt(9)
            src.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    else:
        p = doc.add_paragraph('（暂无與臺灣相關的中東新聞）')

    doc.add_paragraph()

    # ===== 第七章：全球市場影響 =====
    h7 = doc.add_heading('💹 全球市場影響', level=1)
    h7.runs[0].font.color.rgb = BLUE

    market_keywords = ['美股', '亞股', '台股', '股市', '投資', 'Fed', '聯準會', '通膨', 'stock', 'market', 'inflation', 'Fed rate', 'Trump', 'tariff', 'trade']
    market_news = [n for n in unique_news if any(k in n['title'] for k in market_keywords)][:12]

    seen_market = set()
    final_market = []
    for item in market_news:
        key = item.get('title', item.get('url', ''))[:40]
        if key not in seen_market:
            seen_market.add(key)
            final_market.append(item)

    if final_market:
        for item in final_market[:7]:
            p = doc.add_paragraph()
            p.add_run(f'• {item.get("_title_zh", item.get("title", ""))}')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.3)
    else:
        p = doc.add_paragraph('（暫無最新市場資訊）')

    doc.add_paragraph()

    # ===== 資料來源 =====
    h8 = doc.add_heading('📋 資料來源', level=1)
    h8.runs[0].font.color.rgb = BLUE

    sources = [
        '• BBC World News (RSS)',
        '• Al Jazeera (RSS)',
        '• DuckDuckGo 即時搜尋（已停用，暫用 RSS 替代）',
    ]
    for s in sources:
        p = doc.add_paragraph(s)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # 免責聲明
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disc.add_run('【聲明】本報告僅供參考，資料來源為公開資訊，請以官方公布為準。')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    dr.font.italic = True

    # 頁尾
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 中東局勢追蹤即時版')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    output_path = f'/root/.openclaw/reports/daily/中東局勢追蹤_即時版_{date_str}.docx'
    doc.save(output_path)
    return output_path

def send_to_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        f'🌍 中東局勢追蹤（即時版）\n'
        f'{today.strftime("%Y年%m月%d日 %H:%M")}\n\n'
        f'📊 最新局勢摘要\n'
        f'🌏 國際最新動態\n'
        f'⚔️ 美伊核子談判\n'
        f'🇮🇱 以色列與加薩\n'
        f'⚡ 能源市場與航運\n'
        f'🇹🇼 臺灣相關資訊\n'
        f'💹 全球市場影響\n\n'
        f'資料來源：DuckDuckGo 即時搜尋 + RSS\n'
        f'小安製'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    print(f'[{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}] 開始生成中東局勢報告（即時版）...')
    try:
        report_path = generate_report()
        print(f'📄 報告已生成: {report_path}')
        if send_to_telegram(report_path):
            print('✅ 已發送到 Telegram')
        else:
            print('❌ 發送失敗')
        print(f'[{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}] 任務完成！')
    except Exception as e:
        print(f'❌ 錯誤: {e}')