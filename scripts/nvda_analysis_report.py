#!/usr/bin/python3
"""輝達 (NVDA) 財報分析報告"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess, datetime, json, os, time

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'

GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xCC)

def set_color(cell, text, color, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/輝達NVDA財報分析_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('📊 輝達 (NVDA) 財報深度分析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(26)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('報告日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    period_p = doc.add_paragraph()
    period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = period_p.add_run('資料來源：FY Q4 2026財報 (2026/02/25發布) | Yahoo Finance')
    pr.font.size = Pt(10)
    pr.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    doc.add_paragraph()

    # 核心指標
    h1 = doc.add_heading('📌 FY Q4 2026 核心財務指標', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '指標'
    hdr1[1].text = '數據'
    hdr1[2].text = ' YoY變化'
    for c in hdr1:
        c.paragraphs[0].runs[0].bold = True

    data1 = [
        ('EPS (稀釋)', '$0.89', '+72%', True),
        ('營收', '$119.0億', '+78%', True),
        ('數據中心營收', '$97.4億', '+93%', True),
        ('毛利率', '73.5%', '+1%', True),
        ('營業利益率', '64.2%', '+2%', True),
    ]
    for i, (idx, val, change, is_pos) in enumerate(data1):
        table1.rows[i+1].cells[0].text = idx
        table1.rows[i+1].cells[1].text = val
        color = GREEN if is_pos else RED
        set_color(table1.rows[i+1].cells[2], change, color, True)

    doc.add_paragraph()

    # 營收組成
    h2 = doc.add_heading('📈 營收分析', level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '業務類別'
    hdr2[1].text = 'Q4 FY2026'
    hdr2[2].text = 'Q4 FY2025'
    hdr2[3].text = 'YoY'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    data2 = [
        ('數據中心 (Data Center)', '$97.4億', '$50.4億', '+93%', True),
        ('遊戲 (Gaming)', '$14.5億', '$12.7億', '+14%', True),
        ('專業視覺 (Pro Vis)', '$4.8億', '$4.1億', '+17%', True),
        ('車用 (Automotive)', '$2.0億', '$1.4億', '+43%', True),
        ('OEM/其他', '$0.3億', '$0.4億', '-25%', False),
    ]
    for i, (idx, q4, q4_25, yoy, is_pos) in enumerate(data2):
        table2.rows[i+1].cells[0].text = idx
        table2.rows[i+1].cells[1].text = q4
        table2.rows[i+1].cells[2].text = q4_25
        color = GREEN if is_pos else RED
        set_color(table2.rows[i+1].cells[3], yoy, color, True)

    doc.add_paragraph()

    # 股價表現
    h3 = doc.add_heading('💹 股價表現對比', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table3 = doc.add_table(rows=5, cols=4)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '期間'
    hdr3[1].text = '期初'
    hdr3[2].text = '期末'
    hdr3[3].text = '報酬率'
    for c in hdr3:
        c.paragraphs[0].runs[0].bold = True

    data3 = [
        ('Q1 2025 (1-3月)', '$134.29', '$109.67', '-18.33%', False),
        ('Q4 2025 (10-12月)', '$186.58', '$187.54', '+0.51%', True),
        ('Q1 2026 (1-3月)', '$186.50', '$165.17', '-11.44%', False),
        ('近一年 (2025/04~)', '$108.38', '$201.68', '+86.09%', True),
    ]
    for i, (period, start, end, ret, is_pos) in enumerate(data3):
        table3.rows[i+1].cells[0].text = period
        table3.rows[i+1].cells[1].text = start
        table3.rows[i+1].cells[2].text = end
        color = GREEN if is_pos else RED
        set_color(table3.rows[i+1].cells[3], ret, color, True)

    doc.add_paragraph()

    # 關鍵數據
    h4 = doc.add_heading('📊 關鍵營運數據', level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table4 = doc.add_table(rows=7, cols=2)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '指標'
    hdr4[1].text = '數據'
    for c in hdr4:
        c.paragraphs[0].runs[0].bold = True

    data4 = [
        ('52週最高', '$212.19'),
        ('52週最低', '$86.62'),
        ('Q1 2026區間最高', '$197.63'),
        ('Q1 2026區間最低', '$164.27'),
        ('平均日成交量', '~1.8億股'),
        ('市值', '~USD $4,950億'),
    ]
    for i, (idx, val) in enumerate(data4):
        table4.rows[i+1].cells[0].text = idx
        table4.rows[i+1].cells[1].text = val

    doc.add_paragraph()

    # AI需求分析
    h5 = doc.add_heading('🤖 AI需求帶動效應', level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    ai_points = [
        '📈 數據中心營收佔比達82%，年增93%為最大成長動能',
        '🤖 H100/H200 GPU需求持續強勁，Blackwell架構出貨放量',
        '💎 毛利率維持73.5%高檔，顯示定價權強勁',
        '🏭 供應鏈瓶頸緩解，營收預測成長動能延續',
        '☁️ 雲端服務商資本支出創高，支撐AI晶片需求',
    ]
    for point in ai_points:
        p = doc.add_paragraph(point)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # 風險因素
    h6 = doc.add_heading('⚠️ 風險因素', level=1)
    h6.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    risk_points = [
        '📉 Q1 2026股價回檔11%，估值過高引發獲利了結',
        '💰 競爭加劇：AMD MI300X、Intel Gaudi 3搶市佔',
        '🏛️ 中國出口管制：中國營收佔比降至個位數',
        '💵 庫存修正：供應鏈庫存調整影響季度營收',
        '🔄 產品過渡：Blackwell新舊產品交替期的不確定性',
    ]
    for risk in risk_points:
        p = doc.add_paragraph(risk)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # 總結
    h7 = doc.add_heading('📋 投資結論', level=1)
    h7.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    conclusion = [
        ('✅ AI領導地位：數據中心營收年增93%，遙遙領先競爭對手', True),
        ('✅ 毛利率73.5%：高技術壁壘維持強勁定價權', True),
        ('✅ 近一年報酬+86%：長期趨勢仍然向上', True),
        ('⚠️ Q1 2026股價回檔11%：估值修正壓力待消化', False),
        ('⚠️ 競爭加劇：AMD/Intel追趕，需觀察市佔變化', False),
    ]
    for text, is_pos in conclusion:
        p = doc.add_paragraph()
        p.add_run(text)
        p.paragraph_format.space_after = Pt(6)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 輝達NVDA財報分析')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '📊 輝達 (NVDA) 財報深度分析\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '📌 FY Q4 2026 核心指標\n'
        '• EPS: $0.89 (+72% YoY)\n'
        '• 營收: $119.0億 (+78% YoY)\n'
        '• 數據中心營收: $97.4億 (+93% YoY)\n'
        '• 毛利率: 73.5%\n\n'
        '💹 股價表現\n'
        '• Q1 2026: -11.44%\n'
        '• 近一年: +86.09%\n\n'
        '🤖 AI需求持續爆發，數據中心營收佔比82%\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    import requests
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生輝達NVDA財報分析報告...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
