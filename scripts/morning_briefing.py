#!/usr/bin/python3
"""
每日晨間摘要報告
1. 天氣資訊
2. 行事曆（今日+未來30天）
3. Google Tasks 待辦事項
4. Gmail 未讀郵件
"""

import subprocess
import datetime
import os
import re
import json
import time
import imaplib
import email
import urllib.request
import urllib.parse
from email.header import decode_header
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import yfinance

# ========== iCloud CalDAV 設定 ==========
ICLOUD_USER = 'Hkaiyen@icloud.com'
ICLOUD_APP_PASSWORD = 'lpve-dfwe-spxv-pcdh'
CALDAV_HOST = 'https://p116-caldav.icloud.com:443/1056470819/calendars/'

# iCloud 行事曆 UUID（主要行事曆）
PRIMARY_CALENDARS = [
    '5B62C2B1-1709-4A2B-9036-AD2021FD4DF5',  # 行事曆（主要）
    '1E8BE61F-92D9-4D01-9716-1D8AC2FA751B',  # T-EX行事曆
    '33C297E9-D9DA-4C60-9E99-AC81AFFD2044',  # Work
    '6CADDA48-509D-4675-8A4E-087D978D8FDB',  # Routine
]

# ========== 常數 ==========
BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
CALENDARS_DIR = "/Users/hsuehkaiyen/Library/Calendars"

# Gmail 設定
GMAIL_USER = 'Hkaiyen@gmail.com'
GMAIL_APP_PASSWORD = 'hvusuczdoejqecnm'

# ========== 工具函數 ==========
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def make_header_cell(cell, text, bg='1F497D'):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, bg)

def add_data_cell(cell, text, is_positive=None):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ========== 1. 天氣資訊 ==========
def get_weather():
    """取得天氣資料（wttr.in）"""
    try:
        result = subprocess.run(['curl', '-s', 'wttr.in/Taipei?format=j1'],
                              capture_output=True, text=True, timeout=15)
        data = json.loads(result.stdout)
        curr = data['current_condition'][0]
        tom = data['weather'][1]
        tom_hourly = tom.get('hourly', [])
        tom_desc = tom_hourly[0]['weatherDesc'][0]['value'] if (tom_hourly and tom_hourly[0].get('weatherDesc')) else 'N/A'
        
        return {
            'current': {
                'desc': curr['weatherDesc'][0]['value'],
                'temp': f"{curr['temp_C']}°C",
                'feels': f"{curr['FeelsLikeC']}°C",
                'humidity': f"{curr['humidity']}%",
                'rain': f"{curr.get('precipMM', '0')}mm",
                'wind': f"{curr['windspeedKmph']} km/h",
                'uv': curr['uvIndex'],
                'vis': f"{curr['visibility']} km",
            },
            'tomorrow': {
                'desc': tom_desc,
                'max': tom['maxtempC'],
                'min': tom['mintempC'],
            }
        }
    except Exception as e:
        print(f"天氣取得失敗: {e}")
        return None

# ========== 1.5 台灣節日 ==========
def get_taiwan_holidays():
    """取得台灣節日（從政府資料庫）"""
    import datetime
    
    # 2026 年節日（直接列出主要節日）
    holidays_2026 = {
        '2026-01-01': '元旦',
        '2026-01-28': '春節假開始',
        '2026-01-29': '春節',
        '2026-01-30': '春節',
        '2026-01-31': '春節',
        '2026-02-01': '春節',
        '2026-02-02': '春節',
        '2026-02-03': '春節',
        '2026-02-04': '春節',
        '2026-02-14': '春節補假',
        '2026-02-15': '和平紀念日',
        '2026-02-16': '和平紀念日補假',
        '2026-02-28': '228紀念日',
        '2026-03-08': '婦女節',
        '2026-03-29': '清明節',
        '2026-04-03': '兒童節',
        '2026-04-04': '兒童節補假',
        '2026-04-05': '清明節',
        '2026-05-01': '勞動節',
        '2026-06-19': '端午節',
        '2026-06-20': '端午節補假',
        '2026-09-28': '教師節',
        '2026-10-01': '中秋節',
        '2026-10-02': '中秋節補假',
        '2026-10-10': '國慶日',
        '2026-10-25': '台灣光復節',
        '2026-11-12': '國父逝世紀念日',
        '2026-12-25': '聖誕節',
    }
    
    # 嘗試從 API 取得最新節日資料
    try:
        url = 'https://data.gov.ai/api/holidays?year=2026&country=TW'
        result = subprocess.run(['curl', '-s', url], capture_output=True, text=True, timeout=10)
        if result.returncode == 0 and result.stdout:
            data = json.loads(result.stdout)
            if 'holidays' in data:
                for h in data['holidays']:
                    date = h.get('date', '')[:10]
                    name = h.get('name', '')
                    if date and name:
                        holidays_2026[date] = name
    except:
        pass
    
    return holidays_2026

# ========== 股票報價（yfinance）==========
def get_stock_prices():
    """取得主要股票報價（yfinance）"""
    stocks = {
        'TSM': {'name': '台積電 ADR', 'symbol': 'TSM'},
        'NVDA': {'name': '輝達', 'symbol': 'NVDA'},
        'AAPL': {'name': '蘋果', 'symbol': 'AAPL'},
        'MSFT': {'name': '微軟', 'symbol': 'MSFT'},
        'GOOGL': {'name': 'Google', 'symbol': 'GOOGL'},
    }
    prices = {}
    for key, info in stocks.items():
        try:
            ticker = yfinance.Ticker(info['symbol'])
            hist = ticker.history(period='2d')
            if not hist.empty:
                current = hist['Close'].iloc[-1]
                prev = hist['Close'].iloc[-2] if len(hist) > 1 else current
                change = current - prev
                pct = (change / prev * 100) if prev > 0 else 0
                prices[key] = {
                    'name': info['name'],
                    'price': round(current, 2),
                    'change': round(change, 2),
                    'pct': round(pct, 2),
                }
        except Exception as e:
            print(f"股票取得失敗 {key}: {e}")
    return prices

# ========== 2. 行事曆（iCloud CalDAV）==========
def get_calendar_events():
    """取得今日+未來30天的行事曆事件（iCloud CalDAV）"""
    events = []
    today = datetime.datetime.now()
    today_start = today.replace(hour=0, minute=0, second=0, microsecond=0)
    future = today + datetime.timedelta(days=30)
    
    # CalDAV REPORT 查詢（不使用 time-range 過濾）
    body = '''<?xml version="1.0" encoding="UTF-8"?>
<C:calendar-query xmlns:D="DAV:" xmlns:C="urn:ietf:params:xml:ns:caldav">
  <D:prop>
    <D:displayname/>
    <C:calendar-data/>
  </D:prop>
  <C:filter>
    <C:comp-filter name="VCALENDAR">
      <C:comp-filter name="VEVENT"/>
    </C:comp-filter>
  </C:filter>
</C:calendar-query>'''
        
    import base64
    for cal_uuid in PRIMARY_CALENDARS:
        url = f"{CALDAV_HOST}{cal_uuid}/"
        req = urllib.request.Request(url, data=body.encode('utf-8'), method='REPORT')
        req.add_header('Content-Type', 'application/xml; charset=utf-8')
        req.add_header('Depth', '1')
        credentials = f'{ICLOUD_USER}:{ICLOUD_APP_PASSWORD}'.encode('utf-8')
        req.add_header('Authorization', 'Basic ' + base64.b64encode(credentials).decode('utf-8'))
        
        try:
            with urllib.request.urlopen(req, timeout=15) as resp:
                content = resp.read().decode('utf-8')
                
                # Extract CDATA blocks then parse VEVENTs
                cdata_blocks = re.findall(r'<!\[CDATA\[(.*?)\]\]>', content, re.DOTALL)
                for block in cdata_blocks:
                    vevents = re.findall(r'BEGIN:VEVENT.+?END:VEVENT', block, re.DOTALL)
                    for vevent in vevents:
                        dtstart = re.search(r'DTSTART[^:]*:([\dT]+)', vevent)
                        summary = re.search(r'SUMMARY[^:]*:([^\r\n]+)', vevent)
                        
                        if dtstart and summary:
                            dtstart_val = dtstart.group(1).strip()
                            summary_val = summary.group(1).strip()
                            
                            clean_date = re.sub(r'[^0-9]', '', dtstart_val[:15])
                            if len(clean_date) >= 8:
                                try:
                                    event_date = datetime.datetime.strptime(clean_date[:8], '%Y%m%d')
                                    if today_start <= event_date <= future:
                                        display_date = f"{clean_date[4:6]}/{clean_date[6:8]}"
                                        events.append((display_date, summary_val))
                                except:
                                    pass
        except Exception as e:
            print(f"行事曆取得失敗 ({cal_uuid}): {e}")
    
    events = list(set(events))
    events.sort(key=lambda x: x[0])
    return events

# ========== 3. Google Tasks ==========
def get_google_tasks():
    """取得 Google Tasks 待辦事項 """
    import urllib.request
    
    # 讀取 token
    try:
        with open('/root/.openclaw/google_tasks_token.json', 'r') as f:
            token_data = json.load(f)
        
        access_token = token_data.get('access_token')
        if not access_token:
            return None
        
        # 檢查 token 是否過期，需要 refresh
        import datetime
        expires_at = token_data.get('expires_at', 0)
        if datetime.datetime.now().timestamp() >= expires_at:
            # refresh token
            refresh_token = token_data.get('refresh_token')
            if refresh_token:
                client_id = '620667525511-qekqk0quvad4v9mdgv3t9p773fsno3r7.apps.googleusercontent.com'
                client_secret = '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'
                
                url = 'https://oauth2.googleapis.com/token'
                payload = {
                    'client_id': client_id,
                    'client_secret': client_secret,
                    'refresh_token': refresh_token,
                    'grant_type': 'refresh_token'
                }
                
                data = json.dumps(payload).encode('utf-8')
                req = urllib.request.Request(url, data=data, headers={'Content-Type': 'application/json'})
                
                with urllib.request.urlopen(req, timeout=30) as response:
                    result = json.loads(response.read().decode('utf-8'))
                    access_token = result['access_token']
                    token_data['access_token'] = access_token
                    token_data['expires_at'] = datetime.datetime.now().timestamp() + result.get('expires_in', 3600)
                    
                    with open('/root/.openclaw/google_tasks_token.json', 'w') as f:
                        json.dump(token_data, f, indent=2)
        
        headers = {'Authorization': f'Bearer {access_token}'}
        
        # 取得 Tasks
        url = 'https://tasks.googleapis.com/tasks/v1/lists/@default/tasks?maxResults=20'
        req = urllib.request.Request(url, headers=headers)
        
        with urllib.request.urlopen(req, timeout=30) as response:
            result = json.loads(response.read().decode('utf-8'))
            tasks = result.get('items', [])
            
            # 分類
            pending = [t for t in tasks if t.get('status') != 'completed']
            completed = [t for t in tasks if t.get('status') == 'completed']
            
            return {
                'pending': pending[:10],  # 最多10項
                'completed': len(completed),
                'total': len(tasks)
            }
            
    except Exception as e:
        print(f'Google Tasks 讀取失敗: {e}')
        return None

# ========== 4. 未讀郵件 ==========
def get_unread_email():
    """取得未讀郵件（macOS Mail + Gmail）"""
    all_emails = []
    
    # --- Mac Mail ---
    mail_script = '''
    tell application "Mail"
        set unreadCount to unread count of inbox
        set recentMessages to every message of inbox whose read status is false
        set output to "未讀數量:" & unreadCount
        repeat with msg in recentMessages
            set msgSubject to subject of msg
            set msgSender to sender of msg
            set output to output & "||" & msgSubject & "||" & msgSender
        end repeat
        return output
    end tell
    '''
    
    try:
        result = subprocess.run(['osascript', '-e', mail_script],
                              capture_output=True, text=True, timeout=30)
        raw = result.stdout.strip()
        if raw:
            parts = raw.split('||')
            count_text = parts[0].replace('未讀數量:', '')
            count = int(count_text) if count_text.isdigit() else 0
            for i in range(1, min(count + 1, 20)):
                if i < len(parts) - 1:
                    all_emails.append({'subject': parts[i], 'sender': parts[i + 1], 'source': 'Mac Mail'})
    except Exception as e:
        print(f'Mac Mail讀取失敗: {e}')
    
    # --- Gmail ---
    try:
        mail = imaplib.IMAP4_SSL('imap.gmail.com')
        mail.login(GMAIL_USER, GMAIL_APP_PASSWORD)
        mail.select('inbox')
        
        status, messages = mail.search(None, 'UNSEEN')
        email_ids = messages[0].split()
        gmail_count = len(email_ids)
        print(f'Gmail 未讀: {gmail_count} 封')
        
        for eid in email_ids[-20:]:
            status, msg_data = mail.fetch(eid, '(RFC822)')
            msg = email.message_from_bytes(msg_data[0][1])
            
            subject, encoding = decode_header(msg['Subject'])[0]
            if isinstance(subject, bytes):
                subject = subject.decode(encoding or 'utf-8', errors='ignore')
            
            sender = msg['From']
            all_emails.append({'subject': subject, 'sender': sender, 'source': 'Gmail'})
        
        mail.logout()
    except Exception as e:
        print(f'Gmail讀取失敗: {e}')
    
    # 合併並去除重複（以主旨判斷）
    seen = set()
    unique_emails = []
    for e in all_emails:
        key = e['subject'][:50]
        if key not in seen:
            seen.add(key)
            unique_emails.append(e)
    
    total = len(unique_emails)
    raw_output = f'未讀數量:{total}'
    for e in unique_emails:
        raw_output += f"||{e['subject']}||{e['sender']}"
    
    return raw_output

def parse_emails(raw):
    """解析未讀郵件"""
    if not raw:
        return {'count': 0, 'emails': []}
    
    try:
        parts = raw.split('||')
        count_text = parts[0].replace('未讀數量:', '')
        count = int(count_text) if count_text.isdigit() else 0
        
        emails = []
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                subject = parts[i]
                sender = parts[i + 1] if i + 1 < len(parts) else ''
                emails.append({'subject': subject, 'sender': sender})
        
        return {'count': count, 'emails': emails}
    except:
        return {'count': 0, 'emails': []}

# ========== 生成報告 ==========
def generate_report(weather_data, events, tasks_data, emails_data, today):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 標題 =====
    title = doc.add_heading('☀️ 晨間摘要', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(today.strftime('%Y年%m月%d日'))
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # ===== 1. 天氣資訊 =====
    h1 = doc.add_heading('一、天氣資訊', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    if weather_data:
        curr = weather_data['current']
        tom = weather_data['tomorrow']

        # 今日天氣表格
        table1 = doc.add_table(rows=5, cols=4)
        table1.style = 'Table Grid'
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['項目', '數據', '項目', '數據']
        for i, h in enumerate(headers):
            make_header_cell(table1.rows[0].cells[i], h)

        data = [
            ['天氣狀況', curr['desc'], '體感溫度', curr['feels']],
            ['目前溫度', curr['temp'], '濕度', curr['humidity']],
            ['降雨量', curr['rain'], '風速', curr['wind']],
            ['紫外線', curr['uv'], '能見度', curr['vis']],
        ]
        for ri, row_data in enumerate(data):
            row = table1.rows[ri + 1]
            for ci, val in enumerate(row_data):
                add_data_cell(row.cells[ci], val)

        doc.add_paragraph()

        # 明日天氣
        p = doc.add_paragraph()
        p.add_run(f'🌤️ 明日天氣：{tom["desc"]}，氣溫 {tom["min"]}°C ~ {tom["max"]}°C')
        p.runs[0].font.size = Pt(11)
    else:
        p = doc.add_paragraph()
        p.add_run('⚠️ 無法取得天氣資料')
        p.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # ===== 2. 行事曆（含台灣節日）=====
    h2 = doc.add_heading('二、行事曆', level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 取得台灣節日
    taiwan_holidays = get_taiwan_holidays()
    today_str = today.strftime('%Y-%m-%d')
    upcoming_holidays = {k: v for k, v in taiwan_holidays.items() if k >= today_str[:8] + '01' and k <= (today + datetime.timedelta(days=60)).strftime('%Y-%m-%d')}

    # 行事曆
    p = doc.add_paragraph()
    p.add_run('📅 今日及未來30天行程\n').bold = True
    p.runs[0].font.size = Pt(11)

    if events:
        for date_str, summary in events:
            p.add_run(f'• {date_str} — {summary}\n')
    else:
        p.add_run('• 沒有行程安排')

    # 台灣節日
    if upcoming_holidays:
        p.add_run('\n🇹🇼 台灣節日\n').bold = True
        for date_str, name in sorted(upcoming_holidays.items()):
            display_date = f"{date_str[5:7]}/{date_str[8:10]}"
            p.add_run(f'• {display_date} — {name}\n')

    doc.add_paragraph()

    # ===== 3. Google Tasks =====
    h3 = doc.add_heading('三、Google Tasks 待辦事項', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    if tasks_data:
        pending = tasks_data.get('pending', [])
        completed = tasks_data.get('completed', 0)
        total = tasks_data.get('total', 0)

        p = doc.add_paragraph()
        p.add_run(f'📋 待辦事項（共 {total} 項，已完成 {completed} 項）\n').bold = True
        p.runs[0].font.size = Pt(11)

        if pending:
            for task in pending:
                title = task.get('title', '（無標題）')
                due = task.get('due', '')
                if due:
                    due_str = due[:10]
                    p.add_run(f'• {title}（到期：{due_str}）\n')
                else:
                    p.add_run(f'• {title}\n')
        else:
            p.add_run('• 目前沒有待辦事項')
    else:
        p = doc.add_paragraph()
        p.add_run('⚠️ 無法取得 Google Tasks（請確認授權）')
        p.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # ===== 4. 未讀郵件 =====
    h4 = doc.add_heading('四、未讀郵件', level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    count = emails_data['count']
    emails = emails_data['emails']

    p = doc.add_paragraph()
    p.add_run(f'📬 未讀郵件：{count} 封\n').bold = True
    p.runs[0].font.size = Pt(11)

    if emails:
        for email in emails:
            p.add_run(f'• {email["subject"]}\n')
            p.add_run(f'  寄件人：{email["sender"]}\n')
    else:
        p.add_run('• 沒有未讀郵件')

    doc.add_paragraph()

    # ===== 頁尾 =====
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'小安助理 · {today.strftime("%Y年%m月%d日")}')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ===== 儲存 =====
    output_path = f'/root/.openclaw/reports/daily/晨間摘要_{today.strftime("%Y%m%d_%H%M")}.docx'
    doc.save(output_path)
    print(f'✅ 已儲存: {output_path}')
    return output_path

def send_to_telegram(doc_path, today):
    caption = f"☀️ 晨間摘要_{today.strftime('%Y年%m月%d日')}\n\n天氣 · 行事曆 · Tasks · Gmail"
    
    # 重試機制：最多嘗試3次
    for attempt in range(3):
        result = subprocess.run(['curl', '-s', '-X', 'POST',
            f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
            '-F', f'chat_id={CHAT_ID}',
            '-F', f'document=@{doc_path}',
            '-F', f'caption={caption}'],
            capture_output=True, timeout=30)
        
        if result.returncode == 0:
            response_data = json.loads(result.stdout)
            if response_data.get('ok'):
                print('✅ 已發送到 Telegram')
                return True
        
        print(f'⚠️ 發送失敗，第{attempt+1}次嘗試...')
        time.sleep(2)
    
    print('❌ 發送失敗（已重試3次）')
    return False

def main():
    today = datetime.datetime.now()
    print("=" * 50)
    print("📋 開始生成晨間摘要報告...")
    print("=" * 50)

    # 1. 天氣
    print("\n☀️ 取得天氣資訊...")
    weather_data = get_weather()
    if weather_data:
        print(f"   ✅ 天氣取得成功")

    # 2. 行事曆
    print("\n📅 取得行事曆...")
    events = get_calendar_events()
    print(f"   ✅ 找到 {len(events)} 個行程")

    # 3. Google Tasks
    print("\n📋 取得 Google Tasks...")
    tasks_data = get_google_tasks()
    if tasks_data:
        print(f"   ✅ Tasks 取得成功：{tasks_data['total']} 項待辦")
    else:
        print("   ⚠️ 無法取得 Google Tasks")

    # 4. 未讀郵件
    print("\n📬 取得未讀郵件...")
    raw_emails = get_unread_email()
    emails_data = parse_emails(raw_emails)
    print(f"   ✅ 未讀郵件：{emails_data['count']} 封")

    # 生成報告
    print("\n" + "=" * 50)
    print("📝 生成報告...")
    doc_path = generate_report(weather_data, events, tasks_data, emails_data, today)
    send_to_telegram(doc_path, today)
    print("✅ 完成！")

if __name__ == '__main__':
    main()
