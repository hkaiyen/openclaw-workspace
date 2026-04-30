#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
每日健康新知 - 自動執行腳本
==========================================

排程：每天早上 08:00

流程：
1. 小安使用 Groq API 搜尋健康新知
2. 整理成簡報
3. 潔咪生成圖片
4. 發送到 Telegram

"""

import subprocess
import json
import time
import urllib.request
import datetime
import os
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== 常數 =====
BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
REPORT_DIR = '/root/.openclaw/reports/daily'
IMAGE_DIR = '/root/.openclaw/media'
GEMINI_API_KEY = 'AIzaSyCrUO2AlTi69NtKNdwDgvBaYppe2EU7-Jw'
GROQ_API_KEY = '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'

# ===== 健康主題關鍵字 =====
HEALTH_TOPICS = [
    "健康", "養生", "飲食", "運動", "睡眠", "心理", "慢性病", "癌症",
    "心血管", "糖尿病", "高血壓", "骨質疏鬆", "肥胖", "減肥", "美容",
    "老化", "免疫", "維他命", "保健品", "食品安全"
]

def search_health_news():
    """使用 Groq API（小安）搜尋健康新知"""
    print("🔍 小安正在搜尋健康新知...")
    
    import random
    topic = random.choice(HEALTH_TOPICS)
    
    system_prompt = """你是川寶投顧的小安，是一位專業的健康顧問。

⚠️ 所有輸出必須使用繁體中文！

請搜尋並整理今日最新的健康新知，主題是「{topic}」。

請提供：
1. 3-5則重要健康資訊
2. 每則資訊的標題和重點內容
3. 實用的健康建議

請用條列式呈現，內容要實用且有根據。""".format(topic=topic)
    
    payload = {
        "model": "llama-3.3-70b-versatile",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"請搜尋今日關於「{topic}」的健康新知，並整理成3-5則簡短的重點資訊。"}
        ],
        "temperature": 0.7,
        "max_tokens": 2048
    }
    
    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json=payload,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            content = result['choices'][0]['message']['content']
            print(f"✅ 小安健康新知取得成功")
            return content, topic
        else:
            print(f"❌ 小安搜尋失敗: {response.status_code}")
            return f"今日健康新知：{topic}相關資訊整理中，請稍後再試。", topic
            
    except Exception as e:
        print(f"❌ 小安搜尋錯誤: {e}")
        return f"今日健康新知：{topic}相關資訊整理中，請稍後再試。", topic

def generate_gemini_image(prompt, output_path):
    """使用 Gemini API 生成圖片"""
    import base64
    
    print(f"🎨 正在生成圖片...")
    
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.1-flash-image-preview:generateContent?key={GEMINI_API_KEY}"
    
    payload = {
        "contents": [{
            "parts": [{"text": f"請生成一張精美的圖片，主題是：{prompt}"}]
        }]
    }
    
    try:
        data = json.dumps(payload).encode('utf-8')
        req = urllib.request.Request(url, data=data, headers={'Content-Type': 'application/json'})
        
        with urllib.request.urlopen(req, timeout=300) as response:
            result = json.loads(response.read().decode('utf-8'))
        
        if 'candidates' in result:
            parts = result['candidates'][0].get('content', {}).get('parts', [])
            for part in parts:
                if 'inlineData' in part:
                    image_data = part['inlineData']['data']
                    image_bytes = base64.b64decode(image_data)
                    with open(output_path, 'wb') as f:
                        f.write(image_bytes)
                    print(f"✅ 圖片已生成: {output_path}")
                    return True
        
        print(f"❌ 圖片生成失敗")
        return False
        
    except Exception as e:
        print(f"❌ 圖片生成錯誤: {e}")
        return False

def send_telegram_photo(file_path, caption):
    """發送圖片到Telegram"""
    cmd = [
        'curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{BOT_TOKEN}/sendPhoto',
        '-F', f'chat_id={CHAT_ID}',
        '-F', f'photo=@{file_path}',
        '-F', f'caption={caption}'
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    try:
        response = json.loads(result.stdout)
        return response.get('ok', False)
    except:
        return False

def generate_word_report(news_content, topic, image_path):
    """生成Word報告"""
    doc = Document()
    
    # 標題
    title = doc.add_heading('💚 每日健康新知', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_str = datetime.datetime.now().strftime('%Y年%m月%d日')
    doc.add_paragraph(f'日期：{date_str}')
    doc.add_paragraph(f'主題：{topic}')
    doc.add_paragraph('')
    
    # 圖片
    if image_path and os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
    
    doc.add_paragraph('')
    
    # 健康新知內容（小安整理）
    doc.add_heading('📰 小安的健康資訊整理', level=1)
    doc.add_paragraph(news_content)
    doc.add_paragraph('')
    
    # 結尾
    doc.add_paragraph('—' * 20)
    doc.add_paragraph(f'川寶投顧小安助理 | {date_str}')
    doc.add_paragraph('本報告僅供參考，不構成醫療建議。如有健康疑慮，請諮詢專業醫師。')
    
    # 儲存
    os.makedirs(REPORT_DIR, exist_ok=True)
    today_str = datetime.datetime.now().strftime('%Y%m%d')
    report_path = os.path.join(REPORT_DIR, f'健康新知_{today_str}.docx')
    doc.save(report_path)
    print(f"✅ Word報告已生成: {report_path}")
    return report_path

def send_telegram_document(file_path, caption):
    """發送文件到Telegram"""
    cmd = [
        'curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
        '-F', f'chat_id={CHAT_ID}',
        '-F', f'document=@{file_path}',
        '-F', f'caption={caption}'
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    try:
        response = json.loads(result.stdout)
        return response.get('ok', False)
    except:
        return False

# ===== 主程式 =====
def main():
    print("=" * 60)
    print("每日健康新知 - 自動執行")
    print("=" * 60)
    print(f"時間：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    os.makedirs(REPORT_DIR, exist_ok=True)
    os.makedirs(IMAGE_DIR, exist_ok=True)
    
    # 步驟1：小安搜尋健康新知（Groq API）
    news_content, topic = search_health_news()
    
    # 步驟2：生成圖片
    today_str = datetime.datetime.now().strftime('%Y%m%d')
    image_path = os.path.join(IMAGE_DIR, f'health_{today_str}.png')
    
    image_prompt = f"健康養生主題，{topic}相關，清新自然風格，專業醫學插圖"
    if not generate_gemini_image(image_prompt, image_path):
        image_path = None
    
    # 步驟3：生成Word報告
    docx_path = generate_word_report(news_content, topic, image_path)
    
    # 步驟4：發送到Telegram
    if docx_path and os.path.exists(docx_path):
        if image_path and os.path.exists(image_path):
            send_telegram_photo(image_path, f"💚 每日健康新知：{topic}")
        
        caption = f"💚 每日健康新知 {datetime.datetime.now().strftime('%Y年%m月%d日')} | 川寶投顧"
        if send_telegram_document(docx_path, caption):
            print("\n✅ 已發送到Telegram")
        else:
            print("\n❌ 發送失敗")
    
    print("\n" + "=" * 60)
    print("完成！")
    print("=" * 60)

if __name__ == '__main__':
    main()
