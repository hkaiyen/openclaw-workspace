#!/usr/bin/python3
"""
天氣報告生成腳本
- 資料來源：wttr.in
- 格式：與市場統計報表一致（黑字白底、置中）
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import datetime
import json

# ========== 工具函數 ==========

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def make_header_cell(cell, text, bg='FFFFFF'):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.font.size = Pt(10)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def add_data_cell(cell, text, is_positive=None):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_positive is True:
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    elif is_positive is False:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

# ========== 主程式 ==========

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d')

    # 取得天氣資料
    result = subprocess.run(['curl', '-s', 'wttr.in/Taipei?format=j1'], capture_output=True, text=True, timeout=15)
    d = json.loads(result.stdout)
    curr = d['current_condition'][0]
    tom = d['weather'][1]
    tom_hourly = tom.get('hourly', [])
    tom_desc = tom_hourly[0]['weatherDesc'][0]['value'] if (tom_hourly and tom_hourly[0].get('weatherDesc')) else 'N/A'

    doc = Document()

    # 頁面設定
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 標題 =====
    title = doc.add_heading('☀️ 天氣報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.runs[0]
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('台北市天氣預報')
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    sr.font.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run(f'資料日期：{today.strftime("%Y年%m月%d日")}')
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== 今日天氣 =====
    h1 = doc.add_heading('一、今日天氣', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['項目', '數據', '項目', '數據']
    for i, h in enumerate(headers):
        make_header_cell(table1.rows[0].cells[i], h)

    data = [
        ['天氣狀況', curr['weatherDesc'][0]['value'], '體感溫度', f"{curr['FeelsLikeC']}°C"],
        ['目前溫度', f"{curr['temp_C']}°C", '濕度', f"{curr['humidity']}%"],
        ['降雨量', f"{curr.get('precipMM', '0')}mm", '風速', f"{curr['windspeedKmph']} km/h"],
        ['紫外線指數', str(curr['uvIndex']), '能見度', f"{curr['visibility']} km"],
    ]
    for ri, row_data in enumerate(data):
        row = table1.rows[ri + 1]
        for ci, val in enumerate(row_data):
            add_data_cell(row.cells[ci], val)

    doc.add_paragraph()

    # ===== 明日天氣 =====
    h2 = doc.add_heading('二、明日天氣預報', level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table2 = doc.add_table(rows=3, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['項目', '數據', '項目', '數據']):
        make_header_cell(table2.rows[0].cells[i], h)

    row = table2.rows[1]
    add_data_cell(row.cells[0], '天氣狀況')
    add_data_cell(row.cells[1], tom_desc)
    add_data_cell(row.cells[2], '紫外線指數')
    add_data_cell(row.cells[3], tom['uvIndex'])

    row2 = table2.rows[2]
    add_data_cell(row2.cells[0], '最高溫度')
    add_data_cell(row2.cells[1], f"{tom['maxtempC']}°C")
    add_data_cell(row2.cells[2], '最低溫度')
    add_data_cell(row2.cells[3], f"{tom['mintempC']}°C")

    doc.add_paragraph()

    # ===== 小安建議 =====
    h3 = doc.add_heading('三、貼心提醒', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    suggestions = [
        ('☀️ 今日', '天氣不錯，適合出門！體感溫度較高，外出記得補充水分。'),
        ('☂️ 明日', f"天氣：{tom_desc}，出門記得帶把傘。"),
        ('🧴 防曬', '今日紫外線指數低，但戶外活動仍建議適度防曬。'),
    ]
    for label, content in suggestions:
        p = doc.add_paragraph()
        p.add_run(f'{label}：').bold = True
        p.add_run(content)

    doc.add_paragraph()

    # 資料來源
    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = src.add_run('資料來源：天氣預報服務 wttr.in')
    sr.font.size = Pt(9)
    sr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    sr.font.italic = True

    # 聲明
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_r = disc.add_run('本報告僅供參考，出門前請以中央氣象局公告為準')
    disc_r.font.size = Pt(9)
    disc_r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    disc_r.font.italic = True

    # 頁尾
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 小安製')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    output_path = f'/root/.openclaw/reports/daily/天氣報告_{date_str}.docx'
    doc.save(output_path)
    return output_path

def send_to_telegram(file_path):
    bot_token = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
    chat_id = '8779713208'
    today = datetime.datetime.now()
    caption = f"☀️ 天氣報告_{today.strftime('%Y年%m月%d日')}\n\n台北市天氣預報\n小安製"

    result = subprocess.run(
        ['curl', '-s', '-X', 'POST',
         f'https://api.telegram.org/bot{bot_token}/sendDocument',
         '-F', f'chat_id={chat_id}',
         '-F', f'document=@{file_path}',
         '-F', f'caption={caption}'],
        capture_output=True, text=True, timeout=30
    )
    response = json.loads(result.stdout)
    return response.get('ok', False)

if __name__ == '__main__':
    print(f'[{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}] 開始生成天氣報告...')
    try:
        report_path = generate_report()
        print(f'📄 報告已生成: {report_path}')
        if send_to_telegram(report_path):
            print('✅ 已發送到 Telegram')
        else:
            print('❌ 發送失敗')
        print(f'[{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}] 任務完成！')
    except Exception as e:
        print(f'❌ 錯誤: {e}')
