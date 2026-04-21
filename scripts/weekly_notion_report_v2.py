#!/usr/bin/python3
"""
Notion 每週全方位個人報告生成器 - 四助理協力版
由小安、拉瑪、千問、小歐各自撰寫獨立報告
"""

import requests
import urllib.request
import json
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import os

# ========== 基本設定 ==========
NOTION_TOKEN = "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"
TELEGRAM_BOT_TOKEN = "8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw"
TELEGRAM_CHAT_ID = "8779713208"
GROQ_API_KEY = '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'

# 資料庫 ID
DATABASE_IDS = {
    'docs': '927bb7a7-c7e5-4411-9c3a-b9f36e9134e6',
    'business': 'c8db2442-c719-4873-b73b-25dbe106497d',
    'warehouse': '1336a4ae-1760-81b4-ad2d-d39cd28ed090',
    'ideas': '8b347be8-3c42-4e0b-a41c-6dd7585f7f84',
    'ref': 'b1e42614-8e73-42d8-a731-639dad9e1f6a',
    'weekly_review': '1716a4ae-1760-8162-827b-d4e43e695244',
    'diary': '1716a4ae-1760-81cb-aada-dac840549da5',
    'english': '4d7e93eb-2d0b-4af6-ba2f-90856b3f6368',
}

# 四助理設定
ASSISTANTS = {
    'xiaoan': {
        'name': '小安',
        'model': 'xiaoan',
        'emoji': '🐰',
        'personality': '務實、嚴謹、有條理'
    },
    'lama': {
        'name': '拉瑪',
        'model': 'llama-3.3-70b-versatile',
        'emoji': '🐰',
        'personality': '分析深入、數據導向'
    },
    'qianwen': {
        'name': '千問',
        'model': 'qwen/qwen3-32b',
        'emoji': '🐰',
        'personality': '宏觀策略、系統思考'
    },
    'xiaoou': {
        'name': '小歐',
        'model': 'openrouter/free',
        'emoji': '🐰',
        'personality': '創意發散、趨勢洞察'
    }
}

# ========== Notion API 函數 ==========
def get_headers():
    return {
        'Authorization': f'Bearer {NOTION_TOKEN}',
        'Notion-Version': '2022-06-28'
    }

def get_block_children(block_id):
    try:
        req = urllib.request.Request(
            f'https://api.notion.com/v1/blocks/{block_id}/children',
            headers=get_headers()
        )
        resp = urllib.request.urlopen(req, timeout=15)
        return json.loads(resp.read())
    except:
        return {'results': []}

def get_page_content(page_id, max_blocks=100):
    blocks_data = get_block_children(page_id)
    content = []
    for b in blocks_data.get('results', [])[:max_blocks]:
        t = b.get('type', '')
        if t in ['paragraph', 'bulleted_list_item', 'numbered_list_item', 'heading_1', 'heading_2', 'heading_3']:
            texts = b.get(t, {}).get('rich_text', [])
            txt = ''.join(x.get('plain_text', '') for x in texts).strip()
            if txt:
                content.append(txt)
        elif t == 'to_do':
            texts = b.get(t, {}).get('rich_text', [])
            txt = ''.join(x.get('plain_text', '') for x in texts).strip()
            checked = b.get(t, {}).get('checked', False)
            content.append(f"[{'v' if checked else '☐'}] {txt}")
        elif t == 'callout':
            texts = b.get(t, {}).get('rich_text', [])
            txt = ''.join(x.get('plain_text', '') for x in texts).strip()
            content.append(f"📌 {txt}")
    return content

def query_database(db_id, page_size=10):
    try:
        data = {"page_size": page_size}
        req = urllib.request.Request(
            f'https://api.notion.com/v1/databases/{db_id}/query',
            headers=get_headers(),
            data=json.dumps(data).encode('utf-8'),
            method='POST'
        )
        resp = urllib.request.urlopen(req, timeout=15)
        return json.loads(resp.read())
    except:
        return {'results': []}

# ========== 讀取上週複盤 ==========
def get_last_week_review():
    """讀取上週的複盤資料（不是本週）"""
    weekly_data = query_database(DATABASE_IDS['weekly_review'], page_size=10)
    entries = []
    
    for r in weekly_data.get('results', []):
        page_id = r.get('id', '')
        created = r.get('created_time', '')[:10]
        page_content = get_page_content(page_id, max_blocks=50)
        if page_content and len(page_content) > 3:
            entries.append({'date': created, 'content': page_content})
    
    # 按日期排序（最新在前面）
    entries.sort(key=lambda x: x['date'], reverse=True)
    
    # 返回倒數第二筆（上週的）
    if len(entries) >= 2:
        return entries[1]  # 上週
    elif len(entries) == 1:
        return entries[0]  # 只有一筆就用這筆
    return None

# ========== 讀取花園資料 ==========
def get_all_garden_data():
    """讀取花園資料庫所有內容"""
    all_data = {
        'ideas': [], 'business': [], 'warehouse': [],
        'docs': [], 'ref': [], 'diary': [], 'english': []
    }
    
    for key, db_id in DATABASE_IDS.items():
        if key in all_data:
            data = query_database(db_id, page_size=20)
            for r in data.get('results', []):
                page_id = r.get('id', '')
                content = get_page_content(page_id, max_blocks=30)
                if content:
                    all_data[key].append({
                        'id': page_id,
                        'content': content
                    })
    
    return all_data

# ========== Groq API ==========
def groq_query(model, prompt, system="你是專業的助理。"):
    """使用 Groq API 查詢"""
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.5,
        "max_tokens": 3000
    }
    try:
        result = subprocess.run(
            ['curl', '-s', '-X', 'POST',
             'https://api.groq.com/openai/v1/chat/completions',
             '-H', f'Authorization: Bearer {GROQ_API_KEY}',
             '-H', 'Content-Type: application/json',
             '-d', json.dumps(payload)],
            capture_output=True, text=True, timeout=60
        )
        data = json.loads(result.stdout)
        return data['choices'][0]['message']['content']
    except:
        return None

# ========== 文件格式化 ==========
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def make_header_cell(cell, text, bg='1F497D'):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    set_cell_bg(cell, bg)

def add_data_cell(cell, text):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(10)

# ========== 各助理報告生成 ==========
def generate_xiaoan_report(last_week_review, all_data, today_str):
    """🐰 小安報告：務實嚴謹風格"""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    
    # 標題
    title = doc.add_heading('🐰 小安｜每週全方位個人報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(20)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{today_str}｜小安風格：務實嚴謹').font.size = Pt(11)
    
    doc.add_paragraph()
    
    # 上週複盤
    h = doc.add_heading('📋 上週複盤回顧', 1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    if last_week_review:
        p = doc.add_paragraph()
        p.add_run(f"📅 日期：{last_week_review['date']}").bold = True
        for c in last_week_review['content'][:10]:
            doc.add_paragraph(f"• {c[:200]}")
    else:
        doc.add_paragraph("⚠️ 無上週複盤資料")
    
    doc.add_paragraph()
    
    # 小安分析
    h = doc.add_heading('📊 小安分析與建議', 1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    # 根據資料生成建議
    ideas_text = ' '.join(''.join(i['content'][:100]) for i in all_data['ideas'][:5])
    business_text = ' '.join(''.join(i['content'][:100]) for i in all_data['business'][:5])
    
    prompt = f"""你是小安，請根據以下資料，寫一份簡短的每週分析報告（300字以內）：

上週複盤重點：{last_week_review['content'][:200] if last_week_review else '無'}
點子收集：{ideas_text[:300]}
商務任務：{business_text[:300]}

請用小安的風格（務實嚴謹、有條理）寫出：
1. 上週執行檢視
2. 本週重點建議
3. 風險警示

請用繁體中文回覆。"""
    
    response = groq_query("llama-3.3-70b-versatile", prompt,
                         system="你是小安，務實嚴謹的助理。")
    
    if response:
        for line in response.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip()[:200])
    
    # 頁尾
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'小安助理 · {today_str}')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    
    output_path = f'/root/.openclaw/reports/weekly/每週報告_小安_XiaoAn_{today_str.replace("年", "").replace("月", "").replace("日", "")}.docx'
    doc.save(output_path)
    return output_path

def generate_lama_report(last_week_review, all_data, today_str):
    """🐰 拉瑪報告：分析深入風格"""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    
    title = doc.add_heading('🐰 拉瑪｜每週全方位個人報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(20)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{today_str}｜拉瑪風格：分析深入').font.size = Pt(11)
    
    doc.add_paragraph()
    
    # 上週複盤
    h = doc.add_heading('📋 上週複盤回顧', 1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    if last_week_review:
        p = doc.add_paragraph()
        p.add_run(f"📅 日期：{last_week_review['date']}").bold = True
        for c in last_week_review['content'][:10]:
            doc.add_paragraph(f"• {c[:200]}")
    else:
        doc.add_paragraph("⚠️ 無上週複盤資料")
    
    doc.add_paragraph()
    
    # 拉瑪分析
    h = doc.add_heading('📊 拉瑪分析與建議', 1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    ideas_text = ' '.join(''.join(i['content'][:100]) for i in all_data['ideas'][:5])
    business_text = ' '.join(''.join(i['content'][:100]) for i in all_data['business'][:5])
    
    prompt = f"""你是拉瑪，請根據以下資料，寫一份深入的每週分析報告（300字以內）：

上週複盤重點：{last_week_review['content'][:200] if last_week_review else '無'}
點子收集：{ideas_text[:300]}
商務任務：{business_text[:300]}

請用拉瑪的風格（分析深入、數據導向）寫出：
1. 量化檢視
2. 深度分析
3. 具體建議

請用繁體中文回覆。"""
    
    response = groq_query("llama-3.3-70b-versatile", prompt,
                         system="你是拉瑪，分析深入的助理。")
    
    if response:
        for line in response.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip()[:200])
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'拉瑪助理 · {today_str}')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    
    output_path = f'/root/.openclaw/reports/weekly/每週報告_拉瑪_Lama_{today_str.replace("年", "").replace("月", "").replace("日", "")}.docx'
    doc.save(output_path)
    return output_path

def generate_qianwen_report(last_week_review, all_data, today_str):
    """🐰 千問報告：宏觀策略風格"""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    
    title = doc.add_heading('🐰 千問｜每週全方位個人報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(20)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{today_str}｜千問風格：宏觀策略').font.size = Pt(11)
    
    doc.add_paragraph()
    
    # 上週複盤
    h = doc.add_heading('📋 上週複盤回顧', 1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    if last_week_review:
        p = doc.add_paragraph()
        p.add_run(f"📅 日期：{last_week_review['date']}").bold = True
        for c in last_week_review['content'][:10]:
            doc.add_paragraph(f"• {c[:200]}")
    else:
        doc.add_paragraph("⚠️ 無上週複盤資料")
    
    doc.add_paragraph()
    
    # 千問分析
    h = doc.add_heading('📊 千問分析與建議', 1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    ideas_text = ' '.join(''.join(i['content'][:100]) for i in all_data['ideas'][:5])
    business_text = ' '.join(''.join(i['content'][:100]) for i in all_data['business'][:5])
    
    prompt = f"""你是千問，請根據以下資料，寫一份宏觀的每週策略報告（300字以內）：

上週複盤重點：{last_week_review['content'][:200] if last_week_review else '無'}
點子收集：{ideas_text[:300]}
商務任務：{business_text[:300]}

請用千問的風格（宏觀策略、系統思考）寫出：
1. 大局觀察
2. 策略建議
3. 長期建議

請用繁體中文回覆。"""
    
    response = groq_query("qwen/qwen3-32b", prompt,
                         system="你是千問，宏觀策略的助理。")
    
    if response:
        for line in response.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip()[:200])
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'千問助理 · {today_str}')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    
    output_path = f'/root/.openclaw/reports/weekly/每週報告_千問_QianWen_{today_str.replace("年", "").replace("月", "").replace("日", "")}.docx'
    doc.save(output_path)
    return output_path

def generate_xiaoou_report(last_week_review, all_data, today_str):
    """🐰 小歐報告：創意發散風格"""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    
    title = doc.add_heading('🐰 小歐｜每週全方位個人報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(20)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{today_str}｜小歐風格：創意發散').font.size = Pt(11)
    
    doc.add_paragraph()
    
    # 上週複盤
    h = doc.add_heading('📋 上週複盤回顧', 1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    if last_week_review:
        p = doc.add_paragraph()
        p.add_run(f"📅 日期：{last_week_review['date']}").bold = True
        for c in last_week_review['content'][:10]:
            doc.add_paragraph(f"• {c[:200]}")
    else:
        doc.add_paragraph("⚠️ 無上週複盤資料")
    
    doc.add_paragraph()
    
    # 小歐分析
    h = doc.add_heading('📊 小歐分析與建議', 1)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    ideas_text = ' '.join(''.join(i['content'][:100]) for i in all_data['ideas'][:5])
    business_text = ' '.join(''.join(i['content'][:100]) for i in all_data['business'][:5])
    
    prompt = f"""你是小歐，請根據以下資料，寫一份創意的每週觀點報告（300字以內）：

上週複盤重點：{last_week_review['content'][:200] if last_week_review else '無'}
點子收集：{ideas_text[:300]}
商務任務：{business_text[:300]}

請用小歐的風格（創意發散、趨勢洞察）寫出：
1. 創意觀點
2. 趨勢觀察
3. 突破建議

請用繁體中文回覆。"""
    
    response = groq_query("llama-3.3-70b-versatile", prompt,
                         system="你是小歐，創意發散的助理。")
    
    if response:
        for line in response.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip()[:200])
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'小歐助理 · {today_str}')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    
    output_path = f'/root/.openclaw/reports/weekly/每週報告_小歐_XiaoOu_{today_str.replace("年", "").replace("月", "").replace("日", "")}.docx'
    doc.save(output_path)
    return output_path

# ========== 發送到 Telegram ==========
def send_to_telegram(docx_path, assistant_name):
    """發送報告到 Telegram"""
    caption = f"📊 每週全方位報告 - {assistant_name}\n{datetime.now().strftime('%Y年%m月%d日')}"
    
    cmd = [
        'curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendDocument',
        '-F', f'chat_id={TELEGRAM_CHAT_ID}',
        '-F', f'document=@{docx_path}',
        '-F', f'caption={caption}'
    ]
    
    result = subprocess.run(cmd, capture_output=True, text=True)
    try:
        response = json.loads(result.stdout)
        return response.get('ok', False)
    except:
        return False

# ========== 主程式 ==========
def main():
    print("=" * 60)
    print("🐰 四助理 - 每週全方位報告生成器")
    print("=" * 60)
    
    today = datetime.now()
    today_str = today.strftime('%Y年%m月%d日')
    date_id = today.strftime('%Y%m%d')
    
    try:
        # 讀取上週複盤
        print("\n📋 讀取上週複盤資料...")
        last_week_review = get_last_week_review()
        if last_week_review:
            print(f"   ✅ 找到上週複盤：{last_week_review['date']}")
        else:
            print("   ⚠️ 無上週複盤資料")
        
        # 讀取花園資料
        print("\n📚 讀取花園資料...")
        all_data = get_all_garden_data()
        print(f"   ✅ 讀取完成")
        
        reports = []
        
        # 小安報告
        print("\n🐰 生成小安報告...")
        path = generate_xiaoan_report(last_week_review, all_data, today_str)
        reports.append(('小安', path))
        print(f"   ✅ 已儲存: {path}")
        
        # 拉瑪報告
        print("\n🐰 生成拉瑪報告...")
        path = generate_lama_report(last_week_review, all_data, today_str)
        reports.append(('拉瑪', path))
        print(f"   ✅ 已儲存: {path}")
        
        # 千問報告
        print("\n🐰 生成千問報告...")
        path = generate_qianwen_report(last_week_review, all_data, today_str)
        reports.append(('千問', path))
        print(f"   ✅ 已儲存: {path}")
        
        # 小歐報告
        print("\n🐰 生成小歐報告...")
        path = generate_xiaoou_report(last_week_review, all_data, today_str)
        reports.append(('小歐', path))
        print(f"   ✅ 已儲存: {path}")
        
        # 發送到 Telegram
        print("\n📤 發送到 Telegram...")
        for name, path in reports:
            if send_to_telegram(path, name):
                print(f"   ✅ {name}報告已發送")
            else:
                print(f"   ❌ {name}報告發送失敗")
        
        print("\n" + "=" * 60)
        print("✅ 全部完成！")
        print("=" * 60)
        
    except Exception as e:
        print(f"\n❌ 錯誤：{e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
