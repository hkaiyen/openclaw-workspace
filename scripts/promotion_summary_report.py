#!/usr/bin/python3
"""
促銷活動總整理報告 - 五助理協力版（使用 Tavily 搜尋）
- 小咪、小歐、千問、撈仔、拉瑪 同時使用 Tavily 搜尋
- 小安最後彙整、去除重複、生成報告
- 發送到Telegram
"""

import subprocess
import datetime
import json
import os
import sys

# ========== 常數 ==========
BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
REPORTS_DIR = '/root/.openclaw/workspace/reports'

# 五助理搜尋關鍵字
SEARCH_TASKS = {
    "小咪": "2026年5月 台灣 母親節 餐廳優惠 吃到飽 聚餐",
    "小歐": "2026年5月 台灣 銀行信用卡 報稅優惠 支付回饋 LINE Pay",
    "千問": "2026年5月 台灣 百貨公司 母親節 skm points 滿額贈",
    "撈仔": "2026年5月 台灣 超市 超商 全聯 全家 7-11 優惠",
    "拉瑪": "2026年5月 台灣 電商 線上購物 折扣 促銷活動"
}

# 小安靜態資料（5月母親節檔期）
XIAOAN_ITEMS = [
    {"brand": "家樂福", "category": "超市", "campaign": "每月5號會員日", "period": "每月5日～6/5", "details": "OPENPOINT折200元"},
    {"brand": "LINE Pay", "category": "支付", "campaign": "購好券", "period": "即日起～6/30", "details": "最高30% LINE POINTS回饋"},
    {"brand": "LINE Pay×Klook", "category": "支付", "campaign": "找體驗", "period": "即日起～6/30", "details": "景點門票平日10%、週末18%"},
    {"brand": "大全聯", "category": "超市", "campaign": "母親節預購", "period": "母親節檔期", "details": "Gogoro回饋6,000元"},
    {"brand": "新光三越", "category": "百貨", "campaign": "母親節檔期", "period": "4/1-5/10", "details": "滿5,000贈600點"},
    {"brand": "饗食天堂", "category": "餐飲", "campaign": "母親節套餐", "period": "即日起～5/18", "details": "外帶9道+蛋糕4,688元"},
    {"brand": "大米義式餐廳", "category": "餐飲", "campaign": "母親節優惠", "period": "5/1-5/10", "details": "龍蝦吃到飽每人2,088元"},
    {"brand": "玉山銀", "category": "銀行", "campaign": "報稅優惠", "period": "5月", "details": "一次付清0.3%、星宇卡0.6%"},
]

def spawn_search_subagent(name, keyword):
    """使用 sessions_spawn 啟動搜尋子代理"""
    from openclaw import sessions_spawn
    
    task = f"""你是{name}，請用 Tavily 搜尋「{keyword}」並回傳 JSON 格式的促銷資料：

JSON格式：
{{
  "source": "{name} (Tavily)",
  "items": [
    {{"brand": "品牌名", "category": "類別", "campaign": "活動名", "period": "期間", "details": "優惠內容"}}
  ]
}}

請盡量收集各類別至少5-8筆優惠，只回覆JSON，不要其他文字。"""

    try:
        session = sessions_spawn(
            label=f"promo-{name.lower()}",
            mode="run",
            runtime="subagent",
            task=task
        )
        return session.get('childSessionKey')
    except Exception as e:
        print(f"   ❌ {name} 啟動失敗: {e}")
        return None

def parse_items_from_json(json_str):
    """解析 JSON 字串為 items 列表"""
    if not json_str:
        return []
    try:
        start = json_str.find('{')
        end = json_str.rfind('}') + 1
        if start >= 0 and end > start:
            data = json.loads(json_str[start:end])
            if 'items' in data:
                return data['items']
            elif isinstance(data, list):
                return data
    except:
        pass
    return []

def deduplicate_items(all_items):
    """去除重複項目"""
    seen = set()
    merged = []
    for item in all_items:
        key = (item.get('brand', ''), item.get('campaign', ''))
        if key not in seen:
            seen.add(key)
            merged.append(item)
    return merged

def generate_report(all_items):
    """生成 Word 報告"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    today = datetime.datetime.now()
    doc = Document()

    # 設定頁面
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    def set_cell_color(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def make_header_cell(cell, text):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_color(cell, '1F497D')

    def add_data_cell(cell, text):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_table(doc, title, data):
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        if not data:
            p = doc.add_paragraph()
            p.add_run('（尚無資料）').italic = True
            return

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['品牌/通路', '類別', '活動名稱', '活動期間', '優惠內容']
        hdr_cells = table.rows[0].cells
        for i, hdr in enumerate(headers):
            make_header_cell(hdr_cells[i], hdr)

        for item in data:
            row_cells = table.add_row().cells
            add_data_cell(row_cells[0], item.get('brand', ''))
            add_data_cell(row_cells[1], item.get('category', ''))
            add_data_cell(row_cells[2], item.get('campaign', ''))
            add_data_cell(row_cells[3], item.get('period', ''))
            add_data_cell(row_cells[4], item.get('details', ''))

        doc.add_paragraph()

    # 標題
    title = doc.add_heading('🛒 2026年5月母親節檔期 台灣促銷活動總整理', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('（小咪、小歐、千問、撈仔、拉瑪 五助理協力，Tavily搜尋，小安彙整）')
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f'整理日期：{today.strftime("%Y年%m月%d日")}')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # 統計
    stat_p = doc.add_paragraph()
    stat_p.add_run(f'📊 總計：{len(all_items)} 筆促銷資訊（已去除重複）').bold = True
    stat_p.runs[0].font.size = Pt(11)
    stat_p.runs[0].font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    doc.add_paragraph()

    # 分類顯示
    categories = [
        ('餐飲', '🍔 餐飲促銷'),
        ('超市', '🏪 超市/超商促銷'),
        ('銀行', '💳 銀行/支付優惠'),
        ('支付', '💳 支付優惠'),
        ('百貨', '🛍️ 百貨/電商優惠'),
        ('電商', '🛍️ 電商優惠'),
        ('電信', '📱 電信優惠'),
    ]

    for cat_key, cat_title in categories:
        items = [i for i in all_items if cat_key in i.get('category', '')]
        if items:
            add_table(doc, f'{cat_title}（{len(items)}筆）', items)

    # 備註
    p = doc.add_paragraph()
    p.add_run('📝 備註：').bold = True
    p = doc.add_paragraph()
    p.add_run('以上資料為五助理、Tavily搜尋協力蒐集，實際促銷內容及期間可能有所變動，建議消費前至各官方平台確認最新資訊。').italic = True
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 頁尾
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'小安助理 · 五助理協力 · {today.strftime("%Y年%m月%d日")}')
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # 儲存
    os.makedirs(REPORTS_DIR, exist_ok=True)
    output_path = f'{REPORTS_DIR}/促銷活動總整理_五助理版_{today.strftime("%Y%m%d_%H%M")}.docx'
    doc.save(output_path)
    print(f'✅ 已儲存: {output_path}')
    return output_path

def send_to_telegram(doc_path):
    """發送到 Telegram"""
    today = datetime.datetime.now()
    caption = f"🛒 促銷活動總整理_五助理版_{today.strftime('%Y年%m月%d日')}\n\n小咪、小歐、千問、撈仔、拉瑪協力搜尋，小安彙整去重後發送"

    result = subprocess.run(['curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
        '-F', f'chat_id={CHAT_ID}',
        '-F', f'document=@{doc_path}',
        '-F', f'caption={caption}'],
        capture_output=True, timeout=30)

    if result.returncode == 0:
        print('✅ 已發送到 Telegram')
    else:
        print(f'❌ 發送失敗')

def main():
    print("📋 開始生成促銷活動總整理報告（五助理協力版）...")
    print("=" * 50)
    print("🔄 五助理同時使用 Tavily 搜尋...\n")

    all_items = []

    # 小安靜態資料
    for item in XIAOAN_ITEMS:
        all_items.append(item)
    print(f"🐰 小安：已備好 {len(XIAOAN_ITEMS)} 筆母親節優惠")

    # 收集各助理搜尋結果
    # 這裡需要手動執行 Tavily 搜尋並收集結果
    # 由主腳本直接呼叫 Tavily API

    print("\n📝 小安：等待各助理搜尋結果...")
    print("=" * 50)

    # 由外部 main session 呼叫 Tavily 搜尋並彙整
    # 此脚本僅用於報告生成

    return all_items

if __name__ == '__main__':
    # 此脚本主要用於生成報告
    # 實際搜尋由 main agent 透過 sessions_spawn 執行
    print("此脚本公司使用 sessions_spawn 執行搜尋，完成後呼叫 generate_report()")
    print("請由主 agent 統籌執行")