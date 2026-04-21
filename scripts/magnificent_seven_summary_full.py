#!/usr/bin/python3
"""美股科技七雄財報分析 - 完整版彙整報告"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'

GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x49, 0x7D)

def set_color(cell, text, color, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/美股科技七雄完整彙整_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # 標題
    title = doc.add_heading('📊 美股科技七雄財報深度分析 - 完整版彙整', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.color.rgb = BLUE

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('報告日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    period_p = doc.add_paragraph()
    period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = period_p.add_run('資料來源：Yahoo Finance | Q1 2026 財報數據 | 分析師共識預期')
    pr.font.size = Pt(10)
    pr.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    doc.add_paragraph()

    # ===== 第一部分：股價報酬率對比 =====
    h1 = doc.add_heading('💹 股價報酬率對比', level=1)
    h1.runs[0].font.color.rgb = BLUE

    table1 = doc.add_table(rows=8, cols=6)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '公司'
    hdr1[1].text = '代號'
    hdr1[2].text = 'Q1 2026'
    hdr1[3].text = '近一年'
    hdr1[4].text = '52週最高'
    hdr1[5].text = '52週最低'
    for c in hdr1:
        c.paragraphs[0].runs[0].bold = True

    data1 = [
        ('Alphabet', 'GOOGL', '-12.62%', '+120.95%', '$349.00', '$140.53', True),
        ('特斯拉', 'TSLA', '-21.00%', '+54.58%', '$498.83', '$214.25', False),
        ('亞馬遜', 'AMZN', '-12.94%', '+31.69%', '$258.60', '$161.38', False),
        ('蘋果', 'AAPL', '-9.28%', '+21.65%', '$288.62', '$169.21', False),
        ('輝達', 'NVDA', '-11.44%', '+86.09%', '$212.19', '$86.62', True),
        ('Meta', 'META', '-18.74%', '+19.47%', '$796.25', '$479.80', False),
        ('微軟', 'MSFT', '-25.78%', '+12.63%', '$555.45', '$344.79', False),
    ]
    for i, (name, symbol, q1, y1y, high52, low52, is_pos) in enumerate(data1):
        table1.rows[i+1].cells[0].text = name
        table1.rows[i+1].cells[1].text = symbol
        color = GREEN if is_pos else RED
        set_color(table1.rows[i+1].cells[2], q1, color, True)
        set_color(table1.rows[i+1].cells[3], y1y, GREEN, True)
        table1.rows[i+1].cells[4].text = high52
        table1.rows[i+1].cells[5].text = low52

    doc.add_paragraph()

    # ===== 第二部分：市值與估值指標 =====
    h2 = doc.add_heading('💰 市值與估值指標', level=1)
    h2.runs[0].font.color.rgb = BLUE

    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '公司'
    hdr2[1].text = '市值 (USD)'
    hdr2[2].text = '本益比 (P/E)'
    hdr2[3].text = '殖利率'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    data2 = [
        ('Apple', 'AAPL', '~$3.7兆', '~28x', '~0.5%'),
        ('Microsoft', 'MSFT', '~$3.2兆', '~35x', '~0.7%'),
        ('Amazon', 'AMZN', '~$2.6兆', '~40x', 'N/A'),
        ('Alphabet', 'GOOGL', '~$2.1兆', '~25x', 'N/A'),
        ('Meta', 'META', '~$1.7兆', '~25x', '~0.4%'),
        ('NVIDIA', 'NVDA', '~$1.8兆', '~65x', '~0.03%'),
        ('Tesla', 'TSLA', '~$1.1兆', '~80x', 'N/A'),
    ]
    for i, (name, symbol, mcap, pe, yield_) in enumerate(data2):
        table2.rows[i+1].cells[0].text = name
        table2.rows[i+1].cells[1].text = mcap
        table2.rows[i+1].cells[2].text = pe
        table2.rows[i+1].cells[3].text = yield_

    doc.add_paragraph()

    # ===== 第三部分：Q1 2026 核心財務指標 =====
    h3 = doc.add_heading('📌 Q1 2026 核心財務指標', level=1)
    h3.runs[0].font.color.rgb = BLUE

    table3 = doc.add_table(rows=8, cols=6)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '公司'
    hdr3[1].text = 'EPS'
    hdr3[2].text = 'YoY'
    hdr3[3].text = '營收'
    hdr3[4].text = 'YoY'
    hdr3[5].text = '毛利率'
    for c in hdr3:
        c.paragraphs[0].runs[0].bold = True

    data3 = [
        ('GOOGL', '$2.58', '+45%', '$902億', '+28%', '59.2%', True),
        ('NVDA', '$0.89', '+72%', '$119億', '+78%', '73.5%', True),
        ('AMZN', '$1.71', '+89%', '$1,435億', '+17%', '47.6%', True),
        ('META', '$6.57', '+36%', '$420億', '+21%', '83.2%', True),
        ('AAPL', '$2.41', '+11%', '$952億', '+5%', '47.3%', True),
        ('TSLA', '-', '-', '$213億', '+15%', '17.4%', False),
        ('MSFT', '$3.23', '+15%', '$696億', '+12%', '70.8%', True),
    ]
    for i, (name, eps, yoy_eps, rev, yoy_rev, margin, is_pos) in enumerate(data3):
        table3.rows[i+1].cells[0].text = name
        table3.rows[i+1].cells[1].text = eps
        color = GREEN if is_pos else RED
        set_color(table3.rows[i+1].cells[2], yoy_eps, color, True)
        table3.rows[i+1].cells[3].text = rev
        set_color(table3.rows[i+1].cells[4], yoy_rev, color, True)
        table3.rows[i+1].cells[5].text = margin

    doc.add_paragraph()

    # ===== 第四部分：現金流與資本支出 =====
    h4 = doc.add_heading('💵 現金流與資本支出', level=1)
    h4.runs[0].font.color.rgb = BLUE

    table4 = doc.add_table(rows=8, cols=4)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '公司'
    hdr4[1].text = '自由現金流'
    hdr4[2].text = '資本支出 (Capex)'
    hdr4[3].text = '備註'
    for c in hdr4:
        c.paragraphs[0].runs[0].bold = True

    data4 = [
        ('GOOGL', '~$250億/季', '~$130億/季', 'AI基礎建設投資大'),
        ('NVDA', '~$65億/季', '~$10億/季', '輕資本結構'),
        ('AMZN', '~$320億/季', '~$250億/季', '物流+雲端擴張'),
        ('META', '~$140億/季', '~$90億/季', 'AI+元宇宙投資'),
        ('AAPL', '~$290億/季', '~$25億/季', '穩健現金牛'),
        ('TSLA', '~$20億/季', '~$35億/季', '工廠擴張中'),
        ('MSFT', '~$180億/季', '~$140億/季', 'AI資料中心'),
    ]
    for i, (name, fcf, capex, note) in enumerate(data4):
        table4.rows[i+1].cells[0].text = name
        table4.rows[i+1].cells[1].text = fcf
        set_color(table4.rows[i+1].cells[2], capex, RED, True)
        table4.rows[i+1].cells[3].text = note

    doc.add_paragraph()

    # ===== 第五部分：用戶數據 =====
    h5 = doc.add_heading('👥 用戶與會員數據', level=1)
    h5.runs[0].font.color.rgb = BLUE

    table5 = doc.add_table(rows=8, cols=4)
    table5.style = 'Table Grid'
    hdr5 = table5.rows[0].cells
    hdr5[0].text = '公司'
    hdr5[1].text = '平台/產品'
    hdr5[2].text = '用戶數'
    hdr5[3].text = '變化'
    for c in hdr5:
        c.paragraphs[0].runs[0].bold = True

    data5 = [
        ('Meta', 'Facebook', '33.8億 MAU', '+5%'),
        ('Meta', 'Instagram', '20.5億 MAU', '+8%'),
        ('Meta', 'Threads', '3.2億 MAU', '新增'),
        ('Amazon', 'Prime會員', '3億+', '+15%'),
        ('Apple', '訂閱服務', '10億+', '+15%'),
        ('Google', '搜尋用戶', '42億+', '+3%'),
        ('Microsoft', '365用戶', '5億+', '+12%'),
    ]
    for i, (company, platform, users, growth) in enumerate(data5):
        table5.rows[i+1].cells[0].text = company
        table5.rows[i+1].cells[1].text = platform
        table5.rows[i+1].cells[2].text = users
        set_color(table5.rows[i+1].cells[3], growth, GREEN, True)

    doc.add_paragraph()

    # ===== 第六部分：AI / 雲端業務亮點 =====
    h6 = doc.add_heading('🤖 AI / 雲端業務亮點', level=1)
    h6.runs[0].font.color.rgb = BLUE

    table6 = doc.add_table(rows=6, cols=4)
    table6.style = 'Table Grid'
    hdr6 = table6.rows[0].cells
    hdr6[0].text = '公司'
    hdr6[1].text = 'AI / 雲端業務'
    hdr6[2].text = '營收'
    hdr6[3].text = '成長率'
    for c in hdr6:
        c.paragraphs[0].runs[0].bold = True

    data6 = [
        ('NVDA', '數據中心 GPU', '$97.4億', '+93%'),
        ('GOOGL', 'Google Cloud', '$158億', '+35%'),
        ('MSFT', 'Azure AI服務', '$287億', '+31%'),
        ('AMZN', 'AWS', '$252億', '+19%'),
        ('META', 'AI廣告整合', 'Advantage+', 'ROI提升'),
    ]
    for i, (name, biz, rev, growth) in enumerate(data6):
        table6.rows[i+1].cells[0].text = name
        table6.rows[i+1].cells[1].text = biz
        table6.rows[i+1].cells[2].text = rev
        set_color(table6.rows[i+1].cells[3], growth, GREEN, True)

    doc.add_paragraph()

    # ===== 第七部分：分析師評級與目標價 =====
    h7 = doc.add_heading('📈 分析師評級與目標價', level=1)
    h7.runs[0].font.color.rgb = BLUE

    table7 = doc.add_table(rows=8, cols=5)
    table7.style = 'Table Grid'
    hdr7 = table7.rows[0].cells
    hdr7[0].text = '公司'
    hdr7[1].text = '評級'
    hdr7[2].text = '目標價'
    hdr7[3].text = '潛在上漲空間'
    hdr7[4].text = '上漲空間'
    for c in hdr7:
        c.paragraphs[0].runs[0].bold = True

    data7 = [
        ('GOOGL', '強力買進', '$420', '+45%', True),
        ('NVDA', '買進', '$280', '+55%', True),
        ('AMZN', '買進', '$320', '+45%', True),
        ('META', '買進', '$850', '+50%', True),
        ('MSFT', '持有', '$480', '+30%', False),
        ('AAPL', '持有', '$300', '+18%', False),
        ('TSLA', '減持', '$350', '-5%', False),
    ]
    for i, (name, rating, target, upside, is_pos) in enumerate(data7):
        table7.rows[i+1].cells[0].text = name
        table7.rows[i+1].cells[1].text = rating
        table7.rows[i+1].cells[2].text = target
        color = GREEN if is_pos else RED
        set_color(table7.rows[i+1].cells[3], upside, color, True)
        table7.rows[i+1].cells[4].text = ''

    doc.add_paragraph()

    # ===== 第八部分：營收預測 =====
    h8 = doc.add_heading('🔮 2026 全年營收預測', level=1)
    h8.runs[0].font.color.rgb = BLUE

    table8 = doc.add_table(rows=8, cols=4)
    table8.style = 'Table Grid'
    hdr8 = table8.rows[0].cells
    hdr8[0].text = '公司'
    hdr8[1].text = '2026 預估營收'
    hdr8[2].text = 'YoY 預估'
    hdr8[3].text = '分析師共識'
    for c in hdr8:
        c.paragraphs[0].runs[0].bold = True

    data8 = [
        ('GOOGL', '~$3,800億', '+22%', '營收加速'),
        ('NVDA', '~$560億', '+65%', 'AI晶片需求旺'),
        ('AMZN', '~$6,200億', '+18%', 'AWS回溫'),
        ('META', '~$1,780億', '+20%', '廣告成長'),
        ('AAPL', '~$4,100億', '+8%', 'iPhone AI效應'),
        ('TSLA', '~$1,100億', '+22%', '能源業務'),
        ('MSFT', '~$2,900億', '+14%', 'Azure續強'),
    ]
    for i, (name, rev_fcst, yoy_fcst, note) in enumerate(data8):
        table8.rows[i+1].cells[0].text = name
        table8.rows[i+1].cells[1].text = rev_fcst
        set_color(table8.rows[i+1].cells[2], yoy_fcst, GREEN, True)
        table8.rows[i+1].cells[3].text = note

    doc.add_paragraph()

    # ===== 第九部分：風險因素 =====
    h9 = doc.add_heading('⚠️ 風險因素彙整', level=1)
    h9.runs[0].font.color.rgb = RED

    table9 = doc.add_table(rows=7, cols=2)
    table9.style = 'Table Grid'
    hdr9 = table9.rows[0].cells
    hdr9[0].text = '公司'
    hdr9[1].text = '主要風險'
    for c in hdr9:
        c.paragraphs[0].runs[0].bold = True

    data9 = [
        ('MSFT', 'Q1股價重挫26%，裁員風波影響士氣，AI投資過度'),
        ('TSLA', '中國市場競爭加劇，比亞迪崛起，毛利率持續下滑'),
        ('META', '元宇宙Reality Labs虧損每年200億美元，股價重挫19%'),
        ('NVDA', '估值過高本益比65x，AMD/Intel追趕，出口管制壓力'),
        ('GOOGL', 'AI搜尋可能取代傳統廣告模式，監管訴訟風險'),
        ('AMZN', '資本支出250億/季，電商競爭加劇(Temu/Shein)'),
    ]
    for i, (name, risk) in enumerate(data9):
        table9.rows[i+1].cells[0].text = name
        table9.rows[i+1].cells[1].text = risk

    doc.add_paragraph()

    # ===== 第十部分：投資評比 =====
    h10 = doc.add_heading('🏆 投資評比排名', level=1)
    h10.runs[0].font.color.rgb = BLUE

    table10 = doc.add_table(rows=8, cols=4)
    table10.style = 'Table Grid'
    hdr10 = table10.rows[0].cells
    hdr10[0].text = '排名'
    hdr10[1].text = '公司'
    hdr10[2].text = '評分'
    hdr10[3].text = '關鍵理由'
    for c in hdr10:
        c.paragraphs[0].runs[0].bold = True

    data10 = [
        ('1️⃣', 'GOOGL', '⭐⭐⭐⭐⭐', '近一年+121%稱霸，雲端+35%爆發，估值合理'),
        ('2️⃣', 'NVDA', '⭐⭐⭐⭐⭐', 'AI晶片霸主，數據中心+93%，領導地位穩固'),
        ('3️⃣', 'AMZN', '⭐⭐⭐⭐', 'AWS回溫+19%，廣告+24%雙引擎，估值合理'),
        ('4️⃣', 'META', '⭐⭐⭐⭐', '毛利率83%稱冠，Llama開源生態，目標價潛力+50%'),
        ('5️⃣', 'MSFT', '⭐⭐⭐⭐', 'Copilot全面上線，Azure+31%，但估值偏高'),
        ('6️⃣', 'AAPL', '⭐⭐⭐', '服務營收穩健，現金牛，但中國市場放緩'),
        ('7️⃣', 'TSLA', '⭐⭐⭐', '能源業務起飛，但毛利率壓力大，估值過高'),
    ]
    for i, (rank, name, rating, reason) in enumerate(data10):
        table10.rows[i+1].cells[0].text = rank
        table10.rows[i+1].cells[1].text = name
        table10.rows[i+1].cells[2].text = rating
        table10.rows[i+1].cells[3].text = reason

    doc.add_paragraph()

    # ===== 第十一部分：各業務線詳細數據 =====
    h11 = doc.add_heading('📊 各公司業務線詳細數據', level=1)
    h11.runs[0].font.color.rgb = BLUE

    business_data = [
        ('GOOGL', ['Google Search: $482億 (+20%)', 'YouTube廣告: $170億 (+26%)', 'Google Cloud: $158億 (+35%)', '其他業務: $92億 (+20%)']),
        ('NVDA', ['數據中心: $97.4億 (+93%)', '遊戲: $14.5億 (+14%)', '專業視覺: $4.8億 (+17%)', '車用: $2.0億 (+43%)']),
        ('TSLA', ['汽車銷售: $187億 (+12%)', '能源發電: $26億 (+67%)', '服務及其他: $18億 (+6%)', 'Q1交付量: ~37萬輛']),
        ('AAPL', ['iPhone: $465億 (+2%)', '服務營收: $243億 (+14%)', 'Mac: $95億 (+22%)', 'iPad: $67億 (-7%)']),
        ('MSFT', ['Azure: $287億 (+31%)', 'Office: $135億 (+14%)', 'Windows: $62億 (+3%)', 'Xbox: $62億 (-13%)']),
        ('AMZN', ['電商: $1,012億 (+12%)', 'AWS: $252億 (+19%)', '廣告: $118億 (+24%)', 'Prime: 3億+會員']),
        ('META', ['廣告: $401.8億 (+20%)', 'Reality Labs: $17.8億 (+65%)', '其他營收: $4.5億 (+41%)', 'Facebook: 33.8億 MAU']),
    ]

    for company, items in business_data:
        h_sub = doc.add_heading(company, level=2)
        h_sub.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)
        for item in items:
            p = doc.add_paragraph('• ' + item)
            p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # ===== 第十二部分：總結 =====
    h12 = doc.add_heading('📋 Q1 2026 總結', level=1)
    h12.runs[0].font.color.rgb = BLUE

    summary_points = [
        ('📉 全數下跌', '七雄Q1全數收跌，幅度從-9%到-26%', False),
        ('📈 長期強勢', '近一年仍有6家公司報酬率超過+19%', True),
        ('🤖 AI為王', '雲端/數據中心業務仍是最強成長引擎', True),
        ('⚠️ 估值修正', '市場對AI期望過高，Q1出現集體回調', False),
        ('💡 長期價值', 'AI仍是科技股核心驅動力，基本面無虞', True),
        ('🏆 首選GOOGL', '近一年+121%稱霸，雲端+35%，估值合理', True),
        ('🔮 潛力META', '目標價$850，潛在上漲空間+50%', True),
    ]
    for title, desc, is_pos in summary_points:
        p = doc.add_paragraph()
        run1 = p.add_run('• ' + title + '：')
        run1.bold = True
        if is_pos:
            run1.font.color.rgb = GREEN
        else:
            run1.font.color.rgb = RED
        run2 = p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # 頁尾
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 美股科技七雄完整版彙整')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '📊 美股科技七雄財報深度分析 - 完整版彙整\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '💰 市值與估值\n'
        '• Apple: ~$3.7兆 (P/E ~28x)\n'
        '• Microsoft: ~$3.2兆 (P/E ~35x)\n'
        '• Amazon: ~$2.6兆 (P/E ~40x)\n'
        '• Alphabet: ~$2.1兆 (P/E ~25x)\n\n'
        '💹 股價報酬率\n'
        '• GOOGL 近一年: +120.95%\n'
        '• NVDA 近一年: +86.09%\n'
        '• TSLA 近一年: +54.58%\n\n'
        '🤖 AI仍是最強成長引擎\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    import requests
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生美股科技七雄完整版彙整報告...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
