#!/usr/bin/env python3
"""
🚀 創業想法研究報告自動執行
川寶投顧 × 五位助理聯合研究
執行時間：每日 03:00
"""

import subprocess
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 五位助理的創業建議
assistants_data = {
    "小安": {
        "title": "🐰 小安（AI智能助理）創業建議",
        "suggestions": """
1. **AI 客服懶人包服務**
   - 產業：企業服務 / AI 應用
   - 資金門檻：低（15-50萬）
   - 技術需求：現有 AI API 串接、基础程式能力
   - 說明：幫中小商家快速建立 AI 客服機器人，懶人包方案，月費制
   - 優勢：結合 OpenClaw AI 系統技術經驗

2. **手搖飲品牌顧問服務**
   - 產業：餐飲顧問
   - 資金門檻：低（10-30萬）
   - 技術需求：餐飲業know-how、設備評估
   - 說明：結合飲料業加盟研究經驗，提供選點、品牌設定、AI自動化流程顧問服務

3. **自動化報告生成器（SaaS）**
   - 產業：企業服務 / 工具軟體
   - 資金門檻：中（30-80萬）
   - 技術需求：Python、API整合能力
   - 說明：幫企業自動生成每週/月報告，訂閱制

4. **AI 創業課程平台**
   - 產業：教育科技
   - 資金門檻：低（5-20萬）
   - 技術需求：影片製作、簡報能力

5. **外包接案平台（利基市場）**
   - 產業：派遣服務 / 數位接案
   - 資金門檻：低（10-40萬）
"""
    },
    "小歐": {
        "title": "🌍 小歐（國際財經專家）創業建議",
        "suggestions": """
1. **數位內容創作平台** - 資金：低（10-50萬）| 技術：中等
2. **AI 應用服務（SaaS）** - 資金：中（100-300萬）| 技術：高
3. **健康餐盒外送品牌** - 資金：中（50-200萬）| 技術：低~中
4. **ESG 永續顧問公司** - 資金：低（30-100萬）| 技術：中
5. **工業物聯網（IIoT）解決方案** - 資金：高（500萬以上）| 技術：高
"""
    },
    "千問": {
        "title": "🧐 千問（研究分析師）創業建議",
        "suggestions": """
1. **微型電商（自有品牌選品）** - 資金：低（5-15萬）| 技術：低
2. **SaaS工具/訂閱制軟體** - 資金：中（30-100萬）| 技術：中高
3. **特色餐飲（健康/機能飲品）** - 資金：中高（50-200萬）| 技術：中
4. **AI應用服務（垂直領域Agent）** - 資金：中（20-80萬）| 技術：中高
5. **垂直知識付費（專業課程/顧問）** - 資金：低（5-20萬）| 技術：中
"""
    },
    "拉瑪": {
        "title": "🐂 拉瑪（深度研究顧問）創業建議",
        "suggestions": """
1. **健身科技/健康管理平台** - 資金：低至中等 | 技術：中等
2. **永續電商/綠色消費平台** - 資金：中等 | 技術：中
3. **線上教育/技能培訓平台** - 資金：低 | 技術：中
4. **AI顧問服務/企業流程自動化** - 資金：中等至高 | 技術：高
5. **餐飲科技/雲端廚房品牌** - 資金：中等 | 技術：中
"""
    },
    "撈仔": {
        "title": "🐔 撈仔（萬能小幫手）創業建議",
        "suggestions": """
1. **微型植物工廠（科技農業）** - 資金：中（50-200萬）| 技術：中
2. **客製化禮品電商平台** - 資金：低（10-50萬）| 技術：低
3. **AI客服代營運服務** - 資金：低（20-80萬）| 技術：中
4. **線上健身教練平台** - 資金：低（10-50萬）| 技術：中
5. **工業物聯網預測維護系統** - 資金：高（200萬以上）| 技術：高
"""
    }
}

def generate_report():
    doc = Document()
    title = doc.add_heading('🚀 創業想法研究報告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph(f'報告日期：{datetime.datetime.now().strftime("%Y年%m月%d日")}')
    doc.add_paragraph('研究團隊：小安、小歐、千問、拉瑪、撈仔')
    doc.add_paragraph()
    
    # 各助理建議
    for name, data in assistants_data.items():
        doc.add_heading(data['title'], level=1)
        doc.add_paragraph(data['suggestions'])
    
    # 共識統計
    doc.add_heading('📊 五位助理共識統計', level=1)
    doc.add_paragraph("""
| 創業方向 | 小安 | 小歐 | 千問 | 拉瑪 | 撈仔 | 總計 |
|---------|------|------|------|------|------|------|
| AI應用/客服服務 | ✅ | ✅ | ✅ | ✅ | ✅ | 5票 |
| 電商/選品 | — | — | ✅ | ✅ | ✅ | 3票 |
| 餐飲/食品 | — | ✅ | ✅ | ✅ | — | 3票 |
| 知識付費/教育 | ✅ | — | ✅ | ✅ | — | 3票 |
| 健康/健身科技 | — | — | — | ✅ | ✅ | 2票 |
| IoT/工業 | — | ✅ | — | — | ✅ | 2票 |
| ESG/永續 | — | ✅ | — | ✅ | — | 2票 |
| 自動化工具/SaaS | ✅ | — | ✅ | — | — | 2票 |
""")
    
    # 小安結論
    doc.add_heading('🎯 小安綜合結論與建議', level=1)
    doc.add_paragraph("""
【最高共識方向：AI應用/客服服務】5票（全部推薦）

🥇 首選：AI 客服懶人包服務（資金15-50萬）
🥈 次選：手搖飲品牌顧問服務（資金10-30萬）
🥉 第三：自動化報告生成器 SaaS（資金30-80萬）

【創業建議】
- 先做 MVP 驗證需求
- 不要一開始就砸大錢
- 選擇自己擅長且有興趣的領域
""")
    
    doc.add_paragraph()
    doc.add_paragraph(f'報告生成時間：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('川寶投顧 小安匯總')
    
    filename = f'/root/.openclaw/reports/daily/startup_ideas_{datetime.datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    return filename

def send_telegram(docx_path):
    import requests
    url = f"https://api.telegram.org/bot8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw/sendDocument"
    with open(docx_path, 'rb') as f:
        files = {'document': f}
        data = {'chat_id': '8779713208', 'caption': '🚀 創業想法研究報告（每日自動生成）| 川寶投顧×五位助理'}
        r = requests.post(url, files=files, data=data)
    return r.json().get('ok', False)

if __name__ == '__main__':
    print(f"[{datetime.datetime.now()}] 創業研究報告開始生成...")
    docx_path = generate_report()
    print(f"報告已生成：{docx_path}")
    if send_telegram(docx_path):
        print("已發送到 Telegram")
    else:
        print("Telegram 發送失敗")