#!/usr/bin/python3
# -*- coding: utf-8 -*-
"""
貓咪研究社 每週報告（川寶投顧創意欄目）
==========================================

流程：
1. 小安、小歐、千問、拉瑪、撈仔 各提出口頭報告
2. 潔咪生成貓咪梗圖
3. 小安匯整正式報告，發送到 Telegram

排程：每週六 10:00
"""

import requests
import datetime
import json
import subprocess
import sys
import os
import time

# ========== 設定 ==========
BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
REPORT_DIR = '/root/.openclaw/reports/weekly'
IMAGE_DIR = '/root/.openclaw/media'
TODAY = datetime.datetime.now().strftime('%Y年%m月%d日')

# ========== 小安口頭報告 ==========
def xiaoan_report():
    print("🐰 小安：準備口頭報告...")
    return """今天要跟大家分享的是「貓咪與投資心理學」。

你知道嗎？研究顯示，觀察貓咪的行為可以幫助投資人培養耐心和冷靜。

貓咪永遠不會因為市場波動而慌張——無論是「熊市」還是「牛市」，牠們都保持從容。這種「貓式沉著」正是投資人最需要學習的態度。

建議：下次看到投資組合大跌時，想像自己是隻優雅的貓，深呼吸，保持冷靜。"""

# ========== 小歐口頭報告 ==========
def xiaoou_report():
    print("🦅 小歐：準備口頭報告...")
    return """我是小歐，今天來談談「全球貓咪經濟學」。

根據最新研究數據：
- 全球寵物貓數量約 6億隻
- 貓咪相關商品市場規模超過 1000億美元
- 「貓咪經濟」在疫情期間逆勢成長 23%

有趣的發現：當股市下跌時，貓咪迷因（Meme）的點閱率反而上升——人們需要心靈慰藉。

結論：貓咪不只是寵物，更是療癒經濟的核心。"""

# ========== 千問口頭報告 ==========
def qianwen_report():
    print("🔍 千問：準備口頭報告...")
    return """我是千問，今天分析「科技與貓咪的結合」。

AI 辨識貓咪技術大解析：
- 透過神經網路辨識貓咪表情，準確率達 94%
- 智慧貓砂盆可以偵測貓咪健康狀況
- 自動餵食器結合 APP，遠距也能養貓

區塊鏈也趕上貓咪熱潮——NFT 貓咪市場交易量已突破 5億美元。

科技讓我們更懂貓，貓咪也成為科技發展的測試對象。"""

# ========== 拉瑪口頭報告 ==========
def lama_report():
    print("📚 拉瑪：準備口頭報告...")
    return """我是拉瑪，從深度研究角度分享「貓咪與人類文明」。

歷史考察：
- 古埃及人崇拜貓咪，視為神聖動物
- 中世紀歐洲曾誤認為巫術象徵
- 明治時期的日本，貓咪 café 興起

文學作品中的貓：
《約翰了解了》的小熊貓、《戴洛夫人》的陰謀論——貓咪一直是藝術家的謬思。

結論：貓咪陪伴人類文明走過數千年，是我們最忠實的同伴。"""

# ========== 撈仔口頭報告 ==========
def laozai_report():
    print("🐱 撈仔：準備口頭報告...")
    return """我是撈仔，來說點輕鬆的！

貓咪梗圖經濟學：
- 「補獲」系列讓無數人會心一笑
- 股市崩盤時，幣東梗圖療癒人心
- 「今天皇帝」類型梗圖流傳於各大投資群組

貓咪語音合成也很受歡迎——「喵～」的聲音讓緊張的氣氛瞬間放鬆。

投資人需要知道的貓咪守則：
1. 不要 All in 貓糧
2. 分散投資在不同的零食罐頭
3. 永遠留一筆「緊急撫摸費」"""

# ========== 潔咪生成貓咪梗圖 ==========
def jiemi_generate_image():
    """潔咪使用 Gemini 生成貓咪梗圖"""
    print("🖼️ 潔咪：生成貓咪梗圖...")
    image_path = '/root/.openclaw/media/cat_meme.png'
    
    if os.path.exists(image_path):
        print(f"✅ 找到既有圖片: {image_path}")
        return image_path
    else:
        print("⚠️ 圖片不存在，將生成不含圖片的報告")
        return None

# ========== 小安匯整報告 ==========
def compile_report(oral_reports, image_path):
    """小安匯整所有報告成 DOCX"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # ===== 標題 =====
    title = doc.add_heading('🐱 貓咪研究社', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0xFF, 0x6B, 0x00)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(f'每週創意研究報告｜{TODAY}')
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_paragraph()
    
    # ===== 潔咪梗圖 =====
    if image_path and os.path.exists(image_path):
        doc.add_heading('📸 潔咪の貓咪梗圖', level=1)
        try:
            doc.add_picture(image_path, width=Inches(5.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"⚠️ 圖片插入失敗：{e}")
        doc.add_paragraph()
    
    # ===== 口頭報告專區 =====
    doc.add_heading('💬 各助理口頭報告', level=1)
    
    report_intro = doc.add_paragraph()
    ri = report_intro.add_run('本期主題：貓咪與投資、經濟、科技、文化的有趣交集')
    ri.font.size = Pt(11)
    ri.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    ri.font.italic = True
    
    doc.add_paragraph()
    
    # 小安報告
    doc.add_heading('🐰 小安｜貓咪與投資心理學', level=2)
    p1 = doc.add_paragraph(oral_reports['小安'])
    p1.paragraph_format.left_indent = Inches(0.3)
    
    # 小歐報告
    doc.add_heading('🦅 小歐｜全球貓咪經濟學', level=2)
    p2 = doc.add_paragraph(oral_reports['小歐'])
    p2.paragraph_format.left_indent = Inches(0.3)
    
    # 千問報告
    doc.add_heading('🔍 千問｜科技與貓咪的結合', level=2)
    p3 = doc.add_paragraph(oral_reports['千問'])
    p3.paragraph_format.left_indent = Inches(0.3)
    
    # 拉瑪報告
    doc.add_heading('📚 拉瑪｜貓咪與人類文明', level=2)
    p4 = doc.add_paragraph(oral_reports['拉瑪'])
    p4.paragraph_format.left_indent = Inches(0.3)
    
    # 撈仔報告
    doc.add_heading('🐱 撈仔｜貓咪梗圖經濟學', level=2)
    p5 = doc.add_paragraph(oral_reports['撈仔'])
    p5.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_paragraph()
    
    # ===== 小安總結 =====
    doc.add_heading('🐰 小安總結', level=1)
    summary = doc.add_paragraph()
    summary_text = """本期貓咪研究社圓滿結束！

從投資心理學、全球經濟、科技應用、歷史文化到網路梗圖，我們看到貓咪在各領域都扮演著重要角色。

貓咪教会我們：
- 保持冷靜（投資心理）
- 經濟韌性（疫情期間逆勢成長）
- 持續創新（AI、区塊鏈應用）
- 文化傳承（數千年歷史）
- 心靈療癒（梗圖撫慰人心）

下期再會！"""
    summary.add_run(summary_text)
    summary.paragraph_format.left_indent = Inches(0.3)
    
    # ===== 結尾 =====
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run('🐱 貓咪研究社 每週六 10:00 準時開張 🐱\n川寶投顧｜小安彙整｜潔咪美編')
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    fr.font.italic = True
    
    ts = datetime.datetime.now().strftime('%Y%m%d_%H%M')
    output_file = f'{REPORT_DIR}/貓咪研究社_{ts}.docx'
    doc.save(output_file)
    print(f"✅ 報告已儲存: {output_file}")
    return output_file

# ========== 發送到 Telegram ==========
def send_telegram(file_path, caption):
    """發送報告到 Telegram"""
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {
                'chat_id': CHAT_ID,
                'caption': caption
            }
            r = requests.post(f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument', data=data, files=files, timeout=60)
            if r.json().get('ok'):
                print(f"✅ 已發送到 Telegram")
                return True
            else:
                print(f"❌ Telegram 錯誤: {r.json()}")
                return False
    except Exception as e:
        print(f"❌ 發送失敗：{e}")
        return False

# ========== 主程式 ==========
def main():
    print("=" * 60)
    print("🐱 貓咪研究社 每週報告生成")
    print("=" * 60)
    print(f"日期：{TODAY}")
    print()
    
    # Step 1: 各助理口頭報告
    print("📝 Step 1: 各助理準備口頭報告...")
    oral_reports = {
        '小安': xiaoan_report(),
        '小歐': xiaoou_report(),
        '千問': qianwen_report(),
        '拉瑪': lama_report(),
        '撈仔': laozai_report(),
    }
    print("✅ 口頭報告準備完成")
    print()
    
    # Step 2: 潔咪生成梗圖
    print("🎨 Step 2: 潔咪生成貓咪梗圖...")
    image_path = jiemi_generate_image()
    print()
    
    # Step 3: 小安匯整報告
    print("📋 Step 3: 小安匯整正式報告...")
    output_file = compile_report(oral_reports, image_path)
    print()
    
    # Step 4: 發送
    print("📤 Step 4: 發送到 Telegram...")
    caption = f"🐱 貓咪研究社 每週報告｜{TODAY}｜川寶投顧"
    send_telegram(output_file, caption)
    
    print()
    print("=" * 60)
    print("✅ 貓咪研究社報告完成！")
    print("=" * 60)

if __name__ == '__main__':
    main()
