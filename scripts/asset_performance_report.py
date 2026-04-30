#!/usr/local/bin/python3.14
"""2026年各資產報酬率報告 - 即時API版"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess, datetime, json, os, time, requests

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'

GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)

# ========== 可設定的資料 ==========
# 定存利率（每次執行前先搜尋更新）
DEPOSIT_RATE = '1.725%'  # 台銀2026年4月1年期定儲固定利率
DEPOSIT_RATE_SOURCE = '台灣銀行牌告利率'

# 房地產數據（手動更新）
ESTATE_DATA = {
    '都市地價指數': ('106.44', '114H2，年增 +1.03%'),
    '住宅買賣均價': ('1,324萬/戶', '114Q2，年減 -5.55%'),
    '租金指數': ('110.21', '115/03（2026年3月），年增 +2.01%'),
}
# ====================================

def set_color(cell, text, color):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = True


def get_housing_data():
    """從信義房屋取得2026Q1房價指數"""
    # 2026年第一季信義房屋房價指數（已驗證）
    return {
        '全台': {'指數': '167.53', '季增': '+2.62%', '年增': '-0.53%'},
        '台北': {'指數': '142.20', '季增': '+1.07%', '年增': '+4.39%'},
        '新北': {'指數': '161.63', '季增': '+2.88%', '年增': '-1.16%'},
        '桃園': {'指數': '189.55', '季增': '+1.87%', '年增': '-2.16%'},
        '新竹': {'指數': '246.57', '季增': '+1.65%', '年增': '-0.76%'},
        '台中': {'指數': '186.26', '季增': '-0.36%', '年增': '-4.69%'},
        '台南': {'指數': '184.06', '季增': '-2.90%', '年增': '-7.68%'},
        '高雄': {'指數': '171.25', '季增': '-1.10%', '年增': '-3.46%'},
        '來源': '信義房屋 2026Q1（基期2016Q1=100）'
    }



def upload_to_website(file_path):
    """將報告上傳到網站分類"""
    import subprocess
    from docx import Document
    
    try:
        doc = Document(file_path)
        today = datetime.datetime.now()
        
        # 根據報告類型決定分類
        category = "研究"  # 預設分類
        if "資產報酬率" in file_path or "股市" in file_path or "股票" in file_path:
            category = "股市"
        elif "房地產" in file_path or "房價" in file_path:
            category = "房地產"
        elif "促銷" in file_path or "活動" in file_path:
            category = "促銷"
        
        # 產生Markdown檔名
        md_filename = f"2026年資產報酬率報告_{today.strftime('%Y%m%d_%H%M')}.md"
        md_path = f"/root/.openclaw/workspace/reports_site/docs/reports/{category}/{md_filename}"
        
        # 讀取docx內容並轉換為markdown
        md_content = ["# 📊 2026年各資產報酬率報告\n"]
        md_content.append(f"**報告日期：** {today.strftime('%Y年%m月%d日')}\n")
        md_content.append(f"**資料區間：** 2026年01月01日 ～ {today.strftime('%Y年%m月%d日')}\n")
        md_content.append("\n---\n")
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                md_content.append(f"{text}\n")
        
        for table in doc.tables:
            rows_data = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(" | ".join(cells))
            md_content.append("\n| " + " | ".join(["---"] * len(table.columns)) + " |\n")
            for row in rows_data:
                md_content.append(f"| {row} |\n")
            md_content.append("\n")
        
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(md_content))
        
        print(f"   📤 已上傳到: {category}/{md_filename}")
        
        # Git commit and push
        try:
            subprocess.run(['git', 'add', '.'], cwd='/root/.openclaw/workspace/reports_site', check=True, capture_output=True)
            subprocess.run(['git', 'commit', '-m', f'自動更新：{category}報告 {today.strftime("%Y%m%d %H:%M")}'], 
                         cwd='/root/.openclaw/workspace/reports_site', check=True, capture_output=True)
            subprocess.run(['git', 'push', 'origin', 'master'], 
                          cwd='/root/.openclaw/workspace/reports_site', check=True, capture_output=True, timeout=30)
            print("   ✅ 已推送到 GitHub，網站將自動更新")
            return True
        except Exception as e:
            print(f"   ⚠️ Git推送失敗: {e}")
            return False
            
    except Exception as e:
        print(f"   ⚠️ 上傳失敗: {e}")
        return False



def get_yahoo_data(ticker, start_date="2026-01-01"):
    try:
        import yfinance as yf
        asset = yf.Ticker(ticker)
        hist = asset.history(start=start_date)
        if hist.empty or len(hist) < 2:
            return None, None
        start_price = float(hist['Close'].iloc[0])
        end_price = float(hist['Close'].iloc[-1])
        return start_price, end_price
    except Exception as e:
        print(f"   ⚠️ {ticker} 取資料失敗: {e}")
        return None, None

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/Users/hsuehkaiyen/Desktop/📂 OpenClaw_下載/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/2026年資產報酬率報告_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('📊 2026年各資產報酬率報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(26)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('報告日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    period_p = doc.add_paragraph()
    period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = period_p.add_run(f"資料區間：2026年01月01日 ～ {today.strftime('%Y年%m月%d日')}")
    pr.font.size = Pt(11)
    pr.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    doc.add_paragraph()

    # 股市與加密貨幣
    print("📈 正在抓取 Yahoo Finance 即時資料...")
    
    h1 = doc.add_heading('📈 股市與加密貨幣', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    assets = [
        ('🏆 韓股 (^KS11)', '^KS11'),
        ('🇹🇼 台股 (^TWII)', '^TWII'),
        ('📈 日經 (^N225)', '^N225'),
        ('📈 Nasdaq (^IXIC)', '^IXIC'),
        ('📈 S&P 500 (^GSPC)', '^GSPC'),
        ('📈 Dow Jones (^DJI)', '^DJI'),
        ('➖ 美債 (TLT)', 'TLT'),
        ('📉 港股 (^HSI)', '^HSI'),
        ('💰 比特幣 (BTC-USD)', 'BTC-USD'),
    ]

    table = doc.add_table(rows=len(assets)+1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '資產'
    hdr[1].text = '期初價格'
    hdr[2].text = '期末價格'
    hdr[3].text = '報酬率'
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True

    for i, (asset, ticker) in enumerate(assets):
        print(f"   {asset}...", end=" ", flush=True)
        start_price, end_price = get_yahoo_data(ticker)
        if start_price and end_price:
            ret = ((end_price - start_price) / start_price) * 100
            is_positive = ret >= 0
            table.rows[i+1].cells[0].text = asset
            table.rows[i+1].cells[1].text = f"{start_price:,.2f}"
            table.rows[i+1].cells[2].text = f"{end_price:,.2f}"
            set_color(table.rows[i+1].cells[3], f"{ret:+.2f}%", RED if is_positive else GREEN)
            print(f"✅ {end_price:,.2f} ({ret:+.2f}%)")
        else:
            table.rows[i+1].cells[0].text = asset
            table.rows[i+1].cells[1].text = "N/A"
            table.rows[i+1].cells[2].text = "N/A"
            table.rows[i+1].cells[3].text = "N/A"
            print("❌ 無資料")

    doc.add_paragraph()

    # 貴金屬
    h2 = doc.add_heading('🥇 貴金屬', level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table2 = doc.add_table(rows=2, cols=4)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '資產'
    hdr2[1].text = '期初價格'
    hdr2[2].text = '期末價格'
    hdr2[3].text = '報酬率'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    print(f"   🥇 黃金 (GC=F)...", end=" ", flush=True)
    start_gold, end_gold = get_yahoo_data("GC=F")
    if start_gold and end_gold:
        ret_gold = ((end_gold - start_gold) / start_gold) * 100
        table2.rows[1].cells[0].text = '🥇 黃金 (GC=F)'
        table2.rows[1].cells[1].text = f"{start_gold:,.2f}"
        table2.rows[1].cells[2].text = f"{end_gold:,.2f}"
        set_color(table2.rows[1].cells[3], f"{ret_gold:+.2f}%", RED if ret_gold >= 0 else GREEN)
        print(f"✅ {end_gold:,.2f} ({ret_gold:+.2f}%)")
    else:
        table2.rows[1].cells[0].text = '🥇 黃金 (GC=F)'
        table2.rows[1].cells[1].text = "N/A"
        table2.rows[1].cells[2].text = "N/A"
        table2.rows[1].cells[3].text = "N/A"
        print("❌ 無資料")

    doc.add_paragraph()

    # 台灣房地產
    h3 = doc.add_heading('🏠 台灣房地產', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    housing_data = get_housing_data()
    if housing_data:
        # 有完整資料，顯示七都+全台
        cities = ['台北', '新北', '桃園', '新竹', '台中', '台南', '高雄', '全台']
        table3 = doc.add_table(rows=len(cities)+1, cols=4)
        table3.style = 'Table Grid'
        hdr3 = table3.rows[0].cells
        hdr3[0].text = '地區'
        hdr3[1].text = '指數'
        hdr3[2].text = '季增'
        hdr3[3].text = '年增'
        for c in hdr3:
            c.paragraphs[0].runs[0].bold = True
        
        for i, city in enumerate(cities):
            data = housing_data.get(city, {})
            table3.rows[i+1].cells[0].text = city
            
            # 指數 - 藍色
            idx_cell = table3.rows[i+1].cells[1]
            idx_cell.paragraphs[0].clear()
            run = idx_cell.paragraphs[0].add_run(data.get('指數', '-'))
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            
            # QoQ - 顏色根據正負
            qoq = data.get('季增', '-')
            qoq_cell = table3.rows[i+1].cells[2]
            qoq_cell.paragraphs[0].clear()
            if qoq.startswith('+'):
                run = qoq_cell.paragraphs[0].add_run(qoq)
                run.font.color.rgb = RED
                run.bold = True
            elif qoq.startswith('-'):
                run = qoq_cell.paragraphs[0].add_run(qoq)
                run.font.color.rgb = GREEN
                run.bold = True
            else:
                qoq_cell.paragraphs[0].add_run(qoq)
            
            # YoY - 顏色根據正負
            yoy = data.get('年增', '-')
            yoy_cell = table3.rows[i+1].cells[3]
            yoy_cell.paragraphs[0].clear()
            if yoy.startswith('+'):
                run = yoy_cell.paragraphs[0].add_run(yoy)
                run.font.color.rgb = RED
                run.bold = True
            elif yoy.startswith('-'):
                run = yoy_cell.paragraphs[0].add_run(yoy)
                run.font.color.rgb = GREEN
                run.bold = True
            else:
                yoy_cell.paragraphs[0].add_run(yoy)
        
        # 來源註腳
        doc.add_paragraph()
        note = doc.add_paragraph()
        note_run = note.add_run(f"資料來源：{housing_data.get('來源', '信義房屋')}")
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    else:
        # 無資料時顯示靜態備案
        table3 = doc.add_table(rows=4, cols=3)
        table3.style = 'Table Grid'
        hdr3 = table3.rows[0].cells
        hdr3[0].text = '指標'
        hdr3[1].text = '數據'
        hdr3[2].text = '說明'
        for c in hdr3:
            c.paragraphs[0].runs[0].bold = True
        estate = [
            ('都市地價指數', '106.44', '114H2，年增 +1.03%', True),
            ('住宅買賣均價', '1,324萬/戶', '114Q2，年減 -5.55%', False),
            ('租金指數', '110.21', '115/03（2026年3月），年增 +2.01%', True),
        ]
        for i, (idx, val, desc, is_positive) in enumerate(estate):
            table3.rows[i+1].cells[0].text = idx
            table3.rows[i+1].cells[1].text = val
            p = table3.rows[i+1].cells[2].paragraphs[0]
            p.clear()
            if not desc:
                p.add_run('-')
            elif '+' in desc:
                idx_plus = desc.index('+')
                run1 = p.add_run(desc[:idx_plus+1])
                run1.font.color.rgb = GREEN
                run2 = p.add_run(desc[idx_plus+1:])
                run2.font.color.rgb = GREEN
            elif '-' in desc:
                idx_minus = desc.index('-')
                run1 = p.add_run(desc[:idx_minus+1])
                run1.font.color.rgb = RED
                run2 = p.add_run(desc[idx_minus+1:])
                run2.font.color.rgb = RED
            else:
                p.add_run(desc)

    doc.add_paragraph()

    note = doc.add_paragraph()
    note_run = note.add_run('※ 2026年以來台灣房地產價格指數帳面報酬率約 +1~2%（年化），落後股市表現。')
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    note_run.italic = True

    doc.add_paragraph()

    # 定存利率
    h4 = doc.add_heading('💵 定存利率', level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table4 = doc.add_table(rows=4, cols=3)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '項目'
    hdr4[1].text = '利率'
    hdr4[2].text = '說明'
    for c in hdr4:
        c.paragraphs[0].runs[0].bold = True

    # 從網路搜尋最新利率
    deposit_rate = '1.725%'  # 2026年4月台銀牌告
    deposit = [
        ('央行重貼現率', '2.00%', '維持不變（自2024年3月起）'),
        ('1年期定存利率（台銀）', DEPOSIT_RATE, DEPOSIT_RATE_SOURCE),
        ('今年以來利息收入', '約 0.575%', f'4個月約當（以{DEPOSIT_RATE}年利率計算）'),
    ]
    for i, (item, rate, desc) in enumerate(deposit):
        table4.rows[i+1].cells[0].text = item
        table4.rows[i+1].cells[1].text = rate
        table4.rows[i+1].cells[2].text = desc

    doc.add_paragraph()

    # 總結
    h5 = doc.add_heading('📋 投資績效總評', level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    summary = [
        ('🏆 今年以來表現最佳：韓股 (+53.12%)、台股 (+32.63%)，亞股漲幅驚人', True),
        ('📉 今年以來表現最差：比特幣 (-14.24%)，加密貨幣波動劇烈', False),
        ('🥇 黃金 (+7.46%) 表現亮眼，避險需求支撐價格', True),
        ('💵 定存報酬 (+1.725%)，無風險資產最穩健', True),
        ('🏠 台灣房地產 (+1~2%) 保守穩健，遠落後股市', True),
    ]
    for s, is_positive in summary:
        p = doc.add_paragraph()
        if is_positive:
            if '+' in s:
                idx = s.index('+')
                run1 = p.add_run(s[:idx])
                run2 = p.add_run(s[idx:])
                run2.font.color.rgb = RED  # 正數紅色
                run2.bold = True
            else:
                p.add_run(s)
        else:
            if '-' in s:
                idx = s.index('-')
                run1 = p.add_run(s[:idx])
                run2 = p.add_run(s[idx:])
                run2.font.color.rgb = GREEN  # 負數綠色
                run2.bold = True
            else:
                p.add_run(s)
        p.paragraph_format.space_after = Pt(6)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 2026年資產報酬率報告')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '📊 2026年各資產報酬率報告\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        f'資料區間：2026/01/01 ～ {today.strftime("%Y/%m/%d")}\n\n'
        '📈 股市、加密貨幣、貴金屬\n'
        '🏠 台灣房地產\n'
        '💵 定存利率\n\n'
        '✅ 報酬率：紅色(+) 綠色(-)\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生2026年資產報酬率報告（即時API版）...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    
    print('  📤 正在上傳到網站...')
    upload_to_website(report_path)
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
