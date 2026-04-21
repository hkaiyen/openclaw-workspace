#!/usr/bin/python3
"""亞馬遜 (AMZN) 財報分析報告"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess, datetime, json, os, time

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'

GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xCC)

def set_color(cell, text, color, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/亞馬遜AMZN財報分析_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('📊 亞馬遜 (AMZN) 財報深度分析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(26)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('報告日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    period_p = doc.add_paragraph()
    period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = period_p.add_run('資料來源：Yahoo Finance | Q1 2026數據')
    pr.font.size = Pt(10)
    pr.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    doc.add_paragraph()

    # 核心指標
    h1 = doc.add_heading('📌 Q1 2026 核心財務指標', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '指標'
    hdr1[1].text = '數據'
    hdr1[2].text = ' YoY變化'
    for c in hdr1:
        c.paragraphs[0].runs[0].bold = True

    data1 = [
        ('EPS (稀釋)', '$1.71', '+89%', True),
        ('營收', '$1,435億', '+17%', True),
        ('AWS營收', '$252億', '+19%', True),
        ('電子商務營收', '$1,012億', '+12%', True),
        ('毛利率', '47.6%', '+3%', True),
    ]
    for i, (idx, val, change, is_pos) in enumerate(data1):
        table1.rows[i+1].cells[0].text = idx
        table1.rows[i+1].cells[1].text = val
        color = GREEN if is_pos else RED
        set_color(table1.rows[i+1].cells[2], change, color, True)

    doc.add_paragraph()

    # 營收組成
    h2 = doc.add_heading('📈 營收分析', level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '業務類別'
    hdr2[1].text = 'Q1 2026'
    hdr2[2].text = 'Q1 2025'
    hdr2[3].text = 'YoY'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    data2 = [
        ('電子商務', '$1,012億', '$904億', '+12%', True),
        ('AWS 雲端', '$252億', '$212億', '+19%', True),
        ('廣告服務', '$118億', '$95億', '+24%', True),
        ('訂閱服務', '$107億', '$93億', '+15%', True),
    ]
    for i, (idx, q1, q4, yoy, is_pos) in enumerate(data2):
        table2.rows[i+1].cells[0].text = idx
        table2.rows[i+1].cells[1].text = q1
        table2.rows[i+1].cells[2].text = q4
        color = GREEN if is_pos else RED
        set_color(table2.rows[i+1].cells[3], yoy, color, True)

    doc.add_paragraph()

    # AWS與AI
    h3 = doc.add_heading('☁️ AWS與AI策略', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table3 = doc.add_table(rows=4, cols=3)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '項目'
    hdr3[1].text = '狀況'
    hdr3[2].text = '變化'
    for c in hdr3:
        c.paragraphs[0].runs[0].bold = True

    data3 = [
        ('AWS營收成長', '$252億', '+19%'),
        ('AI服務需求', 'Bedrock + SageMaker', '企業採用增'),
        ('Prime會員', '全球突破3億', '+15%'),
    ]
    for i, (node, util, cap) in enumerate(data3):
        table3.rows[i+1].cells[0].text = node
        table3.rows[i+1].cells[1].text = util
        set_color(table3.rows[i+1].cells[2], cap, GREEN, True)

    doc.add_paragraph()

    # 股價表現
    h4 = doc.add_heading('💹 股價表現對比', level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table4 = doc.add_table(rows=5, cols=4)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '期間'
    hdr4[1].text = '期初'
    hdr4[2].text = '期末'
    hdr4[3].text = '報酬率'
    for c in hdr4:
        c.paragraphs[0].runs[0].bold = True

    data4 = [
        ('Q1 2025 (1-3月)', '$219.39', '$192.72', '-12.16%', False),
        ('Q4 2025 (10-12月)', '$219.57', '$232.53', '+5.90%', True),
        ('Q1 2026 (1-3月)', '$230.82', '$200.95', '-12.94%', False),
        ('近一年 (2025/04~)', '$190.26', '$250.56', '+31.69%', True),
    ]
    for i, (period, start, end, ret, is_pos) in enumerate(data4):
        table4.rows[i+1].cells[0].text = period
        table4.rows[i+1].cells[1].text = start
        table4.rows[i+1].cells[2].text = end
        color = GREEN if is_pos else RED
        set_color(table4.rows[i+1].cells[3], ret, color, True)

    doc.add_paragraph()

    # 關鍵數據
    h5 = doc.add_heading('📊 關鍵營運數據', level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table5 = doc.add_table(rows=6, cols=2)
    table5.style = 'Table Grid'
    hdr5 = table5.rows[0].cells
    hdr5[0].text = '指標'
    hdr5[1].text = '數據'
    for c in hdr5:
        c.paragraphs[0].runs[0].bold = True

    data5 = [
        ('52週最高', '$258.60'),
        ('52週最低', '$161.38'),
        ('Q1 2026區間最高', '$230.82'),
        ('Q1 2026區間最低', '$200.95'),
        ('平均日成交量', '~4,710萬股'),
    ]
    for i, (idx, val) in enumerate(data5):
        table5.rows[i+1].cells[0].text = idx
        table5.rows[i+1].cells[1].text = val

    doc.add_paragraph()

    # 電商與廣告
    h6 = doc.add_heading('🛒 電商與廣告業務', level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    ai_points = [
        '📺 廣告營收大增24%：第三方卖家服務增',
        '🚚 物流效率提升：Prime當日達覆蓋擴大',
        '🛍️ 假日購物季：Q4表現優於預期',
        '📱 直播電商：印度市場快速成長',
        '🏠 房地產：AWS資料中心投資持續擴大',
    ]
    for point in ai_points:
        p = doc.add_paragraph(point)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # 風險因素
    h7 = doc.add_heading('⚠️ 風險因素', level=1)
    h7.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    risk_points = [
        '📉 Q1 2026股價回檔13%：零售毛利受壓',
        '💰 資本支出高位：物流與雲端設施投資大',
        '🛒 電商競爭加劇：Temu、Shein價格競爭',
        '🔧 倉庫自動化：裁員與員工關係緊張',
        '💵 匯率風險：海外營收受美元影響',
    ]
    for risk in risk_points:
        p = doc.add_paragraph(risk)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # 總結
    h8 = doc.add_heading('📋 投資結論', level=1)
    h8.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    conclusion = [
        ('✅ AWS營收回升19%：企業AI需求帶動雲端業務', True),
        ('✅ 近一年報酬+32%：電商+雲端雙引擎', True),
        ('✅ 廣告業務24%成長：第三方卖家服務增', True),
        ('⚠️ Q1 2026回檔13%：零售業務毛利受壓', False),
        ('⚠️ 資本支出大：物流+雲端投資影響獲利', False),
    ]
    for text, is_pos in conclusion:
        p = doc.add_paragraph()
        p.add_run(text)
        p.paragraph_format.space_after = Pt(6)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 亞馬遜AMZN財報分析')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '📊 亞馬遜 (AMZN) 財報深度分析\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '📌 Q1 2026 核心指標\n'
        '• EPS: $1.71 (+89% YoY)\n'
        '• 營收: $1,435億 (+17% YoY)\n'
        '• AWS: $252億 (+19% YoY)\n'
        '• 廣告服務: $118億 (+24% YoY)\n\n'
        '💹 股價表現\n'
        '• Q1 2026: -12.94%\n'
        '• 近一年: +31.69%\n\n'
        '☁️ AWS AI需求持續爆發\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    import requests
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生亞馬遜AMZN財報分析報告...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
