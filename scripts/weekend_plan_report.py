#!/usr/bin/python3
"""
週末行程規劃報告
每週五下午4:00執行
讀取行事曆（CalDAV）+ 天氣預報 + 自動建議行程
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import datetime
import requests
from requests.auth import HTTPBasicAuth
import re

# ========== iCloud CalDAV 設定 ==========
ICLOUD_ID = 'Hkaiyen@iCloud.com'
ICLOUD_PASSWORD = 'lpve-dfwe-spxv-pcdh'

# 行事曆 ID
CALENDARS = {
    'Work': '33C297E9-D9DA-4C60-9E99-AC81AFFD2044',
    'Family': '209C2820-1EE1-45A2-9E18-51CAB6393454',
    '台灣節日': 'C7C84572-8C29-40C4-97BE-216356344150'
}

TELEGRAM_TOKEN = "8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw"
TELEGRAM_CHAT_ID = "8779713208"

# CalDAV auth
auth = HTTPBasicAuth(ICLOUD_ID, ICLOUD_PASSWORD)
BASE_URL = 'https://caldav.icloud.com/1056470819/calendars'

def parse_vevent(content):
    """解析 VEVENT 取得時間和標題"""
    patterns = [
        r'DTSTART;VALUE=DATE:(\d{8})',
        r'DTSTART;.*?:(\d{8}T\d{6})',
        r'DTSTART:(\d{8}T\d{6})',
        r'DTSTART;VALUE=DATE:(\d{8})',
    ]
    
    dt = None
    for pattern in patterns:
        m = re.search(pattern, content)
        if m:
            dt_str = m.group(1)
            try:
                if 'T' in dt_str:
                    dt = datetime.datetime.strptime(dt_str, '%Y%m%dT%H%M%S')
                else:
                    dt = datetime.datetime.strptime(dt_str, '%Y%m%d')
                break
            except:
                pass
    
    m = re.search(r'SUMMARY[:]?(.*)', content)
    summary = m.group(1).strip() if m else '(無標題)'
    
    return dt, summary

def get_calendar():
    """取得週六週日的行事曆（CalDAV）"""
    today = datetime.datetime.now()
    
    # 找到本週六和週日
    days_until_sat = (6 - today.weekday()) % 7
    if days_until_sat == 0 and today.weekday() == 6:  # 如果今天是週日
        days_until_sat = 6
    sat = today + datetime.timedelta(days=days_until_sat)
    sat_start = sat.replace(hour=0, minute=0, second=0, microsecond=0)
    sat_end = sat_start + datetime.timedelta(days=2)
    
    sun_start = sat_start + datetime.timedelta(days=1)
    sun_end = sat_start + datetime.timedelta(days=2)
    
    all_events = []
    
    headers = {'Depth': '1', 'Content-Type': 'application/xml'}
    propfind = '''<?xml version="1.0" encoding="UTF-8"?>
<dav:propfind xmlns:dav="DAV:"><dav:prop><dav:displayname/></dav:prop></dav:propfind>'''
    
    for cal_name, cal_id in CALENDARS.items():
        url = f'{BASE_URL}/{cal_id}/'
        
        try:
            resp = requests.request('PROPFIND', url, auth=auth, headers=headers, 
                                   data=propfind.encode(), timeout=30)
            
            if resp.status_code == 207:
                hrefs = re.findall(r'href>([^<]+\.ics)<', resp.text)
                
                for href in hrefs[:50]:
                    ics_url = f'https://caldav.icloud.com{href}'
                    try:
                        ics_resp = requests.get(ics_url, auth=auth, timeout=5)
                        if ics_resp.status_code == 200:
                            vevent = re.search(r'BEGIN:VEVENT\s*(.*?)\s*END:VEVENT', 
                                              ics_resp.text, re.DOTALL)
                            if vevent:
                                dt, summary = parse_vevent(vevent.group(1))
                                if dt:
                                    # 檢查是否在週六或週日
                                    if sat_start <= dt < sat_end:
                                        all_events.append((dt, summary, '週六'))
                                    elif sun_start <= dt < sun_end:
                                        all_events.append((dt, summary, '週日'))
                    except:
                        pass
        except Exception as e:
            print(f"Error fetching {cal_name}: {e}")
    
    # 排序
    all_events.sort(key=lambda x: x[0])
    
    if not all_events:
        return "✅ 週末沒有既定行程，可以自由安排！", False
    
    output = ""
    current_date = None
    for dt, summary, day in all_events:
        if dt.date() != current_date:
            current_date = dt.date()
            output += f"📅 [{day}] {dt.strftime('%Y/%m/%d')}\n"
        time_str = dt.strftime('%H:%M') if dt.hour or dt.minute else '全天'
        output += f"  {time_str} - {summary}\n"
    
    return output, True

def get_weather():
    """取得週末天氣（Taipei）"""
    try:
        url = "https://wttr.in/Taipei?format=%c+%t&lang=zh"
        result = subprocess.run(['curl', '-s', url], capture_output=True, text=True, timeout=10)
        if result.stdout:
            return result.stdout.strip()
    except:
        pass
    
    return "天氣資料讀取失敗"

def suggest_activities(weather_text, has_calendar):
    """根據天氣和行事曆建議活動"""
    suggestions = []
    
    if '雨' in weather_text or 'rain' in weather_text.lower():
        suggestions.append("🌧️ 室內活動建議：")
        suggestions.append("  • 逛百貨公司（SOGO、新光三越）")
        suggestions.append("  • 電影欣賞（威秀、國賓）")
        suggestions.append("  • 餐廳美食（母親節優惠中）")
        suggestions.append("  • 博物館/展覽（台北當代藝術館）")
        suggestions.append("  • 按摩/SPA 放鬆")
    else:
        suggestions.append("☀️ 戶外活動建議：")
        suggestions.append("  • 陽明山國家公園（賞花/健行）")
        suggestions.append("  • 淡水老街/八里騎單車")
        suggestions.append("  • 象山步道（看101夜景）")
        suggestions.append("  • 河濱公園野餐")
        suggestions.append("  • 動物園/兒童新樂園")
    
    if not has_calendar:
        suggestions.append("")
        suggestions.append("📝 自由時間建議：")
        suggestions.append("  • 閱讀一本好書")
        suggestions.append("  • 學習新技能（烹飪、語言）")
        suggestions.append("  • 整理房間/斷捨離")
        suggestions.append("  • 運動健身（健身房、游泳）")
    
    return suggestions

def generate_report():
    today = datetime.datetime.now()
    
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # 標題
    title = doc.add_heading(f'🗓️ 週末行程規劃', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(f'{today.strftime("%Y年%m月%d日")} 週五六日行程建議')
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    sr.font.italic = True
    
    doc.add_paragraph()
    
    # ===== 天氣預報 =====
    h1 = doc.add_heading('🌤️ 週末天氣預報', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0xC0)
    
    weather = get_weather()
    p = doc.add_paragraph()
    p.add_run(f'台北天氣：{weather}')
    p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    # ===== 行事曆（CalDAV）=====
    h2 = doc.add_heading('📅 週末行事曆', level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    calendar_text, has_events = get_calendar()
    
    if has_events:
        for line in calendar_text.split('\n'):
            if line.strip():
                p = doc.add_paragraph()
                p.add_run(line)
                p.paragraph_format.space_after = Pt(3)
                p.runs[0].font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        p.add_run('✅ 週末沒有既定行程，可以自由安排！')
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    
    doc.add_paragraph()
    
    # ===== 行程建議 =====
    h3 = doc.add_heading('💡 行程建議', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0xC0, 0x80, 0x00)
    
    suggestions = suggest_activities(weather, has_events)
    for line in suggestions:
        if line.strip():
            p = doc.add_paragraph()
            p.add_run(line)
            p.paragraph_format.space_after = Pt(3)
            p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # 頁尾
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run(f'小安助理 · {today.strftime("%Y年%m月%d日")}')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    
    return doc

def send_telegram(docx_path, caption):
    url = f'https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendDocument'
    with open(docx_path, 'rb') as f:
        files = {'document': f}
        data = {'chat_id': TELEGRAM_CHAT_ID, 'caption': caption}
        resp = requests.post(url, files=files, data=data, timeout=30)
    return resp.status_code == 200

if __name__ == '__main__':
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    
    print("🗓️ 開始生成週末行程規劃報告...")
    doc = generate_report()
    
    filename = f'週末行程規劃_{date_str}.docx'
    filepath = f'/root/.openclaw/reports/weekly/{filename}'
    doc.save(filepath)
    print(f"✅ 已儲存: {filepath}")
    
    caption = f"""🗓️ 週末行程規劃報告

{today.strftime('%Y年%m月%d日')}

🌤️ 天氣預報 · 📅 行事曆
💡 行程建議

由 小安助理 每週五下午自動生成"""

    if send_telegram(filepath, caption):
        print("✅ 已發送到 Telegram")
    else:
        print("⚠️ 發送失敗")