#!/usr/bin/python3
# -*- coding: utf-8 -*-
"""
飆股篩選系統 v1（川寶投顧專用）
================================

功能：
1. 強勢股10檔：依成交量 + 漲跌幅加權評分
2. 飆股10檔：依漲幅排序（從上市上櫃全市場）

資料來源：TWSE（台灣證券交易所）
"""

import requests
import datetime
import json
import os

# ========== 設定 ==========
BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
REPORT_DIR = '/root/.openclaw/reports/daily'
SCREENER_FILE = os.path.join(REPORT_DIR, 'screener_history.json')


def parse_change(row):
    """解析漲跌方向和價差"""
    color_html = str(row[9]) if row[9] else ''
    if 'color:red' in color_html:
        direction = 1  # 漲
    elif 'color:green' in color_html:
        direction = -1  # 跌
    else:
        direction = 0  # 平盤
    
    change_str = str(row[10]).replace(',', '') if row[10] and str(row[10]) not in ['--', ''] else '0'
    change = float(change_str) * direction if change_str and change_str not in ['--', ''] else 0
    return direction, change


def get_all_twse_stocks():
    """從 TWSE 取得所有上市股票 """
    try:
        date_str = datetime.datetime.now().strftime('%Y%m%d')
        url = f'https://www.twse.com.tw/rwd/zh/afterTrading/MI_INDEX?date={date_str}&type=ALL&response=json'
        headers = {'User-Agent': 'Mozilla/5.0'}
        r = requests.get(url, headers=headers, timeout=30)
        
        if r.status_code == 200:
            data = r.json()
            if data.get('stat') == 'OK':
                tables = data.get('tables', [])
                
                for table in tables:
                    rows = table.get('data', [])
                    if len(rows) > 1000:
                        stocks = []
                        for row in rows:
                            if len(row) >= 11:
                                symbol = str(row[0]).strip()
                                if symbol.isdigit() and len(symbol) == 4:
                                    try:
                                        name = str(row[1]).strip()
                                        close_str = str(row[8]).replace(',', '') if row[8] and str(row[8]) not in ['--', ''] else '0'
                                        vol_str = str(row[2]).replace(',', '') if row[2] and str(row[2]) not in ['--', ''] else '0'
                                        
                                        if close_str and close_str not in ['--', '']:
                                            close = float(close_str)
                                            direction, change = parse_change(row)
                                            volume = int(vol_str.replace(',', '')) if vol_str and vol_str not in ['--', ''] else 0
                                            
                                            if close > 0 and volume > 0:
                                                prev_close = close - change
                                                change_pct = (change / prev_close * 100) if prev_close > 0 else 0
                                                
                                                stocks.append({
                                                    'symbol': f'{symbol}.TW',
                                                    'name': name,
                                                    'price': close,
                                                    'change': change,
                                                    'change_pct': round(change_pct, 2),
                                                    'volume': volume
                                                })
                                    except:
                                        pass
                        return stocks
    except Exception as e:
        print(f"TWSE API 錯誤: {e}")
    return []


def screen_stocks():
    """篩選強勢股和飆股"""
    print("=" * 60)
    print("📊 飆股篩選系統啟動（全市場）")
    print("=" * 60)
    print(f"時間: {datetime.datetime.now()}")
    
    print("\n📥 抓取上市股票...")
    all_stocks = get_all_twse_stocks()
    print(f"  取得 {len(all_stocks)} 檔股票")
    
    if not all_stocks:
        print("\n❌ 無法取得股票資料")
        return [], []
    
    # ===== 強勢股：成交量 + 漲幅加權評分 =====
    print("\n🏆 評分強勢股...")
    scored = []
    for s in all_stocks:
        if s['volume'] > 0 and s['price'] > 0:
            score = s['change_pct'] * 3 + min(s['volume'] / 1000000, 15)
            if s['change_pct'] >= 9.5:
                score += 50
            scored.append({
                'symbol': s['symbol'],
                'name': s['name'],
                'price': s['price'],
                'change': s['change'],
                'change_pct': s['change_pct'],
                'volume': s['volume'],
                'score': round(score, 2)
            })
    
    scored.sort(key=lambda x: x['score'], reverse=True)
    strong_top10 = scored[:10]
    
    # ===== 飆股：純漲幅排序 =====
    gainers = [s for s in all_stocks if s['change_pct'] > 0]
    gainers.sort(key=lambda x: x['change_pct'], reverse=True)
    rockets_top10 = gainers[:10]
    
    # 顯示結果
    print("\n" + "=" * 50)
    print("🏆 強勢股 Top 10（成交量+動能評分）：")
    for i, s in enumerate(strong_top10, 1):
        flag = '🔥' if s['change_pct'] >= 9.5 else ''
        print(f"  {i:2}. {s['name']}({s['symbol']}) {flag} | ${s['price']} | {s['change_pct']:+.2f}%")
    
    print("\n🚀 飆股 Top 10（純漲幅）：")
    for i, s in enumerate(rockets_top10, 1):
        flag = '🔥' if s['change_pct'] >= 9.5 else ''
        print(f"  {i:2}. {s['name']}({s['symbol']}) {flag} | ${s['price']} | {s['change_pct']:+.2f}%")
    
    save_result(strong_top10, rockets_top10)
    
    return strong_top10, rockets_top10


def save_result(strong, rockets):
    """儲存結果"""
    os.makedirs(REPORT_DIR, exist_ok=True)
    
    history = {}
    if os.path.exists(SCREENER_FILE):
        with open(SCREENER_FILE, 'r') as f:
            history = json.load(f)
    
    today = datetime.datetime.now().strftime('%Y-%m-%d')
    history[today] = {'strong': strong, 'rockets': rockets}
    history['last_update'] = today
    
    with open(SCREENER_FILE, 'w') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


def generate_report(strong_top10, rockets_top10):
    """生成報告"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # 標題
    title = doc.add_heading('📈 飆股篩選報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run(datetime.datetime.now().strftime('%Y年%m月%d日'))
    dr.font.size = Pt(14)
    dr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    
    # ===== 強勢股 =====
    doc.add_heading('🏆 強勢股 Top 10', level=1)
    p = doc.add_paragraph()
    p.add_run('評分原則：').bold = True
    p.add_run('漲幅×3 + 成交量加成 + 漲停額外加分')
    
    table1 = doc.add_table(rows=11, cols=5)
    table1.style = 'Table Grid'
    headers = ['排名', '代號', '名稱', '價格', '漲跌']
    for i, h in enumerate(headers):
        table1.rows[0].cells[i].text = h
        table1.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    for i, s in enumerate(strong_top10, 1):
        table1.rows[i].cells[0].text = str(i)
        table1.rows[i].cells[1].text = s['symbol']
        table1.rows[i].cells[2].text = s['name']
        table1.rows[i].cells[3].text = f"${s['price']:.2f}"
        
        change_text = f"{s['change_pct']:+.2f}%"
        if s['change_pct'] >= 9.5:
            change_text += ' 🔥'
        table1.rows[i].cells[4].text = change_text
        
        clr = RGBColor(0x00, 0x80, 0x00) if s['change_pct'] > 0 else RGBColor(0xFF, 0x00, 0x00)
        table1.rows[i].cells[4].paragraphs[0].runs[0].font.color.rgb = clr
    
    doc.add_paragraph()
    
    # ===== 飆股 =====
    doc.add_heading('🚀 飆股 Top 10（純漲幅）', level=1)
    p2 = doc.add_paragraph()
    p2.add_run('篩選原則：').bold = True
    p2.add_run('從上市上櫃全市場，依今日漲幅排序')
    
    table2 = doc.add_table(rows=11, cols=4)
    table2.style = 'Table Grid'
    headers2 = ['排名', '代號', '名稱', '今日漲跌']
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    for i, s in enumerate(rockets_top10, 1):
        table2.rows[i].cells[0].text = str(i)
        table2.rows[i].cells[1].text = s['symbol']
        table2.rows[i].cells[2].text = s['name']
        
        change_text = f"{s['change_pct']:+.2f}%"
        if s['change_pct'] >= 9.5:
            change_text = '🔥 ' + change_text
        table2.rows[i].cells[3].text = change_text
        
        clr = RGBColor(0x00, 0x80, 0x00) if s['change_pct'] > 0 else RGBColor(0xFF, 0x00, 0x00)
        table2.rows[i].cells[3].paragraphs[0].runs[0].font.color.rgb = clr
    
    doc.add_paragraph()
    
    # 聲明
    disc = doc.add_paragraph()
    disc_text = disc.add_run('本報告僅供參考，不構成投資建議。投資有風險，請自行評估。｜川寶投顧｜小安製作')
    disc_text.font.size = Pt(10)
    disc_text.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    disc_text.font.italic = True
    
    return doc


def send_telegram_report(doc):
    """發送報告"""
    docx_path = f"{REPORT_DIR}/飆股篩選_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    doc.save(docx_path)
    print(f"✅ DOCX 已儲存: {docx_path}")
    
    with open(docx_path, 'rb') as f:
        files = {'document': f}
        data = {
            'chat_id': CHAT_ID,
            'caption': f"📈 飆股篩選報告 {datetime.datetime.now().strftime('%Y.%m.%d')}｜川寶投顧"
        }
        r = requests.post(f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument', data=data, files=files, timeout=60)
        print(f"Telegram: {'✅' if r.json().get('ok') else '❌'}")
    
    return docx_path


def main():
    print("\n" + "=" * 60)
    print("🚀 飆股篩選系統 - 川寶投顧（全市場）")
    print("=" * 60)
    
    strong, rockets = screen_stocks()
    
    if strong and rockets:
        doc = generate_report(strong, rockets)
        output = send_telegram_report(doc)
        print(f"\n✅ 完成！")
    else:
        print("\n❌ 無法取得資料，請稍後再試")


if __name__ == '__main__':
    main()
