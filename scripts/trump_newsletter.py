#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
川普電子報 - 每日自動執行腳本（完整版 v3）
==========================================

排程：每天早上 03:00

流程：
1. sessions_spawn 小歐+千問+拉瑪+撈仔+小安 同時研究（5個subagent並行）
2. 小安蒐集並整理成完整報告
3. 小安生成每日金句
4. 潔咪用 Gemini 生成早安圖
5. 小安合併成 Word 檔發送到 Telegram

"""

import subprocess
import json
import time
import urllib.request
import urllib.parse
import datetime
import random
import os
import requests
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== 常數 =====
BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
REPORT_DIR = '/root/.openclaw/reports/daily'
OPENCLAW_DIR = os.path.expanduser('~/.openclaw/workspace')
# Gemini API Key
GEMINI_API_KEY = ''  # 請填入您的 Gemini API Key
IMAGE_DIR = '/root/.openclaw/media'

# ===== 每日金句庫 =====
DAILY_QUOTES = [
    "知識是最可靠的資產。",
    "時間是最好的老師。",
    "專注是成功的關鍵。",
    "改變是成長的開始。",
    "學習是永恆的主題。",
    "信任是合作的基礎。",
    "創新是突破的泉源。",
    "謹慎是投資的原則。",
    "耐心是財富的階梯。",
    "纪律是自由的保障。",
    "洞察是決策的依據。",
    "平衡是生活的藝術。",
    "效率是時間的價值。",
    "規劃是未來的藍圖。",
    "自律是自由的起點。",
    "行動是夢想的階梯。",
    "謙遜是智慧的開始。",
    "選擇比努力更重要。",
    "失敗是成功的養分。",
    "方向比速度更關鍵。",
    "累積小的成功，大的改變。",
    "持續優化，而非追求完美。",
    "知識改變命運，學習改變人生。",
    "時間花在哪，成就就在哪。",
    "專注於你能改變的事。",
    "今天的努力，明天的收獲。",
    "行動勝於空談。",
    "學習是最佳的投資。",
    "勇敢踏出舒適圈。",
    "保持好奇，保持成長。",
]

IMAGE_PROMPTS = [
    "一本打開的書漂浮在星空背景上，書頁散發柔和光芒，文字清晰可讀，極簡風格，深藍色夜空襯托。",
    "一座燈塔佇立在海浪中，光芒穿透雲層照亮遠方，寫實與夢幻結合。",
    "一棵樹苗在陽光下茁壯成長，背景是藍天白雲，文青清新風。",
    "K線圖向上延伸變成一棵金色樹，數字與圖表交織，金融科技風。",
    "一枚硬幣在月光下閃閃發光，遠處城市天際線剪影，奢華低調風。",
    "清晨薄霧中的森林小路，光線穿透樹葉形成耶穌光，自然寧靜風。",
    "一隻狐狸站在雪山山頂眺望遠方，背景壯麗山景，電影感寬幅。",
    "俯瞰城市夜景萬家燈火，霓虹燈光交織，未來科技感。",
    "一位行人撐傘走在東京街頭，雨滴在路面積水倒映，日系電影感。",
    "一幅畫架上的油畫，油畫內容是抽象的彩色線條，背景畫室暖光，藝術氣息。",
    "復古打字機放在木桌上，紙張寫著文字，背景老式書房，懷舊文學風。",
    "太空人漂浮在星空中，地球在背景藍色光芒，科幻電影海報風。",
    "銀河系壯觀鳥瞰視角，各種星球散落，唯美星空風。",
]

def get_daily_quote():
    """根據當天日期選擇金句（每天不同）"""
    today = datetime.datetime.now()
    day_of_year = today.timetuple().tm_yday
    return DAILY_QUOTES[day_of_year % len(DAILY_QUOTES)]

def get_daily_image_prompt(quote):
    """根據當天日期選擇圖片生成提示（每天不同風格）"""
    today = datetime.datetime.now()
    day_of_year = today.timetuple().tm_yday
    prompt_template = IMAGE_PROMPTS[day_of_year % len(IMAGE_PROMPTS)]
    return f"每日金句：「{quote}」。設計：{prompt_template}"

def generate_gemini_image(prompt, output_path):
    """使用 Gemini API 生成圖片"""
    import base64
    
    print(f"🎨 正在生成圖片...")
    
    GEMINI_API_KEY = 'AIzaSyCrUO2AlTi69NtKNdwDgvBaYppe2EU7-Jw'
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.1-flash-image-preview:generateContent?key={GEMINI_API_KEY}"
    
    payload = {
        "contents": [{
            "parts": [{"text": f"請生成一張精美的圖片，主題是：{prompt}"}]
        }]
    }
    
    try:
        data = json.dumps(payload).encode('utf-8')
        req = urllib.request.Request(url, data=data, headers={'Content-Type': 'application/json'})
        
        with urllib.request.urlopen(req, timeout=300) as response:
            result = json.loads(response.read().decode('utf-8'))
        
        if 'candidates' in result:
            parts = result['candidates'][0].get('content', {}).get('parts', [])
            for part in parts:
                if 'inlineData' in part:
                    image_data = part['inlineData']['data']
                    image_bytes = base64.b64decode(image_data)
                    with open(output_path, 'wb') as f:
                        f.write(image_bytes)
                    print(f"✅ 圖片已生成: {output_path}")
                    return True
        
        print(f"❌ 圖片生成失敗: {result}")
        return False
        
    except Exception as e:
        print(f"❌ 圖片生成錯誤: {e}")
        return False

def send_telegram_document(file_path, caption):
    """發送文件到Telegram"""
    cmd = [
        'curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
        '-F', f'chat_id={CHAT_ID}',
        '-F', f'document=@{file_path}',
        '-F', f'caption={caption}'
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    try:
        response = json.loads(result.stdout)
        return response.get('ok', False)
    except:
        return False

def send_telegram_photo(file_path, caption):
    """發送圖片到Telegram"""
    cmd = [
        'curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{BOT_TOKEN}/sendPhoto',
        '-F', f'chat_id={CHAT_ID}',
        '-F', f'photo=@{file_path}',
        '-F', f'caption={caption}'
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    try:
        response = json.loads(result.stdout)
        return response.get('ok', False)
    except:
        return False

def generate_word_report(quote, image_path, results_dict):
    """生成Word報告（含圖片）"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 封面標題
        title = doc.add_heading('🦅 川普電子報', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_str = datetime.datetime.now().strftime('%Y年%m月%d日')
        doc.add_paragraph(f'日期：{date_str}')
        doc.add_paragraph('')
        
        # 插入圖片
        if image_path and os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"⚠️ 圖片插入失敗: {e}")
        
        doc.add_paragraph('')
        
        # 每日金句
        doc.add_heading(f'💬 每日金句', level=1)
        quote_para = doc.add_paragraph(quote)
        quote_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')
        
        # 各助理報告
        sections = [
            ('🐰 小安｜總體經濟分析師', 'xiaoan'),
            ('💬 小歐｜財務分析師', 'xiaoou'),
            ('🔍 千問｜技術分析師', 'qianwen'),
            ('📚 拉瑪｜趨勢分析師', 'lama'),
            ('🐂 撈仔｜創意分析師', 'laozai'),
        ]
        
        for section_title, key in sections:
            doc.add_heading(section_title, level=1)
            if key in results_dict:
                doc.add_paragraph(results_dict[key])
            else:
                doc.add_paragraph(f'{key} 報告待補充...')
            doc.add_paragraph('')
        
        # 結尾
        doc.add_paragraph('—' * 30)
        doc.add_paragraph(f'川寶投顧小安助理 | {date_str}')
        doc.add_paragraph('本報告僅供參考，不構成投資建議。')
        
        # 儲存
        today_str = datetime.datetime.now().strftime('%Y%m%d')
        report_path = os.path.join(REPORT_DIR, f'川普電子報_{today_str}.docx')
        doc.save(report_path)
        print(f"✅ Word報告已生成: {report_path}")
        return report_path
    except Exception as e:
        print(f"❌ Word生成失敗: {e}")
        return None

# ===== 主程式 =====
def main():
    print("=" * 60)
    print("川普電子報 - 每日自動執行（完整版 v3）")
    print("=" * 60)
    print(f"時間：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # 取得今日金句
    quote = get_daily_quote()
    print(f"\n💬 今日金句：{quote}")
    
    # 生成圖片
    today_str = datetime.datetime.now().strftime('%Y%m%d')
    image_path = os.path.join(IMAGE_DIR, f'daily_quote_{today_str}.png')
    
    prompt = get_daily_image_prompt(quote)
    if generate_gemini_image(prompt, image_path):
        print(f"✅ 圖片生成成功")
    else:
        print(f"⚠️ 圖片生成失敗，將使用無圖片版本")
        image_path = None
    
    # 生成Word報告（這裡假設已經有各助理的結果）
    # 實際執行時需要 sessions_spawn 各助理後再彙整
    # 目前用 placeholder
    
    results = {
        'xiaoan': '（等待助理研究結果）',
        'xiaoou': '（等待助理研究結果）',
        'qianwen': '（等待助理研究結果）',
        'lama': '（等待助理研究結果）',
        'laozai': '（等待助理研究結果）',
    }
    
    # 生成Word
    docx_path = generate_word_report(quote, image_path, results)
    
    if docx_path:
        # 發送圖片
        if image_path and os.path.exists(image_path):
            send_telegram_photo(image_path, f"💬 每日金句圖 | {quote}")
        
        # 發送Word
        caption = f"🦅 川普電子報 {datetime.datetime.now().strftime('%Y年%m月%d日')} | 川寶投顧"
        if send_telegram_document(docx_path, caption):
            print("✅ 已發送到Telegram")
        else:
            print("❌ 發送失敗")
    
    print("\n" + "=" * 60)
    print("完成！")
    print("=" * 60)

if __name__ == '__main__':
    main()