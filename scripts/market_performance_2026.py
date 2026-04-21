#!/usr/bin/python3
"""
2026年股市績效檢視
- 年初（2026/01/02）到目前的報酬率追蹤
- 台股：TWSE API
- 美股：Finnhub API
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess, datetime, json, os, time

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
FINNHUB_API_KEY = os.environ.get('FINNHUB_API_KEY', 'd7btfs1r01quh9fbn7m0d7btfs1r01quh9fbn7mg')

# Yahoo Finance timestamp for 2026-01-02
YAHOO_START = 1767283200
YAHOO_END = 1767542400

# ========== 工具函數 ==========
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def make_header_cell(cell, text, bg='1F497D'):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def add_data_cell(cell, text, is_positive=None, size=10):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(size)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_positive is True:
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    elif is_positive is False:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

# ========== Yahoo Finance (Historical) ==========
def get_yahoo_base_price(symbol):
    """Get historical price from Yahoo Finance for 2026-01-02"""
    try:
        url = f'https://query1.finance.yahoo.com/v8/finance/chart/{symbol}?period1={YAHOO_START}&period2={YAHOO_END}&interval=1d'
        result = subprocess.run(['curl', '-s', '-H', 'User-Agent: Mozilla/5.0', url],
                             capture_output=True, text=True, timeout=15)
        data = json.loads(result.stdout)
        result_data = data.get('chart', {}).get('result', [])
        if result_data:
            opens = result_data[0].get('indicators', {}).get('quote', [{}])[0].get('open', [])
            for o in opens:
                if o is not None:
                    return o
    except:
        pass
    return None

# ========== Finnhub ==========
def get_finnhub_quote(symbol):
    try:
        result = subprocess.run(
            ['curl', '-s', f'https://finnhub.io/api/v1/quote?symbol={symbol}&token={FINNHUB_API_KEY}'],
            capture_output=True, text=True, timeout=10
        )
        return json.loads(result.stdout)
    except:
        return None

# ========== TWSE ==========
def get_twse_data(stock_no, date_str):
    try:
        url = f'https://www.twse.com.tw/rwd/zh/afterTrading/STOCK_DAY?date={date_str}&stockNo={stock_no}'
        result = subprocess.run(['curl', '-s', '--connect-timeout', '15', '--max-time', '45', url],
                                capture_output=True, text=True, timeout=50)
        if not result.stdout:
            return None
        data = json.loads(result.stdout)
        if data.get('stat') == 'OK':
            return data['data']
        return None
    except:
        return None

def get_twse_index(date_str):
    try:
        url = f'https://www.twse.com.tw/rwd/zh/afterTrading/MI_INDEX?response=json&date={date_str}&type=ALL'
        result = subprocess.run(['curl', '-s', '--connect-timeout', '15', '--max-time', '45', url],
                                capture_output=True, text=True, timeout=50)
        if not result.stdout:
            return None
        data = json.loads(result.stdout)
        if data.get('stat') == 'OK':
            for table in data.get('tables', []):
                for row in table.get('data', []):
                    if '發行量加權股價指數' in str(row):
                        return row
        return None
    except:
        return None

def get_weekday(date_str):
    import datetime as dt
    year, month, day = int(date_str[:4]), int(date_str[4:6]), int(date_str[6:8])
    d = dt.date(year, month, day)
    while d.weekday() >= 5:
        d -= dt.timedelta(days=1)
    return d.strftime('%Y%m%d')

# ========== 主程式 ==========
def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    base_date_str = '20260102'
    now = datetime.datetime.now()
    today_str = now.strftime('%Y%m%d')
    twse_date = today_str if now.hour >= 14 else (now - datetime.timedelta(days=1)).strftime('%Y%m%d')
    twse_date = get_weekday(twse_date)
    twse_base = get_weekday(base_date_str)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 標題 =====
    title = doc.add_heading('📈 2026年股市績效檢視', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('台股大盤 · 美股大盤 · 熱門個股與ETF')
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    sr.font.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run(f'績效期間：2026年1月2日 → {today.strftime("%Y年%m月%d日")}')
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== 台股大盤 =====
    h_tw = doc.add_heading('一、台股大盤', level=1)
    h_tw.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    tw_now = get_twse_index(twse_date)
    tw_base = get_twse_index(twse_base)

    table_tw = doc.add_table(rows=2, cols=5)
    table_tw.style = 'Table Grid'
    table_tw.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['指數', '年初價', '目前價', '報酬率', '狀態']):
        make_header_cell(table_tw.rows[0].cells[i], h)

    r = table_tw.rows[1]
    r.cells[0].text = '加權指數'
    if tw_now and tw_base and len(tw_now) >= 5 and len(tw_base) >= 5:
        base_val = float(tw_base[1].replace(',', ''))
        curr_val = float(tw_now[1].replace(',', ''))
        pct = (curr_val - base_val) / base_val * 100
        pos = pct >= 0
        r.cells[1].text = f'{base_val:,.0f}'
        add_data_cell(r.cells[2], f'{curr_val:,.0f}')
        add_data_cell(r.cells[3], f'{pct:+.2f}%', pos)
        r.cells[4].text = '✅ 正報酬' if pos else '❌ 負報酬'
    else:
        for c in range(1, 5):
            r.cells[c].text = 'N/A'

    doc.add_paragraph()

    # ===== 美股大盤 =====
    h1 = doc.add_heading('二、美股大盤', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Get US base prices from Yahoo Finance (2026-01-02)
    us_base_prices = {}
    for sym in ['SPY', 'QQQ', 'DIA', 'AAPL', 'MSFT', 'GOOGL', 'AMZN', 'META', 'NVDA', 'TSLA']:
        price = get_yahoo_base_price(sym)
        if price:
            us_base_prices[sym] = price
            print(f'  {sym} base: {price:.2f}')

    us_etfs = [
        ('S&P 500 (SPY)', 'SPY'),
        ('Nasdaq (QQQ)', 'QQQ'),
        ('Dow Jones (DIA)', 'DIA'),
    ]
    table1 = doc.add_table(rows=4, cols=5)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['指數/ETF', '年初價', '目前價', '報酬率', '趨勢']):
        make_header_cell(table1.rows[0].cells[i], h)

    for ri, (name, sym) in enumerate(us_etfs):
        q = get_finnhub_quote(sym)
        row = table1.rows[ri + 1]
        row.cells[0].text = name
        base_est = us_base_prices.get(sym, 0)
        if q and q.get('o', 0) > 0 and base_est > 0:
            curr = q['o']  # 使用開盤價
            pct = (curr - base_est) / base_est * 100
            pos = pct >= 0
            row.cells[1].text = f'{base_est:.2f}'
            add_data_cell(row.cells[2], f'{curr:.2f}')
            add_data_cell(row.cells[3], f'{pct:+.2f}%', pos)
            row.cells[4].text = '✅ 正報酬' if pos else '❌ 負報酬'
        else:
            for c in range(1, 5):
                row.cells[c].text = 'N/A'

    doc.add_paragraph()

    # ===== 台股熱門標的 =====
    h3 = doc.add_heading('三、台股熱門標的（0050 · 0056 · 2330）', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    note3 = doc.add_paragraph()
    nr = note3.add_run('※ 資料來源：台灣證券交易所（TWSE 2026/01/02 開盤價）')
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    nr.font.italic = True

    tw_info = [
        ('元大台灣50 (0050)', '0050', 66.00),
        ('元大高股息 (0056)', '0056', 36.77),
        ('台積電 (2330.TW)', '2330', 1555.0),
    ]
    table3 = doc.add_table(rows=4, cols=5)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['標的', '年初價', '目前價', '報酬率', '狀態']):
        make_header_cell(table3.rows[0].cells[i], h)

    for ri, (name, sym, base_est) in enumerate(tw_info):
        data = get_twse_data(sym, twse_date)
        row = table3.rows[ri + 1]
        row.cells[0].text = name
        if data and len(data) > 0:
            curr = float(data[-1][3].replace(',', ''))  # 目前開盤價
            pct = (curr - base_est) / base_est * 100
            pos = pct >= 0
            row.cells[1].text = f'{base_est:,.2f}'
            add_data_cell(row.cells[2], f'{curr:,.2f}')
            add_data_cell(row.cells[3], f'{pct:+.2f}%', pos)
            row.cells[4].text = '✅ 正報酬' if pos else '❌ 負報酬'
        else:
            for c in range(1, 5):
                row.cells[c].text = 'N/A'

    doc.add_paragraph()

    # ===== 科技七雄 =====
    h4 = doc.add_heading('四、美股科技七雄（Mag 7）', level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    note4 = doc.add_paragraph()
    nr4 = note4.add_run('※ 資料來源：Finnhub 開盤報價 + Yahoo Finance 2026/01/02 開盤價')
    nr4.font.size = Pt(9)
    nr4.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    nr4.font.italic = True

    mag7 = [
        ('Apple (AAPL)', 'AAPL'),
        ('Microsoft (MSFT)', 'MSFT'),
        ('Google (GOOGL)', 'GOOGL'),
        ('Amazon (AMZN)', 'AMZN'),
        ('Meta (META)', 'META'),
        ('Nvidia (NVDA)', 'NVDA'),
        ('Tesla (TSLA)', 'TSLA'),
    ]
    table4 = doc.add_table(rows=8, cols=5)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['公司', '年初價 (USD)', '目前價 (USD)', '報酬率', '狀態']):
        make_header_cell(table4.rows[0].cells[i], h)

    for ri, (name, sym) in enumerate(mag7):
        q = get_finnhub_quote(sym)
        row = table4.rows[ri + 1]
        row.cells[0].text = name
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        base_est = us_base_prices.get(sym, 0)
        if q and q.get('o', 0) > 0 and base_est > 0:
            curr = q['o']  # 使用開盤價
            pct = (curr - base_est) / base_est * 100
            pos = pct >= 0
            row.cells[1].text = f'${base_est:.2f}'
            add_data_cell(row.cells[2], f'${curr:.2f}')
            add_data_cell(row.cells[3], f'{pct:+.2f}%', pos)
            row.cells[4].text = '✅ 正報酬' if pos else '❌ 負報酬'
            row.cells[4].paragraphs[0].runs[0].font.size = Pt(10)
        else:
            for c in range(1, 5):
                row.cells[c].text = 'N/A'

    doc.add_paragraph()

    # ===== 資料來源 =====
    h6 = doc.add_heading('五、資料來源', level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    sources = [
        ('台灣證券交易所 (TWSE)', '台股大盤、0050、0056、2330 報價'),
        ('Finnhub API', '美股大盤、科技七雄報價'),
    ]
    for src, desc in sources:
        p = doc.add_paragraph()
        p.add_run(f'• {src}：').bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disclaimer.add_run('【聲明】本報告僅供參考，不構成投資建議。年初價為估算值，實際報酬率以券商成交價為準。')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    dr.font.italic = True

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 小安製')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    output_path = f'/root/.openclaw/reports/daily/2026年股市績效檢視_{date_str}.docx'
    doc.save(output_path)
    return output_path

def send_to_telegram(file_path):
    today = datetime.datetime.now()
    caption = f'📈 2026年股市績效檢視\n\n績效期間：2026年1月2日 → {today.strftime("%Y年%m月%d日")}\n\n台股大盤 · 美股大盤 · 熱門個股與ETF\n\n小安製'
    for attempt in range(3):
        result = subprocess.run(
            ['curl', '-s', '-X', 'POST',
             f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
             '-F', f'chat_id={CHAT_ID}',
             '-F', f'document=@{file_path}',
             '-F', f'caption={caption}'],
            capture_output=True, text=True, timeout=30
        )
        try:
            if json.loads(result.stdout).get('ok'):
                return True
        except:
            pass
        time.sleep(2)
    return False

if __name__ == '__main__':
    print(f'[{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}] 開始生成2026年股市績效檢視...')
    try:
        report_path = generate_report()
        print(f'📄 報告已生成: {report_path}')
        if send_to_telegram(report_path):
            print('✅ 已發送到 Telegram')
        else:
            print('❌ 發送失敗')
        print(f'[{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}] 任務完成！')
    except Exception as e:
        print(f'❌ 錯誤: {e}')
