#!/usr/bin/python3
"""
每日簡報生成腳本
整合：天氣、新聞、股市行情、行事曆
"""

import subprocess
import datetime
import os
import sys
import requests
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# API Keys
GROQ_API_KEY = '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'
GROQ_URL = 'https://api.groq.com/openai/v1/chat/completions'
BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'

def get_weather():
    """取得天氣資訊"""
    try:
        resp = requests.get('https://wttr.in/Taipei?format=j1', timeout=10)
        if resp.status_code == 200:
            data = resp.json()
            current = data['current_condition'][0]
            return {
                'temp': current['temp_C'],
                'condition': current['weatherDesc'][0]['value'],
                'humidity': current['humidity'],
            }
    except:
        pass
    return None

def get_news_summary():
    """使用 Groq AI 生成新聞摘要"""
    prompt = """你是台灣財經助理，請根據以下今日新聞標題，生成一份簡短的早晨簡報。

請用繁體中文回答，格式如下：
1. 今日頭條（1-2句）
2. 財經重點（2-3條）
3. 國際要聞（1-2條）

注意：請基於事實回答，不要捏造新聞內容。"""

    try:
        resp = requests.post(GROQ_URL, headers={
            'Authorization': f'Bearer {GROQ_API_KEY}',
            'Content-Type': 'application/json'
        }, json={
            'model': 'qwen/qwen3-32b',
            'messages': [{'role': 'user', 'content': prompt}],
            'max_tokens': 500,
            'temperature': 0.5
        }, timeout=60)
        if resp.status_code == 200:
            return resp.json()['choices'][0]['message']['content']
    except:
        pass
    return "新聞摘要取得失敗。"

def get_market_summary():
    """取得股市摘要"""
    prompt = """請用繁體中文提供今日（2026年4月15日）台股早盤簡報：
1. 預估開盤狀況
2. 重點關注個股/族群
3. 法人動向預測

請用2-3句話簡短回答。"""

    try:
        resp = requests.post(GROQ_URL, headers={
            'Authorization': f'Bearer {GROQ_API_KEY}',
            'Content-Type': 'application/json'
        }, json={
            'model': 'qwen/qwen3-32b',
            'messages': [{'role': 'user', 'content': prompt}],
            'max_tokens': 300,
            'temperature': 0.5
        }, timeout=60)
        if resp.status_code == 200:
            return resp.json()['choices'][0]['message']['content']
    except:
        pass
    return "股市摘要取得失敗。"

def create_briefing_doc(weather, news, market, today):
    """建立每日簡報 Word 文件"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 標題
    title = doc.add_heading('🌅 每日晨間簡報', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 日期
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(today.strftime('%Y年%m月%d日 %A'))
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()

    # 天氣
    if weather:
        p = doc.add_paragraph()
        p.add_run('🌤️ 今日天氣\n').bold = True
        p.add_run(f"台北市：{weather['temp']}°C，{weather['condition']}，濕度{weather['humidity']}%")
    else:
        p = doc.add_paragraph()
        p.add_run('🌤️ 今日天氣\n').bold = True
        p.add_run('天氣資料取得失敗')

    doc.add_paragraph()

    # 新聞摘要
    p = doc.add_paragraph()
    p.add_run('📰 晨間新聞\n').bold = True
    p.add_run(news)

    doc.add_paragraph()

    # 股市摘要
    p = doc.add_paragraph()
    p.add_run('📈 台股晨報\n').bold = True
    p.add_run(market)

    doc.add_paragraph()

    # 勵志語
    p = doc.add_paragraph()
    p.add_run('💪 每日正能量\n').bold = True
    p.add_run('新的一天，新的機會！祝您今天事事順利。')

    doc.add_paragraph()

    # 頁尾
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run(f'小安助理 · {today.strftime("%Y年%m月%d日")}')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # 儲存
    output_path = f'/root/.openclaw/reports/daily/每日簡報_{today.strftime("%Y%m%d_%H%M")}.docx'
    doc.save(output_path)
    return output_path

def send_to_telegram(doc_path, today):
    """發送到 Telegram"""
    caption = f"🌅 每日晨間簡報_{today.strftime('%Y年%m月%d日')}\n\n由小安助理為您整理"
    subprocess.run(['curl', '-s', '-X', 'POST',
        f'https://api.telegram.org/bot{BOT_TOKEN}/sendDocument',
        '-F', f'chat_id={CHAT_ID}',
        '-F', f'document=@{doc_path}',
        '-F', f'caption={caption}'],
        capture_output=True, timeout=30)

def main():
    today = datetime.datetime.now()

    print("📋 開始生成每日簡報...")

    # 取得各項資料
    print("🌤️ 取得天氣資訊...")
    weather = get_weather()

    print("📰 生成新聞摘要...")
    news = get_news_summary()

    print("📈 生成股市摘要...")
    market = get_market_summary()

    # 建立文件
    print("📄 建立簡報文件...")
    doc_path = create_briefing_doc(weather, news, market, today)

    # 發送
    print("📤 發送到 Telegram...")
    send_to_telegram(doc_path, today)

    print("✅ 每日簡報已完成！")

if __name__ == '__main__':
    main()
