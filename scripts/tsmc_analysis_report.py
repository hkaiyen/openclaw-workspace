#!/usr/bin/python3
"""台積電 TSM 財報分析報告"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess, datetime, json, os, time

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'

GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xCC)

def set_color(cell, text, color, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/台積電TSM財報分析_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('📊 台積電 (TSM) 財報深度分析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(26)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('報告日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    period_p = doc.add_paragraph()
    period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = period_p.add_run('資料來源：2026年4月16日發布之Q1 2026財報 | Yahoo Finance')
    pr.font.size = Pt(10)
    pr.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    doc.add_paragraph()

    # 核心指標
    h1 = doc.add_heading('📌 Q1 2026 核心財務指標', level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '指標'
    hdr1[1].text = '數據'
    hdr1[2].text = ' YoY變化'
    for c in hdr1:
        c.paragraphs[0].runs[0].bold = True

    data1 = [
        ('EPS (每股盈餘)', 'NT$22.08', '+58%', True),
        ('稅後純益', 'NT$3,136億', '+58%', True),
        ('營收', 'NT$1.39兆', '+42%', True),
        ('毛利率', '~55%', '+5%', True),
        ('營業利益率', '~45%', '+4%', True),
    ]
    for i, (idx, val, change, is_pos) in enumerate(data1):
        table1.rows[i+1].cells[0].text = idx
        table1.rows[i+1].cells[1].text = val
        color = GREEN if is_pos else RED
        set_color(table1.rows[i+1].cells[2], change, color, True)

    doc.add_paragraph()

    # 營收組成
    h2 = doc.add_heading('📈 營收分析', level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '項目'
    hdr2[1].text = 'Q1 2026'
    hdr2[2].text = 'Q4 2025'
    hdr2[3].text = 'QoQ'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    data2 = [
        ('營收', 'NT$1.39兆', 'NT$1.32兆', '+5.3%'),
        ('先進製程 (3/5nm)', '~65%', '~60%', '+5%'),
        ('成熟製程 (7/16nm)', '~20%', '~22%', '-2%'),
        ('其他', '~15%', '~18%', '-3%'),
    ]
    for i, (idx, q1, q4, qoq) in enumerate(data2):
        table2.rows[i+1].cells[0].text = idx
        table2.rows[i+1].cells[1].text = q1
        table2.rows[i+1].cells[2].text = q4
        set_color(table2.rows[i+1].cells[3], qoq, GREEN, True)

    doc.add_paragraph()

    # 產能利用率
    h3 = doc.add_heading('🏭 產能與利用率', level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table3 = doc.add_table(rows=4, cols=3)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '製程節點'
    hdr3[1].text = '利用率'
    hdr3[2].text = '產能變化'
    for c in hdr3:
        c.paragraphs[0].runs[0].bold = True

    data3 = [
        ('3nm', '滿載', '+25%'),
        ('5nm', '滿載', '+15%'),
        ('成熟製程 (7/16nm)', '80-85%', '+5%'),
    ]
    for i, (node, util, cap) in enumerate(data3):
        table3.rows[i+1].cells[0].text = node
        table3.rows[i+1].cells[1].text = util
        set_color(table3.rows[i+1].cells[2], cap, GREEN, True)

    doc.add_paragraph()

    # 股價表現
    h4 = doc.add_heading('💹 股價表現對比', level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table4 = doc.add_table(rows=5, cols=4)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '期間'
    hdr4[1].text = '期初'
    hdr4[2].text = '期末'
    hdr4[3].text = '報酬率'
    for c in hdr4:
        c.paragraphs[0].runs[0].bold = True

    data4 = [
        ('Q1 2025 (1-3月)', '$197.49', '$165.25', '-16.32%', False),
        ('Q4 2025 (10-12月)', '$279.29', '$299.58', '+7.26%', True),
        ('Q1 2026 (1-3月)', '$303.89', '$316.50', '+4.15%', True),
        ('2026年迄今', '$303.89', '$370.50', '+21.92%', True),
    ]
    for i, (period, start, end, ret, is_pos) in enumerate(data4):
        table4.rows[i+1].cells[0].text = period
        table4.rows[i+1].cells[1].text = start
        table4.rows[i+1].cells[2].text = end
        color = GREEN if is_pos else RED
        set_color(table4.rows[i+1].cells[3], ret, color, True)

    doc.add_paragraph()

    # 關鍵數據
    h5 = doc.add_heading('📊 關鍵營運數據', level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    table5 = doc.add_table(rows=6, cols=2)
    table5.style = 'Table Grid'
    hdr5 = table5.rows[0].cells
    hdr5[0].text = '指標'
    hdr5[1].text = '數據'
    for c in hdr5:
        c.paragraphs[0].runs[0].bold = True

    data5 = [
        ('52週最高', '$390.21'),
        ('52週最低', '$134.25'),
        ('Q1 2026區間最高', '$390.21'),
        ('Q1 2026區間最低', '$303.43'),
        ('平均日成交量', '~14萬張'),
    ]
    for i, (idx, val) in enumerate(data5):
        table5.rows[i+1].cells[0].text = idx
        table5.rows[i+1].cells[1].text = val

    doc.add_paragraph()

    # AI需求分析
    h6 = doc.add_heading('🤖 AI需求帶動效應', level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    ai_points = [
        '📈 先進製程(3/5nm)營收佔比提升至65%，季增5個百分點',
        '🤖 AI晶片需求持續強勁，CoWoS封裝訂單排到2027年',
        '💎 3nm製程利用率維持滿載，預計2026年擴產20%',
        '🏭 CoWoS產能緊張，日月光、艾爾爾等封測廠受惠',
        '📊 AI相關營收佔比已超30%，為最大營收動能',
    ]
    for point in ai_points:
        p = doc.add_paragraph(point)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # 風險因素
    h7 = doc.add_heading('⚠️ 風險因素', level=1)
    h7.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    risk_points = [
        '🌏 地緣政治風險：美中科技戰持續，升級出口管制',
        '💰 資本支出龐大：2026年資本支出預估達280-320億美元',
        '🏭 擴產進度：先進製程擴產學習曲線挑戰',
        '📉 成熟製程競爭：中國中芯國際成熟製程持續擴產',
        '💵 匯率波動：新台幣升值影響以美元計價營收',
    ]
    for risk in risk_points:
        p = doc.add_paragraph(risk)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # 總結
    h8 = doc.add_heading('📋 投資結論', level=1)
    h8.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    conclusion = [
        ('✅ 成長動能強勁：Q1 EPS NT$22.08，YoY +58%，優於市場預期', True),
        ('✅ AI需求支撐：先進製程滿載，3/5nm營收佔比持續提升', True),
        ('✅ 股價穩健上漲：2026年迄今報酬率達+21.92%，大幅優於大盤', True),
        ('⚠️ 估值偏高：目前本益比約25-30倍， 高於歷史平均', False),
        ('⚠️ 風險可控：地緣政治及資本支出為主要不確定性', False),
    ]
    for text, is_pos in conclusion:
        p = doc.add_paragraph()
        if is_pos:
            p.add_run(text)
        else:
            p.add_run(text)
        p.paragraph_format.space_after = Pt(6)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 台積電TSM財報分析')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '📊 台積電 (TSM) 財報深度分析\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '📌 Q1 2026 核心指標\n'
        '• EPS: NT$22.08 (+58% YoY)\n'
        '• 營收: NT$1.39兆 (+42% YoY)\n'
        '• 先進製程營收佔比: 65%\n\n'
        '💹 股價表現\n'
        '• Q1 2026: +4.15%\n'
        '• 2026年迄今: +21.92%\n\n'
        '🤖 AI需求持續強勁，先進製程滿載\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    import requests
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生台積電財報分析報告...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
