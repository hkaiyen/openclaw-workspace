#!/usr/bin/python3
"""
全方位新聞快報
全方位新聞快報 RSS
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import datetime
import json
import re
import urllib.request
import xml.etree.ElementTree as ET

# ========== 翻譯函數 ==========
def translate_to_chinese(text):
    """使用 MyMemory API 翻譯為中文"""
    if not text or len(text) < 2:
        return text
    try:
        encoded = urllib.parse.quote(text[:500])
        url = f"https://api.mymemory.translated.net/get?q={encoded}&langpair=en|zh-TW"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode())
            result = data.get("responseData", {}).get("translatedText", "")
            if result and result != text:
                return result
    except:
        pass
    return text

import urllib.parse

# ========== 新聞分類RSS ==========
LTN_CATEGORIES = {
    '🌏 國際': 'https://feeds.bbci.co.uk/news/world/rss.xml',
    '📈 財經': 'https://feeds.bbci.co.uk/news/business/rss.xml',
    '🏛️ 政治': 'https://tw.news.yahoo.com/rss/politics',
    '🏠 社會': 'https://tw.news.yahoo.com/rss/society',
    '⚽ 體育': 'https://tw.news.yahoo.com/rss/sports',
    '🎬 娛樂': 'https://tw.news.yahoo.com/rss/entertainment',
    '🏘️ 地方': 'https://tw.news.yahoo.com/rss/local',
    '🎨 生活': 'https://tw.news.yahoo.com/rss/life',
}

# ========== 工具函數 ==========

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def make_header_cell(cell, text, bg='FFFFFF'):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.font.size = Pt(10)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_data_cell(cell, text):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# ========== 抓取 RSS ==========

def fetch_category(category_name, url, count=5):
    """抓取單一分類的新聞"""
    news_list = []
    
    try:
        result = subprocess.run(
            ['curl', '-s', '-L', '--max-time', '15', url],
            capture_output=True, text=True, timeout=20
        )
        
        if result.stdout:
            from bs4 import BeautifulSoup
            from bs4 import XMLParsedAsHTMLWarning
            import warnings
            warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)
            
            soup = BeautifulSoup(result.stdout, 'html.parser')
            items = soup.find_all('item')
            
            for item in items[:count]:
                title = item.find('title')
                if title is not None and title.text:
                    # 處理 CDATA 的內容
                    title_text = title.get_text(strip=True)
                    if title_text.startswith('<![CDATA['):
                        title_text = title_text[9:]
                    if title_text.endswith(']]>'):
                        title_text = title_text[:-3]
                    title_text = re.sub(r'<[^>]+>', '', title_text)
                    if title_text and len(title_text) > 5:
                        # BBC 來源需要翻譯成中文
                        if 'bbci.co.uk' in url:
                            title_text = translate_to_chinese(title_text)
                        news_list.append(title_text)
    except:
        pass
    
    return news_list

def fetch_all_news():
    """抓取所有分類新聞"""
    all_news = {}
    
    for category, url in LTN_CATEGORIES.items():
        news = fetch_category(category, url, 5)
        all_news[category] = news
    
    return all_news

# ========== 主程式 ==========

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 標題 =====
    title = doc.add_heading('📰 全方位新聞快報', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.runs[0]
    tr.font.size = Pt(24)
    tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('多元新聞來源')
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    sr.font.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run('報告日期：' + today.strftime('%Y年%m月%d日 %H:%M'))
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== 抓取新聞 =====
    all_news = fetch_all_news()

    # ===== 各分類內容 ======
    section_num = 1
    
    for emoji, url in LTN_CATEGORIES.items():
        news_list = all_news.get(emoji, [])
        
        h = doc.add_heading(f'{section_num}、{emoji}', level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        
        if news_list:
            for news in news_list:
                p = doc.add_paragraph()
                p.add_run('• ' + news)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Cm(0.3)
        else:
            p = doc.add_paragraph()
            p.add_run('（今日無最新消息）')
        
        doc.add_paragraph()
        section_num += 1


    # 頁尾
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 小安製')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    output_path = f'/root/.openclaw/reports/daily/全方位新聞快報_{date_str}.docx'
    doc.save(output_path)
    return output_path

def send_to_telegram(file_path):
    bot_token = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
    chat_id = '8779713208'
    today = datetime.datetime.now()
    
    # 動態生成分類列表
    cats = ' · '.join([emoji for emoji, _ in LTN_CATEGORIES.items()])
    
    caption = f"📰 全方位新聞快報_{today.strftime('%Y年%m月%d日 %H:%M')}\n\n多元新聞來源\n\n{cats}\n\n小安製"

    result = subprocess.run(
        ['curl', '-s', '-X', 'POST',
         f'https://api.telegram.org/bot{bot_token}/sendDocument',
         '-F', f'chat_id={chat_id}',
         '-F', f'document=@{file_path}',
         '-F', f'caption={caption}'],
        capture_output=True, text=True, timeout=30
    )
    try:
        response = json.loads(result.stdout)
        return response.get('ok', False)
    except:
        return False

if __name__ == '__main__':
    print(f'[{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}] 開始生成新聞快報...')
    try:
        report_path = generate_report()
        print(f'📄 報告已生成: {report_path}')
        if send_to_telegram(report_path):
            print('✅ 已發送到 Telegram')
        else:
            print('❌ 發送失敗')
        print(f'[{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}] 任務完成！')
    except Exception as e:
        print(f'❌ 錯誤: {e}')
        import traceback
        traceback.print_exc()