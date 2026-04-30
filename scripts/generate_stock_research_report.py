#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
股票投資風險管理與獲利策略研究報告產生器
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_run_font(run, font_name='微軟正黑體', font_size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_with_style(doc, text, level=1, font_size=16, bold=True, color=(0, 51, 102)):
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold, color=color)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para

def add_normal_para(doc, text, font_size=12, bold=False, indent=False):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Cm(1)
    run = para.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def add_bullet(doc, text, font_size=12, indent_level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
    run = para.add_run(text)
    set_run_font(run, font_size=font_size)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def add_divider(doc):
    para = doc.add_paragraph()
    run = para.add_run('─' * 50)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_report():
    doc = Document()

    # 設定預設字體
    style = doc.styles['Normal']
    style.font.name = '微軟正黑體'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    # 設定段落間距
    doc.styles['Normal'].paragraph_format.space_after = Pt(6)
    doc.styles['Normal'].paragraph_format.space_before = Pt(3)

    # ===== 封面 =====

    # 標題
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(60)
    title_run = title_para.add_run('📈 股票投資風險管理與獲利策略 📉')
    set_run_font(title_run, font_size=24, bold=True, color=(0, 51, 102))

    # 副標題
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(20)
    sub_run = sub_para.add_run('深度研究報告')
    set_run_font(sub_run, font_size=16, bold=False, color=(80, 80, 80))

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(40)
    today = datetime.now().strftime('%Y年%m月%d日')
    date_run = date_para.add_run(today)
    set_run_font(date_run, font_size=14, color=(100, 100, 100))

    # 團隊
    team_para = doc.add_paragraph()
    team_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_para.paragraph_format.space_before = Pt(20)
    team_run = team_para.add_run('研究團隊：小歐｜千問｜拉瑪｜撈仔｜小咪｜小安')
    set_run_font(team_run, font_size=12, color=(100, 100, 100))

    doc.add_page_break()

    # ===== 目錄 =====

    add_heading_with_style(doc, '📋 研究大綱', level=1, font_size=18)

    toc_items = [
        '一、研究背景與目的',
        '二、風險管理的核心原則',
        '三、股票投資風險的類型',
        '四、常見獲利策略分析',
        '五、風險與報酬的平衡',
        '六、散戶實務操作建議',
        '七、心理風險管理',
        '八、綜合結論與建議'
    ]

    for item in toc_items:
        add_bullet(doc, item, font_size=13)

    doc.add_page_break()

    # ===== 研究背景 =====

    add_heading_with_style(doc, '一、研究背景與目的', level=1, font_size=16, color=(0, 51, 102))
    add_divider(doc)

    add_normal_para(doc,
        '股票投資是現代社會中最常見的資產配置方式之一，但同時也伴隨著顯著的風險。'
        '根據行為金融學研究，許多散戶投資人因為缺乏系統性的風險管理觀念，'
        '往往在市場波動中遭受不必要的損失。本報告旨在深入探討股票投資中的風險管理原則與獲利策略，'
        '幫助投資人建立穩健的投資紀律。')

    add_normal_para(doc,
        '本研究由五位助理（小歐/千問/拉瑪/撈仔/小咪）同步進行文獻研究與市場分析，'
        '整合國內外最新投資理論與實務經驗，最終由小安彙整形成結論報告。')

    doc.add_paragraph()

    # ===== 風險管理核心原則 =====

    add_heading_with_style(doc, '二、風險管理的核心原則', level=1, font_size=16, color=(0, 51, 102))
    add_divider(doc)

    add_heading_with_style(doc, '2.1 風險管理四大要件', level=2, font_size=14)

    add_normal_para(doc,
        '根據專業投資研究，風險管理應從以下四個要件著手：', font_size=12)

    add_bullet(doc, '【要件一】控管持股比率：根據市場趨勢動態調整整體持股比重，採取「金字塔型」操作，'
                '低檔時加重持股，高檔時獲利了結。單一持股建議不超過總部位的20%~30%。')

    add_bullet(doc, '【要件二】限定融資金額：當報酬率達到15%~20%時應實現獲利，同時嚴格控制融資比率，'
                '避免過度槓桿操作。持股市值上限不應過度擴增。')

    add_bullet(doc, '【要件三】嚴格執行停損：設定個人能容忍的虧損上限（一般建議20%為警戒線），'
                '果斷執行停損，避免虧損持續擴大。')

    add_bullet(doc, '【要件四】適時停利：當獲利回吐達到30%時應果斷停利，避免「來去一場空」。'
                '停利比停損更難，因為挑戰的是人性貪婪。')

    doc.add_paragraph()

    add_heading_with_style(doc, '2.2 分散投資原則', level=2, font_size=14)

    add_normal_para(doc,
        '「不要把所有的雞蛋放在同一個籃子裡」是投資最基本的原則。分散投資的重點包括：')

    add_bullet(doc, '跨資產類別：股票、債券、現金、黃金等多元化配置')
    add_bullet(doc, '跨產業配置：避免過度集中於單一產業，降低產業系統性風險')
    add_bullet(doc, '跨地區配置：適當配置不同市場，降低單一市場風險')
    add_bullet(doc, '持有10~20檔低相關性股票，可顯著降低投資組合風險')

    add_normal_para(doc,
        '根據橋水基金創辦人達里奧(Ray Dalio)的研究，持有15~20種零相關性資產，'
        '可以將風險報酬率提升至1.25，大幅降低輸錢機率。', bold=False)

    doc.add_paragraph()

    # ===== 股票投資風險類型 =====

    add_heading_with_style(doc, '三、股票投資風險的類型', level=1, font_size=16, color=(0, 51, 102))
    add_divider(doc)

    add_heading_with_style(doc, '3.1 系統性風險（不可分散風險）', level=2, font_size=14)

    add_bullet(doc, '【市場風險】整體股票市場下跌的風險，影響所有股票，包括藍籌股')
    add_bullet(doc, '【利率風險】市場利率變動對股票價值的影響，利率上升通常對股市不利')
    add_bullet(doc, '【通貨膨脹風險】物價上漲導致貨幣購買力下降，實質報酬被侵蝕')
    add_bullet(doc, '【政治/地緣政治風險】政府政策、戰爭、國際衝突對市場的衝擊')

    doc.add_paragraph()

    add_heading_with_style(doc, '3.2 非系統性風險（可分散風險）', level=2, font_size=14)

    add_bullet(doc, '【企業風險】個別公司經營不善、財務造假、管理問題等')
    add_bullet(doc, '【信用風險】公司無法按時償還債務，可能導致違約或破產')
    add_bullet(doc, '【流動性風險】股票难以在理想價位變現的風險，特別是小市值股票')
    add_bullet(doc, '【集中風險】過度集中於單一股票或產業的風險')

    doc.add_paragraph()

    add_heading_with_style(doc, '3.3 其他重要風險', level=2, font_size=14)

    add_bullet(doc, '【報酬順序風險】投資組合早期若遭遇負報酬，對長期財富影響甚鉅')
    add_bullet(doc, '【期限風險】因突發事件被迫提前變現長期投資，可能遭受損失')
    add_bullet(doc, '【從眾行為風險】盲目跟隨市場熱點或消息面，容易高買低賣')

    doc.add_paragraph()

    # ===== 常見獲利策略分析 =====

    add_heading_with_style(doc, '四、常見獲利策略分析', level=1, font_size=16, color=(0, 51, 102))
    add_divider(doc)

    add_heading_with_style(doc, '4.1 基本面分析策略', level=2, font_size=14)

    add_normal_para(doc,
        '基本面分析著眼於企業的內在價值，適合中長期投資。主要方法包括：')

    add_bullet(doc, '【價值投資】選擇被市場低估的股票，等待價值回歸。代表人物：巴菲特')
    add_bullet(doc, '【成長投資】選擇具有高成長潛力的公司，著眼未來營收與獲利增長')
    add_bullet(doc, '【財務指標】關注EPS、ROE、本益比、本淨比、負債比等關鍵數據')
    add_bullet(doc, '【由下而上法】從個別公司研究出發，深入了解企業經營體質')

    doc.add_paragraph()

    add_heading_with_style(doc, '4.2 技術分析策略', level=2, font_size=14)

    add_normal_para(doc,
        '技術分析透過觀察價格、成交量與市場趨勢來預測未來走勢，適合短線與波段操作。')

    add_bullet(doc, '【趨勢型指標】移動平均線(MA)、一目均衡表、布林通道(BOLL)——順勢追蹤')
    add_bullet(doc, '【震盪型指標】RSI、MACD——區間高拋低吸，單邊行情需謹慎')
    add_bullet(doc, '【型態學】頭肩頂/底、雙重頂(W型)、三角收斂、箱型整理等')
    add_bullet(doc, '【支撐/壓力線】找出價格關鍵位置，作為買賣決策參考')

    add_normal_para(doc,
        '重要提醒：技術分析是「歸納法」與「統計學」，歷史不會100%重演，'
        '不宜過度依賴單一指標，應搭配風險管理使用。')

    doc.add_paragraph()

    add_heading_with_style(doc, '4.3 籌碼面分析策略', level=2, font_size=14)

    add_normal_para(doc,
        '籌碼面分析觀察大戶、法人與大股東的動向，原理是他們通常比散戶掌握更多資訊。')

    add_bullet(doc, '【三大法人買賣超】追蹤外资、投信、自營商的買賣方向')
    add_bullet(doc, '【千張大戶持股】觀察主力大戶的持股變化')
    add_bullet(doc, '【分點進出】追蹤券商分點的買賣情況')
    add_bullet(doc, '【注意】收盤後才能看到當日籌碼，有滞后性，需配合其他分析')

    doc.add_paragraph()

    add_heading_with_style(doc, '4.4 資產配置策略', level=2, font_size=14)

    add_normal_para(doc,
        '四四三三法則是一種簡單有效的資產配置策略：')

    add_bullet(doc, '【40%股票】追求資本增值與成長潛力')
    add_bullet(doc, '【40%債券】提供穩定收益與市場下行的保護')
    add_bullet(doc, '【10%黃金】避險資產，對沖通膨與市場風險')
    add_bullet(doc, '【10%現金】保持流動性，可用於加碼或應急')

    add_normal_para(doc,
        '此配置可根據年齡、風險偏好與市場環境進行調整。年輕時可提高股票比重，'
        '年長後則應增加債券與現金比例。')

    doc.add_paragraph()

    # ===== 風險與報酬平衡 =====

    add_heading_with_style(doc, '五、風險與報酬的平衡', level=1, font_size=16, color=(0, 51, 102))
    add_divider(doc)

    add_heading_with_style(doc, '5.1 風險與報酬的基本關係', level=2, font_size=14)

    add_normal_para(doc,
        '投資的基本原則是：潛在報酬愈高，潛在風險通常也愈大。這適用於所有投資產品與市場情況。'
        '沒有任何投資策略可以同時確保最高報酬與最低風險，投資人必須在兩者之間找到適合自己的平衡點。')

    doc.add_paragraph()

    add_heading_with_style(doc, '5.2 風險承受能力評估', level=2, font_size=14)

    add_bullet(doc, '【風險能力】在不改變生活品質的前提下，能承受多少損失？')
    add_bullet(doc, '【風險偏好】面對虧損時的情緒反應，是保守還是進取？')
    add_bullet(doc, '【投資目標】是追求資本增值、穩定收益，還是保本？')
    add_bullet(doc, '【投資期限】資金閒置的時間越長，可承受的波動越大')

    doc.add_paragraph()

    add_heading_with_style(doc, '5.3 常見資產配置比例建議', level=2, font_size=14)

    add_normal_para(doc, '根據不同風險類型投資人的配置建議：')

    add_bullet(doc, '【積極型】股票50%、創業投資25%、加密貨幣15%、大宗商品10%')
    add_bullet(doc, '【平衡型】股票40%、債券30%、私募股權20%、現金10%')
    add_bullet(doc, '【保守型】黃金40%、債券40%、現金10%、防禦型股票10%')

    add_normal_para(doc,
        '重要原則：定期檢視投資組合，根據市場變化與個人狀況進行動態調整，'
        '而非一旦設定就永遠不變。')

    doc.add_paragraph()

    # ===== 散戶實務操作建議 =====

    add_heading_with_style(doc, '六、散戶實務操作建議', level=1, font_size=16, color=(0, 51, 102))
    add_divider(doc)

    add_heading_with_style(doc, '6.1 散戶常見錯誤認知', level=2, font_size=14)

    add_bullet(doc, '【錯誤1】把大師策略直接複製：巴菲特、张磊等人拥有普通散户无法获得的优势（资金规模、信息渠道、企业决策权等），邯郸学步反而容易亏损')

    add_bullet(doc, '【錯誤2】追求過高報酬：10萬元漲5倍與50萬元漲1倍的絕對收益相同，但風險程度完全不同。不追求過高收益，能讓風險敞口更小')

    add_bullet(doc, '【錯誤3】過度集中與頻繁交易：過度自信導致頻繁進出，增加交易成本的同時也降低了報酬率')

    add_bullet(doc, '【錯誤4】只看股價不看業績：過度關注短期股價波動，忽視企業基本面變化')

    doc.add_paragraph()

    add_heading_with_style(doc, '6.2 散戶八大投資原則', level=2, font_size=14)

    add_bullet(doc, '【原則1】只買自己能理解的股票：買股票就是買公司，要了解企業的產品、客戶與商業模式')

    add_bullet(doc, '【原則2】不追求過高的收益率：有时候慢就是快，安全性好的股票更能让人安心持有')

    add_bullet(doc, '【原則3】明白自己的能力邊際：清楚知道哪些股票可以操作，哪些與自己無關，不懂的堅決不做')

    add_bullet(doc, '【原則4】與國運緊密結合：選擇與國家經濟成長緊密相關的產業，如銀行、保險、消費等')

    add_bullet(doc, '【原則5】買股就要買龍頭：龍頭股確定性好、安全性高、波動率小，不容易被低估太久')

    add_bullet(doc, '【原則6】多關注業績，少關心股價：股價短期漲跌很難預測，專注企業價值才是根本')

    add_bullet(doc, '【原則7】長線才會少是非：時間是好公司的朋友，是爛公司的敵人')

    add_bullet(doc, '【原則8】建立能力圈同心圓：選定原點行業，逐步向外拓展，每個新領域至少觀察3年')

    doc.add_paragraph()

    add_heading_with_style(doc, '6.3 具體操作策略', level=2, font_size=14)

    add_bullet(doc, '【定期定額】忽視短期波動，紀律性地在固定時間投入固定金額，長期累積複利效果')
    add_bullet(doc, '【金字塔型操作】低檔時加重持股，高檔時逐步了結，避免倒金字塔型（高檔才加碼）')
    add_bullet(doc, '【平均成本法】將大額資金分批投入，避免一次性投入後遇到市場下跌')
    add_bullet(doc, '【設定止損點】預先設定止損價位，避免情緒化決策導致虧損擴大')

    doc.add_paragraph()

    # ===== 心理風險管理 =====

    add_heading_with_style(doc, '七、心理風險管理', level=1, font_size=16, color=(0, 51, 102))
    add_divider(doc)

    add_heading_with_style(doc, '7.1 常見行為金融偏誤', level=2, font_size=14)

    add_bullet(doc, '【損失厭惡】損失的痛苦感受大於等額盈利的快樂，導致「處分效果」：過早賣出獲利部位，卻遲遲不肯停損虧損部位')

    add_bullet(doc, '【過度自信】高估自己的判斷能力與資訊優勢，導致頻繁交易與過度集中')

    add_bullet(doc, '【錨定效應】過度看重買入價格或過去高點，以「回到成本價」做為賣出決策依據')

    add_bullet(doc, '【從眾行為】盲目跟隨市場熱點或群眾 decision，導致追高殺低')

    add_bullet(doc, '【近因效應】過度重視近期事件，忽視長期趨勢')

    doc.add_paragraph()

    add_heading_with_style(doc, '7.2 克服心理偏誤的方法', level=2, font_size=14)

    add_bullet(doc, '【寫投資日誌】記錄每次交易的決策過程與情緒，事後檢討哪些是被情緒左右')

    add_bullet(doc, '【制定明確紀律】在冷靜時預先訂定買賣原則，並嚴格執行，不因臨場情緒改變')

    add_bullet(doc, '【運用工具輔助】使用智慧下單系統、條件單、自動停利停損單，減少人為情緒干擾')

    add_bullet(doc, '【定期定額被動投資】減少主動猜測市場的頻率，用系統性紀律取代情緒化決策')

    add_bullet(doc, '【培養耐心】接受投資沒有捷徑的事實，用馬拉松心態而非百米衝刺心態')

    add_normal_para(doc,
        '股神巴菲特的師父班傑明·葛拉漢曾说：「投資者最大的問題，甚至是他最大的敵人，可能就是他自己。」'
        '了解自己的心理弱點，是改善投資決策的第一步。')

    doc.add_paragraph()

    # ===== 結論與建議 =====

    add_heading_with_style(doc, '八、綜合結論與建議', level=1, font_size=16, color=(0, 51, 102))
    add_divider(doc)

    add_heading_with_style(doc, '8.1 核心發現', level=2, font_size=14)

    add_bullet(doc, '風險管理是投資成功的關鍵，而非追求最高報酬')
    add_bullet(doc, '分散投資與資產配置是降低風險最有效的方式')
    add_bullet(doc, '基本面、技術面、籌碼面應相互搭配，避免偏廢其一')
    add_bullet(doc, '行為金融學的認知偏誤是散戶虧損的主因之一')
    add_bullet(doc, '長期投資與定期定額策略適合大多數散戶投資人')

    doc.add_paragraph()

    add_heading_with_style(doc, '8.2 行動建議', level=2, font_size=14)

    add_bullet(doc, '【立即行動】設定個人風險承受上限與停損點，嚴格遵守投資紀律')
    add_bullet(doc, '【短中期】建立多元化的投資組合，避免過度集中於單一股票或產業')
    add_bullet(doc, '【長期目標】培養自己的能力圈，選擇理解且信任的企業進行長期投資')
    add_bullet(doc, '【持續學習】關注行為金融學與風險管理的最新研究，不斷提升投資素養')

    doc.add_paragraph()

    add_heading_with_style(doc, '8.3 給不同投資人的建議', level=2, font_size=14)

    add_bullet(doc, '【新手投資人】從定期定額ETF或台灣50著手，先建立投資紀律再追求報酬')
    add_bullet(doc, '【有經驗投資人】強化基本面研究能力，建立自己的能力圈，減少對消息面的依賴')
    add_bullet(doc, '【高階投資人】可考慮多因子選股策略，並適當運用避險工具管理風險')

    doc.add_paragraph()

    add_divider(doc)

    # 署名
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_para.paragraph_format.space_before = Pt(30)
    end_run = end_para.add_run('本報告由小安彙整｜研究助理：小歐｜千問｜拉瑪｜撈仔｜小咪')
    set_run_font(end_run, font_size=11, color=(120, 120, 120))

    final_para = doc.add_paragraph()
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_run = final_para.add_run(f'生成日期：{today}')
    set_run_font(final_run, font_size=11, color=(120, 120, 120))

    return doc

if __name__ == '__main__':
    doc = create_report()
    output_path = '/root/.openclaw/workspace/reports/股票投資風險管理與獲利策略_20260428.docx'
    doc.save(output_path)
    print(f'✅ 報告已成功生成：{output_path}')