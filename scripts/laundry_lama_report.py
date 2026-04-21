#!/usr/bin/python3
"""拉瑪版：台北市文山區自助洗衣店創業市場研究報告"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

BOT_TOKEN = '8704642969:AAERVfjKsxcHExGOfZP9h5412w9Sp1TtABw'
CHAT_ID = '8779713208'
BLUE = RGBColor(0x1F, 0x49, 0x7D)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)
ORANGE = RGBColor(0xCC, 0x88, 0x00)

def set_color(cell, text, color, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = bold

def generate_report():
    today = datetime.datetime.now()
    date_str = today.strftime('%Y%m%d_%H%M')
    out_dir = '/root/.openclaw/reports/daily'
    os.makedirs(out_dir, exist_ok=True)
    output_path = out_dir + '/自主洗衣店創業研究_拉瑪版_' + date_str + '.docx'

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面
    title = doc.add_heading('🧺 台北市文山區自助洗衣店創業市場研究報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run('拉瑪專業市場分析版')
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run('研究日期：' + today.strftime('%Y年%m月%d日'))
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== 一、市場概況分析 =====
    h1 = doc.add_heading('一、市場概況分析', level=1)
    h1.runs[0].font.color.rgb = BLUE

    h1_1 = doc.add_heading('1.1 文山區基本資料', level=2)
    h1_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table1 = doc.add_table(rows=7, cols=2)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = '項目'
    hdr1[1].text = '數據'
    for c in hdr1:
        c.paragraphs[0].runs[0].bold = True

    data1 = [
        ('總人口', '252,897人（2026年3月）'),
        ('戶數', '110,852戶'),
        ('面積', '31.5090 平方公里'),
        ('人口密度', '8,026人/平方公里'),
        ('里別', '43里'),
        ('郵遞區號', '116'),
    ]
    for i, (項目, 數據) in enumerate(data1):
        table1.rows[i+1].cells[0].text = 項目
        table1.rows[i+1].cells[1].text = 數據

    doc.add_paragraph()

    h1_2 = doc.add_heading('1.2 主要大學進駐', level=2)
    h1_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    universities = [
        '• 國立政治大學（約25,000名師生）',
        '• 世新大學（約10,000名師生）',
        '• 中國科技大學（約6,000名師生）',
        '• 臺灣警察專科學校（約5,000名師生）',
    ]
    for u in universities:
        doc.add_paragraph(u)

    doc.add_paragraph()

    h1_3 = doc.add_heading('1.3 現有自助洗衣店數量與分佈', level=2)
    h1_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '分佈熱區'
    hdr2[1].text = '店家集中度'
    for c in hdr2:
        c.paragraphs[0].runs[0].bold = True

    table2.rows[1].cells[0].text = '景美商圈周邊（景美夜市、景美橋頭、溪口街）'
    table2.rows[1].cells[1].text = '高'
    table2.rows[2].cells[0].text = '萬隆/景美交界、羅斯福路五~六段沿線'
    table2.rows[2].cells[1].text = '中高'

    doc.add_paragraph()

    p_count = doc.add_paragraph()
    p_count.add_run('現有店家估計：').bold = True
    p_count.add_run('45-55家')

    doc.add_paragraph()

    h1_4 = doc.add_heading('1.4 市場競爭態勢', level=2)
    h1_4.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    competition = [
        '• 每店服務戶數約為2,000-2,500戶，已接近飽和',
        '• 晚上21:00後仍營業的店家不足',
        '• 具備舒適等候區、冷氣、免費Wi-Fi的「升級型」店家仍屬少數',
        '• APP預約機台、手機支付等數位化服務落後',
    ]
    for c in competition:
        doc.add_paragraph(c)

    doc.add_paragraph()

    # ===== 二、目標客群分析 =====
    h2 = doc.add_heading('二、目標客群分析', level=1)
    h2.runs[0].font.color.rgb = BLUE

    h2_1 = doc.add_heading('2.1 主要客層', level=2)
    h2_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table3 = doc.add_table(rows=5, cols=3)
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '客層'
    hdr3[1].text = '特徵'
    hdr3[2].text = '消費頻率'
    for c in hdr3:
        c.paragraphs[0].runs[0].bold = True

    data3 = [
        ('學生族群', '宿舍禁止晾衣，高度依賴自助洗衣', '每週1-2次'),
        ('租屋族', '小套房、無陽台住宅，洗衣需求剛性', '每週1-3次'),
        ('家庭族群', '床單、被套、制服等大物洗滌需求', '每月2-4次'),
        ('特定職業', '餐飲業/服務業制服、健身族運動服', '每週1-3次'),
    ]
    for i, (客層, 特徵, 頻率) in enumerate(data3):
        table3.rows[i+1].cells[0].text = 客層
        table3.rows[i+1].cells[1].text = 特徵
        table3.rows[i+1].cells[2].text = 頻率

    doc.add_paragraph()

    h2_2 = doc.add_heading('2.2 潛在客群規模估算', level=2)
    h2_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table4 = doc.add_table(rows=5, cols=3)
    table4.style = 'Table Grid'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '客層'
    hdr4[1].text = '估算人數'
    hdr4[2].text = '依據'
    for c in hdr4:
        c.paragraphs[0].runs[0].bold = True

    data4 = [
        ('大專院校學生', '約35,000人', '政治+世新+中科大+警專'),
        ('租屋族', '約38,000-50,000人', '252,897人×15-20%'),
        ('家庭洗滌需求者', '約60,000人', '110,852戶×55%×1.5人'),
        ('合計潛在客群', '約13-14萬人', ''),
    ]
    for i, (客層, 人數, 依據) in enumerate(data4):
        table4.rows[i+1].cells[0].text = 客層
        table4.rows[i+1].cells[1].text = 人數
        table4.rows[i+1].cells[2].text = 依據

    doc.add_paragraph()

    h2_3 = doc.add_heading('2.3 消費行為特徵', level=2)
    h2_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table5 = doc.add_table(rows=6, cols=2)
    table5.style = 'Table Grid'
    hdr5 = table5.rows[0].cells
    hdr5[0].text = '行為面向'
    hdr5[1].text = '觀察與分析'
    for c in hdr5:
        c.paragraphs[0].runs[0].bold = True

    data5 = [
        ('時段高峰', '19:00-22:00（下班下課後）為黃金時段'),
        ('等待容忍度', '學生族可接受30分鐘；上班族希望20分鐘以內'),
        ('支付偏好', '多元支付（悠遊卡、LINE Pay）接受度高'),
        ('季節性差異', '冬季與梅雨季為旺季；夏季為淡季'),
        ('平均客單', '洗衣$40-80/次，烘衣$30-50/次，平均$80-130'),
    ]
    for i, (面向, 分析) in enumerate(data5):
        table5.rows[i+1].cells[0].text = 面向
        table5.rows[i+1].cells[1].text = 分析

    doc.add_paragraph()

    # ===== 三、SWOT 分析 =====
    h3 = doc.add_heading('三、SWOT 分析', level=1)
    h3.runs[0].font.color.rgb = BLUE

    h3_1 = doc.add_heading('3.1 優勢（Strengths）', level=2)
    h3_1.runs[0].font.color.rgb = GREEN

    strengths = [
        ('S1 大學城聚落效應', '政治大學、世新大學等匯聚穩定學生客源'),
        ('S2 住宅密集且老舊', '30年以上公寓佔比高，許多住宅無陽台'),
        ('S3 租屋市場活絡', '文山區為台北市租屋熱區'),
        ('S4 營業時間彈性', '自助性質可24小時營業'),
        ('S5 邊際成本低', '無需大量人力，機台折舊後仍可持續運作'),
    ]
    for code, desc in strengths:
        p = doc.add_paragraph()
        p.add_run('• ' + code + '：').bold = True
        p.add_run(desc)

    doc.add_paragraph()

    h3_2 = doc.add_heading('3.2 劣勢（Weaknesses）', level=2)
    h3_2.runs[0].font.color.rgb = RED

    weaknesses = [
        ('W1 市場趨近飽和', '每2,000-2,500戶支撐一家'),
        ('W2 初期設備投資高', '洗衣機$80,000-150,000/台，乾燥機$50,000-80,000/台'),
        ('W3 水電成本沉重', '電費約佔營收25-35%'),
        ('W4 地點租金壓力', '景美、木柵成熟商圈店面租金$30,000-80,000/月'),
        ('W5 設備故障風險', '機台故障影響服務體驗'),
    ]
    for code, desc in weaknesses:
        p = doc.add_paragraph()
        p.add_run('• ' + code + '：').bold = True
        p.add_run(desc)

    doc.add_paragraph()

    h3_3 = doc.add_heading('3.3 機會（Opportunities）', level=2)
    h3_3.runs[0].font.color.rgb = BLUE

    opportunities = [
        ('O1 消費升級需求', '消費者願意多付費換取更好等候環境'),
        ('O2 數位化服務缺口', '多數店家缺乏APP預約、手機支付'),
        ('O3 永續環保趨勢', '節水機型、環保洗劑符合ESG趨勢'),
        ('O4 與外送平台合作', '部分店家已推出洗衣服務外送'),
    ]
    for code, desc in opportunities:
        p = doc.add_paragraph()
        p.add_run('• ' + code + '：').bold = True
        p.add_run(desc)

    doc.add_paragraph()

    h3_4 = doc.add_heading('3.4 威脅（Threats）', level=2)
    h3_4.runs[0].font.color.rgb = ORANGE

    threats = [
        ('T1 家用洗衣機普及', '部分取代自助店需求'),
        ('T2 連鎖品牌規模優勢', '集採機台、行銷資源，個體店難以抗衡'),
        ('T3 景氣波動影響消費', '經濟不佳時降低非必要性外出洗衣頻率'),
        ('T4 租金持續上漲', '房價租金上揚壓縮創業利潤空間'),
        ('T5 能源價格上漲', '電費、水費調漲直接壓縮毛利'),
    ]
    for code, desc in threats:
        p = doc.add_paragraph()
        p.add_run('• ' + code + '：').bold = True
        p.add_run(desc)

    doc.add_paragraph()

    h3_5 = doc.add_heading('3.5 SWOT 矩陣對策摘要', level=2)
    h3_5.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table6 = doc.add_table(rows=5, cols=2)
    table6.style = 'Table Grid'
    hdr6 = table6.rows[0].cells
    hdr6[0].text = '組合'
    hdr6[1].text = '策略建議'
    for c in hdr6:
        c.paragraphs[0].runs[0].bold = True

    data6 = [
        ('S+O（優勢×機會）', '打造「升級型」自助洗衣，主打數位預約＋舒適環境'),
        ('S+T（優勢×威脅）', '以大學周邊為核心區位，鎖定無洗衣機設置空間的租屋族'),
        ('W+O（劣勢×機會）', '以精緻裝潢與服務體驗差異化，避開價格競爭'),
        ('W+T（劣勢×威脅）', '選擇租金合理但人流充足的副都心位置，選用節能設備'),
    ]
    for i, (組合, 策略) in enumerate(data6):
        table6.rows[i+1].cells[0].text = 組合
        table6.rows[i+1].cells[1].text = 策略

    doc.add_paragraph()

    # ===== 四、區位分析建議 =====
    h4 = doc.add_heading('四、區位分析建議', level=1)
    h4.runs[0].font.color.rgb = BLUE

    h4_1 = doc.add_heading('4.1 建議地點與理由', level=2)
    h4_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    locations = [
        ('⭐ 首選區位A：景美商圈邊緣（羅斯福路六段沿線）', [
            '景美夜市每日人潮可觀，洗衣需求穩定',
            '羅斯福路為主要交通動線，機車族易於停留',
            '現有店家設備老舊，升級版新店具替代優勢',
            '租金相對市區核心店面合理',
        ]),
        ('⭐ 次選區位B：政治大學後山生活圈（指南路、忠勤路）', [
            '政大學生為最穩定剛性客群',
            '指南路沿線缺乏大型自助洗衣旗艦店',
            '學期期間需求穩定',
            '可與影印店、早餐店形成學生生活機能聚落',
        ]),
        ('⭐ 潛力區位C：興隆路/秀明路住宅區', [
            '住宅密度高，家庭客群潛力大',
            '現有自助洗衣店家數相對不足',
            '可鎖定大件衣物洗滌（被子、窗簾）需求',
            '社區型經營，建立口碑後客戶忠誠度高',
        ]),
    ]
    for loc_title, loc_points in locations:
        p = doc.add_paragraph()
        p.add_run(loc_title).bold = True
        for point in loc_points:
            doc.add_paragraph('  • ' + point)

    doc.add_paragraph()

    h4_2 = doc.add_heading('4.2 選址關鍵條件檢核表', level=2)
    h4_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table7 = doc.add_table(rows=9, cols=3)
    table7.style = 'Table Grid'
    hdr7 = table7.rows[0].cells
    hdr7[0].text = '條件'
    hdr7[1].text = '重要性'
    hdr7[2].text = '說明'
    for c in hdr7:
        c.paragraphs[0].runs[0].bold = True

    data7 = [
        ('臨路性（三角窗或雙向人流匯聚處）', '★★★★★', '自助店高度依賴路過隨機客群'),
        ('機車停車方便', '★★★★★', '文山區以機車為主要交通工具'),
        ('室內空間至少25-35坪', '★★★★☆', '容納6-10台洗衣機+4-6台乾燥機+等候區'),
        ('通風與排水條件佳', '★★★★★', '洗衣業務排水量大，排風不佳影響環境'),
        ('租金合理（$40,000以下為佳）', '★★★★☆', '租金佔成本比過高將壓縮利潤'),
        ('附近無直接競爭對手', '★★★★☆', '避免與現有強勢店家直接競爭'),
        ('可用執照（商業登記）', '★★★★★', '確認土地使用分區可做自助洗衣業態'),
        ('變電容量足夠', '★★★★☆', '多台大型洗衣機同時運轉，建議三相電'),
    ]
    for i, (條件, 重要性, 說明) in enumerate(data7):
        table7.rows[i+1].cells[0].text = 條件
        set_color(table7.rows[i+1].cells[1], 重要性, BLUE, True)
        table7.rows[i+1].cells[2].text = 說明

    doc.add_paragraph()

    # ===== 五、結論與建議 =====
    h5 = doc.add_heading('五、結論與建議', level=1)
    h5.runs[0].font.color.rgb = BLUE

    h5_1 = doc.add_heading('5.1 總結', level=2)
    h5_1.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    doc.add_paragraph('文山區自助洗衣市場呈現「整體穩定、局部缺口」態勢。252,897名居民與超過4萬名大學師生撐起扎實的剛性需求基礎，惟市場已有45-55家業者進駐，留給新進者的空間有限，但差異化定位與精準選址仍是可行的突破策略。')

    doc.add_paragraph()

    h5_2 = doc.add_heading('5.2 核心建議', level=2)
    h5_2.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    core_recs = [
        ('① 差異化定位', '打造「精品自助洗衣」：冷氣、Wi-Fi、舒適座椅、免費手機充電；導入APP預約；使用環保洗劑。預估每筆訂單可較傳統店提高30-50%'),
        ('② 首選區位', '政治大學後山生活圈——學生剛性需求最穩定，消費頻率最高，現有競爭最少'),
        ('③ 務實財務規劃', '單店設備投資預算：$200萬-280萬；目標月營收：$12萬-18萬；損益兩平：18-24個月'),
        ('④ 行銷策略', '初期：校園周邊發放優惠券；會員系統：儲值$500送$50；口碑操作：提供免費局部清洗吸引初次到店'),
    ]
    for title, desc in core_recs:
        p = doc.add_paragraph()
        p.add_run(title + '——').bold = True
        p.add_run(desc)

    doc.add_paragraph()

    h5_3 = doc.add_heading('5.3 最終評估', level=2)
    h5_3.runs[0].font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    table8 = doc.add_table(rows=6, cols=3)
    table8.style = 'Table Grid'
    hdr8 = table8.rows[0].cells
    hdr8[0].text = '評估項目'
    hdr8[1].text = '分數'
    hdr8[2].text = '備註'
    for c in hdr8:
        c.paragraphs[0].runs[0].bold = True

    data8 = [
        ('市場需求強度', '★★★★☆', '大學城+高租屋率支撐'),
        ('競爭激烈程度', '★★★★☆', '市場已趨飽和，需差異化'),
        ('進入障礙', '★★★☆☆', '設備投資適中，技術門檻低'),
        ('長期營運可行性', '★★★★☆', '剛性需求明確，現金流穩定'),
        ('總體推薦度', '★★★★☆', '值得進場，但選址與定位至關重要', True),
    ]
    for i, (項目, 分數, 備註, *extra) in enumerate(data8):
        table8.rows[i+1].cells[0].text = 項目
        set_color(table8.rows[i+1].cells[1], 分數, BLUE, True)
        table8.rows[i+1].cells[2].text = 備註

    doc.add_paragraph()
    doc.add_paragraph()

    # 聲明
    p_disclaimer = doc.add_paragraph()
    p_disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = p_disclaimer.add_run('*本報告由拉瑪專業市場分析師產生*\n*資料基準日：2026年4月*\n*報告內容僅供參考，實際投資前請自行進行更完整之盡職調查*')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()

    # 頁尾
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('小安智能助理｜老闆的專屬理財幫手  ◎ 拉瑪市場研究報告')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    return output_path

def send_telegram(file_path):
    today = datetime.datetime.now()
    caption = (
        '🧺 台北市文山區自助洗衣店創業市場研究報告\n'
        + today.strftime('%Y年%m月%d日') + '\n\n'
        '🐰 拉瑪專業市場分析版\n\n'
        '📊 市場概況：252,897人，現有45-55家店\n'
        '🎯 潛在客群：13-14萬人（學生+租屋族）\n'
        '📍 首選區位：景美/政大後山生活圈\n\n'
        '看完記得分享給需要的人！'
    )
    try:
        with open(file_path, 'rb') as f:
            files = {'document': f}
            data = {'chat_id': CHAT_ID, 'caption': caption}
            resp = requests.post(
                'https://api.telegram.org/bot' + BOT_TOKEN + '/sendDocument',
                data=data, files=files, timeout=30)
            return resp.json().get('ok', False)
    except:
        return False

if __name__ == '__main__':
    import requests
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 產生拉瑪版自助洗衣店創業市場研究報告...')
    report_path = generate_report()
    print('  📄 報告已生成: ' + report_path)
    print('  發送到Telegram...')
    if send_telegram(report_path):
        print('  ✅ 已發送到Telegram')
    else:
        print('  ❌ 發送失敗')
    print('[' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S') + '] 任務完成！')
