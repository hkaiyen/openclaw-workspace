#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    if size:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color

# 標題
title = doc.add_heading('', level=0)
run = title.add_run('遠距工作與自由工作者生活趨勢研究報告')
set_run_font(run, size=22, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副標題
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('「遠距工作與自由工作者生活趨勢」深度研究報告')
set_run_font(run2, size=14, color=RGBColor(0x4A, 0x4A, 0x4A))

# 日期
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_para.add_run('研究日期：2026年4月29日')
set_run_font(run3, size=11, color=RGBColor(0x80, 0x80, 0x80))

doc.add_paragraph()

# ========== 研究團隊 ==========
team_heading = doc.add_heading('', level=2)
team_run = team_heading.add_run('📋 研究團隊')
set_run_font(team_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

p = doc.add_paragraph()
run = p.add_run('本報告由以下五位助理同步研究後整合：')
set_run_font(run)

team_members = [
    ('👑 小咪（主管 / 第二層把關）', '研究統籌與品質把關'),
    ('📚 拉瑪（深度研究顧問）', '全球趨勢與數據分析'),
    ('🧠 千問（技術分析師）', '科技對工作型態影響分析'),
    ('💡 小歐（財報分析師）', '市場規模與產業分析'),
    ('🔍 撈仔（創意分析師）', '趨勢觀察與案例蒐集'),
]

for name, role in team_members:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{name}　—　{role}')
    set_run_font(run)

doc.add_paragraph()

# ========== 研究重點 ==========
heading1 = doc.add_heading('', level=2)
h1_run = heading1.add_run('📊 一、研究背景與市場規模')
set_run_font(h1_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

p = doc.add_paragraph()
run = p.add_run('全球遠距工作與自由工作者市場正經歷前所未有的成長：')
set_run_font(run)

stats = [
    ('全球自由工作者人數', '2025年約達1.57億人，相較傳統就業型態成長驚人'),
    ('美國自由工作者人數', '2020至2024年間成長90%，已超過6,400萬人'),
    ('自由平台市場規模', '2024年估值54億美元，2030年預計增至133億美元，年複合成長率達16.1%'),
    ('數位遊牧民族人數', '2020年約1,090萬人，2024年大幅成長至3,500萬人，超越台灣總人口'),
    ('遠距工作者滿意度', '遠端工作者滿意度高達79%，獨立工作者成長20%'),
    ('Gen Z 自由工作者', '61%認為獨立工作讓他們更能掌控職涯成長'),
]

for label, value in stats:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

doc.add_paragraph()

# ========== 遠距工作全球趨勢 ==========
heading2 = doc.add_heading('', level=2)
h2_run = heading2.add_run('🌍 二、遠距工作與自由工作者的全球發展趨勢')
set_run_font(h2_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

p = doc.add_paragraph()
run = p.add_run('後疫情時代，遠距工作已成為常態而非例外：')
set_run_font(run)

trends = [
    ('美國遠距工作者', '2025年約3,260萬人，占美國勞動力22%，雖較疫情高峰回落，但相較疫情前仍大幅上升'),
    ('混合辦公主流化', '全球83%員工偏好混合辦公模式，結合辦公室與遠端工作，平衡彈性與協作需求'),
    ('完全遠距企業', '僅16%企業採用完全遠距模式，44%仍要求員工全職進辦公室'),
    ('WEF 預測', '2030年前全球數位可遠距工作崗位將成長25%，達9,200萬人'),
    ('台灣排名', '全球數位遊牧適合度排名第12名，為亞洲第一'),
    ('科技業領導', '電腦及IT產業為2024年及2025年遠距工作最蓬勃的產業'),
]

for label, value in trends:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

doc.add_paragraph()

# ========== 科技進步影響 ==========
heading3 = doc.add_heading('', level=2)
h3_run = heading3.add_run('💻 三、科技進步對工作型態的影響')
set_run_font(h3_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

p = doc.add_paragraph()
run = p.add_run('科技發展正在深刻重塑工作型態：')
set_run_font(run)

tech_points = [
    ('AI工具普及', '74%自由工作者使用AI工具工作，61%表示AI為他們節省時間並提升產出'),
    ('AI取代效應', '2026年科技業裁員風暴加劇，Meta裁撤逾20%員工、Atlassian砍1,600職位，AI正在取代內容寫作、客戶服務、初階程式開發等遠端工作崗位'),
    ('新型AI職位崛起', '「AI工具整合與管理」成為新興需求，需要同時理解技術與業務的跨領域人才'),
    ('協作平台成熟', '視訊會議平台整合AI即時轉錄、語言翻譯、沉浸式3D會議環境'),
    ('零信任資安', '分散式團隊資料安全需求增加，零信任框架、端到端加密成為標準配備'),
    ('AI技能搜尋暴增', 'LinkedIn數據顯示2026年Q1「AI工具操作」相關技能搜尋量成長340%'),
]

for label, value in tech_points:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

doc.add_paragraph()

# ========== 自由工作者挑戰 ==========
heading4 = doc.add_heading('', level=2)
h4_run = heading4.add_run('⚠️ 四、自由工作者的挑戰與機遇')
set_run_font(h4_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

p = doc.add_paragraph()
run = p.add_run('自由工作者與遠距工作者面臨獨特的挑戰與機遇：')
set_run_font(run)

challenges_heading = doc.add_paragraph()
cr = challenges_heading.add_run('【挑戰】')
set_run_font(cr, bold=True)

challenges = [
    ('收入不穩定', '自由工作者收入可能面臨大幅波動，特別是在AI衝擊下部分職位消失'),
    ('可見度危機', '遠端員工存在感薄弱，裁員決策時易被忽略'),
    ('合約關係脆弱', '自由工作者不受勞動法規大量解僱保護，企業可隨時終止合作'),
    ('職業倦怠風險', '居家工作者個人生活與職業生活界線模糊，倦怠感容易急劇上升'),
    ('溝通鴻溝', '29%遠端工作者將溝通障礙列為首要難題，22%感到孤獨'),
    ('執行型職位萎縮', '內容寫作、資料處理等標準化任務正被AI快速取代'),
]

for label, value in challenges:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

opp_heading = doc.add_paragraph()
orun = opp_heading.add_run('【機遇】')
set_run_font(orun, bold=True)

opportunities = [
    ('工作生活平衡', '70%自由工作者認為彈性工作顯著改善生活品質，82%表示對心理健康有益'),
    ('技能多元發展', '67%自由工作者表示能接觸更多元的專案類型'),
    ('AI紅利', '能有效運用AI工具的工作者，生產力可達不會用AI者的3至5倍'),
    ('地理套利優勢', '在高收入市場賺錢、在低生活成本地區花費，月生活費可能只有紐約的三分之一'),
    ('多元收入可能', '可建立「核心服務 + 數位產品 + 被動收入」的多元收入結構'),
    ('數位遊牧簽證', '西班牙、荷蘭等國家提供數位遊牧簽證，部分甚至開放公民申請'),
]

for label, value in opportunities:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

doc.add_paragraph()

# ========== 生活品質與工作平衡 ==========
heading5 = doc.add_heading('', level=2)
h5_run = heading5.add_run('⚖️ 五、生活品質與工作平衡')
set_run_font(h5_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

p = doc.add_paragraph()
run = p.add_run('遠距與自由工作對生活品質的影響：')
set_run_font(run)

balance_points = [
    ('快樂指數', '82.9%自由工作者對整體工作生活感到快樂與滿足'),
    ('通勤時間節省', '遠距工作者平均每週節省8小時原本耗費在交通的時間'),
    ('金錢節省', '員工每年平均節省6,000美元交通、餐飲與服裝費用；雇主每名員工節省達11,000美元'),
    ('環境效益', '遠距工作每年可減少5,400萬噸溫室氣體排放'),
    ('自主權', '73.2%受訪者將「彈性」列為選擇自由工作的首要動機'),
    ('心理健康', '79%遠端專業人士表示壓力水平降低'),
]

for label, value in balance_points:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

doc.add_paragraph()

# ========== 企業適應策略 ==========
heading6 = doc.add_heading('', level=2)
h6_run = heading6.add_run('🏢 六、企業與個人適應策略')
set_run_font(h6_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

p = doc.add_paragraph()
run = p.add_run('面對工作型態變革，企業與個人都需要制定新的適應策略：')
set_run_font(run)

corp_heading = doc.add_paragraph()
crun = corp_heading.add_run('【企業端策略】')
set_run_font(crun, bold=True)

corp_strategies = [
    ('優先投資AI培訓', '加強AI技能培訓以因應人才市場變化'),
    ('心理健康支持', '提供心理健康資源、健身房會員、健康管理計劃等員工福祉方案'),
    ('混合辦公配套', '建置高品質視訊會議設備，確保遠端與現場員工會議公平性'),
    ('目標導向管理', '採用非同步溝通、明確目標設定與技術驅動問責機制'),
    ('培訓支出增加', '2025年培訓支出預計成長11.7%，領導力發展支出成長達13.3%'),
    ('員工敬業度', '76%受訪者表示員工敬業度正在改善，需持續優化相關措施'),
]

for label, value in corp_strategies:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

pers_heading = doc.add_paragraph()
prun = pers_heading.add_run('【個人端策略】')
set_run_font(prun, bold=True)

pers_strategies = [
    ('從執行者升級為整合者', '提供完整解決方案，而非單一技能執行'),
    ('精通AI工具', '學習至少3種以上主流AI工具，目標從「會用」到「精通」'),
    ('建立多元收入', '核心服務 + 數位產品（線上課程、電子書）+ 被動收入'),
    ('善用地理套利', '選擇數位遊牧簽證且稅務優惠國家，利用時區差異服務跨區客戶'),
    ('投資人際網絡', '參與線上社群、共享工作空間活動，擴大人脈安全網'),
    ('定期自我評估', '本週內評估目前工作內容有多少比例可被AI完成，若超過50%需立即轉型'),
]

for label, value in pers_strategies:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

doc.add_paragraph()

# ========== AI時代哪些職位安全 ==========
heading7 = doc.add_heading('', level=2)
h7_run = heading7.add_run('🔒 七、AI時代哪些遠端工作最安全')
set_run_font(h7_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

p = doc.add_paragraph()
run = p.add_run('根據研究分析，以下類型的工作在AI時代相對安全：')
set_run_font(run)

safe_jobs = [
    ('AI工具整合與管理者', '需同時理解技術與業務，AI本身無法取代'),
    ('策略顧問與高階分析', '需要綜合判斷力、產業經驗和人際洞察力'),
    ('創意指導與品牌策略', '能定義品牌靈魂，AI能執行但無法定義方向'),
    ('複雜系統架構與資深技術', '理解企業需求、評估技術風險、做出架構決策的高階技術能力'),
    ('人際關係密集型', '教練、心理諮商、銷售、社群經營等需要真實人際互動的工作'),
    ('數位遊牧者獨有優勢', '跨文化觀點、極強適應力、地理套利彈性、全球社群網絡'),
]

for label, value in safe_jobs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

doc.add_paragraph()

# ========== 結論 ==========
heading8 = doc.add_heading('', level=2)
h8_run = heading8.add_run('📝 八、研究結論與建議')
set_run_font(h8_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

conclusions = [
    ('趨勢確立', '遠距工作與自由工作者已成為不可逆的全球趨勢，企業必須適應這一新常態'),
    ('AI雙面刃', 'AI既帶來職位取代的威脅，也提供生產力提升的工具，關鍵在於如何有效運用'),
    ('門檻提高', '企業期待遠端員工具備主動發現問題、有效運用AI、跨時區協作等能力'),
    ('安全區塊', '策略型、整合型、人際關係密集型工作相對安全，執行型、標準化工作最容易被取代'),
    ('地理彈性', '數位遊牧者透過地理套利可在高收入市場賺錢、低成本地區生活，維持生活品質'),
    ('行動優先', '決定命運的不是你在哪裡工作，而是你做什麼樣的工作，以及你有多快的適應速度'),
]

for label, value in conclusions:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：{value}')
    set_run_font(run)

doc.add_paragraph()

# ========== 參考來源 ==========
ref_heading = doc.add_heading('', level=2)
ref_run = ref_heading.add_run('📚 參考來源')
set_run_font(ref_run, size=14, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

sources = [
    'Jobbers.io – Remote Work vs Freelancing: The 2025 Workforce Revolution',
    'Digital Nomad Press – AI裁員潮下的遠端工作者生存指南（2026年3月）',
    'World Economic Forum – Digital Jobs Report',
    'Cisco – Global Hybrid Work Study',
    'MBO Partners – 2024 Digital Nomad Report',
    'Global Citizen Solutions – Global Digital Nomad Report',
    'Forbes Advisor – Top Remote Work Statistics And Trends',
    'Neat – 2025 Remote Work Yearbook',
    'Blanchard – 2025年人力資源/學習發展趨勢調查報告',
    'Payoneer – The 2023 Freelancer Report',
    'FlexJobs – Remote Work Index: Trends & Statistics 2026',
]

for source in sources:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(source)
    set_run_font(run)

doc.add_paragraph()

# 底部署名
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
f_run = footer.add_run('🐰 小安研究團隊　出品')
set_run_font(f_run, size=10, color=RGBColor(0x80, 0x80, 0x80))

# 儲存
doc_path = '/root/.openclaw/workspace/reports/遠距工作與自由工作者生活趨勢_20260429.docx'
doc.save(doc_path)
print(f'報告已儲存至：{doc_path}')

# 更新pending狀態
pending_path = '/root/.openclaw/workspace/scripts/.pending_research.txt'
with open(pending_path, 'w') as f:
    f.write('done')
print(f'狀態已更新為 done')
